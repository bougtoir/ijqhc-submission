#!/usr/bin/env python3
"""
Create BMJ-format docx - Japanese version.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

def add_heading_bmj(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    return p

def add_table_row(table, cells, bold=False):
    row = table.add_row()
    for i, cell_text in enumerate(cells):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(cell_text))
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.bold = bold
    return row

def add_figure(path, caption, width=Inches(6)):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.italic = True
    else:
        add_para(f'[図: {path} が見つかりません]', italic=True)

# ═══════════════════════════════════════════════════════════
# 表紙
# ═══════════════════════════════════════════════════════════

title_text = (
    "日本の335二次医療圏における麻酔診療の地域差：\n"
    "全国レセプトデータを用いた横断的分析"
)
p = doc.add_paragraph()
run = p.add_run(title_text)
run.font.name = 'Arial'
run.font.size = Pt(16)
run.bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
add_para('[著者名]', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('[所属]', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('[責任著者：氏名、住所、メールアドレス]', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
add_para('本文語数：約4,000語（英語版）', italic=True)
add_para('表：2', italic=True)
add_para('図：9（2次元コロプレスマップ6枚 + 3次元押出マップ3枚）', italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 構造化抄録
# ═══════════════════════════════════════════════════════════

add_heading_bmj('抄録', level=1)

p = doc.add_paragraph()
run = p.add_run('目的 ')
run.bold = True
p.add_run(
    '日本の二次医療圏における麻酔診療の地域差を定量化し、観察された差異が保険査定の人為的産物か、'
    '真の臨床実践の差異を反映するものかを検証する。'
)

p = doc.add_paragraph()
run = p.add_run('研究デザイン ')
run.bold = True
p.add_run('公開された全国レセプトデータを用いた横断的生態学的研究。')

p = doc.add_paragraph()
run = p.add_run('対象 ')
run.bold = True
p.add_run('全47都道府県の335二次医療圏。')

p = doc.add_paragraph()
run = p.add_run('主要評価項目 ')
run.bold = True
p.add_run(
    '全身麻酔（L008）、硬膜外麻酔（L002）、持続硬膜外注入（L003）、脊椎麻酔（L004）の'
    '年齢・性別調整済み標準化レセプト出現比（SCR）。保険査定の差異が観察された地域差を'
    '説明するか否かを3つの感度分析で検証した。'
)

p = doc.add_paragraph()
run = p.add_run('結果 ')
run.bold = True
p.add_run(
    '全ての麻酔コードで顕著な地域差が認められた：変動係数（CV）は全身麻酔54.6%、硬膜外麻酔83.2%、'
    '持続硬膜外注入64.9%であった。3つの感度分析は、査定の差異が観察された地域差を説明するには'
    '不十分であることを示した：'
    '（1）L008の分散の85.8%は査定方針が統一的な同一県内で生じていた；'
    '（2）全身麻酔と脊椎麻酔は逆相関（r=\u22120.506）を示し、臨床的代替パターンと整合的であった'
    '（査定による再分類でも逆相関は生じうるが、査定率の規模を考慮するとその影響は限定的である）；'
    '（3）査定率差（0.2ポイント）はSCR変動の1%未満しか説明できない。'
    '大学病院所在の有無が最も強い予測因子であり、L008の全分散の38.5%を説明した（Cohen\'s d=1.78）。'
    'GA+硬膜外併用の指標であるL003は大学病院圏で1.73倍高かった。'
)

p = doc.add_paragraph()
run = p.add_run('結論 ')
run.bold = True
p.add_run(
    '日本における麻酔診療の地域差は構造的なものであり、保険査定の人為的産物ではない。'
    '大学病院の近接性が主要な決定因子である。区域麻酔併用技法と腫瘍学的転帰の関連を考慮すると、'
    'この地域差は修正可能な健康格差の源泉である可能性がある。'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 既知の知見 / 本研究の新知見
# ═══════════════════════════════════════════════════════════

add_heading_bmj('既に知られていること', level=2)
doc.add_paragraph(
    '手術・麻酔診療の地域差は多くの国で報告されているが、臨床的選好、施設要因、'
    '保険査定などの管理的メカニズムの相対的寄与は不明であった。',
    style='List Bullet'
)
doc.add_paragraph(
    '日本の皆保険制度は全てのレセプトを標準化された審査にかけるため、'
    '管理的変動と臨床的変動を区別する自然実験の場となる。',
    style='List Bullet'
)

add_heading_bmj('本研究の新知見', level=2)
doc.add_paragraph(
    '3つの独立した感度分析により、日本の麻酔診療の地域差は構造的であり、'
    '保険査定の産物ではないことが示された。',
    style='List Bullet'
)
doc.add_paragraph(
    '大学病院の有無（単一の二値変数）が全身麻酔量の分散の38.5%を説明し、'
    'この効果は全47都道府県で一貫していた。',
    style='List Bullet'
)
doc.add_paragraph(
    '腫瘍学的利益が示唆されるGA+硬膜外併用法（L003）が最大の地域間不平等（CV=64.9%）を示し、'
    '健康格差縮小の修正可能なターゲットとなりうる。',
    style='List Bullet'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 緒言
# ═══════════════════════════════════════════════════════════

add_heading_bmj('緒言', level=1)

doc.add_paragraph(
    '日本の国民皆保険制度は1961年に確立され、診療報酬点数表（しんりょうほうしゅうてんすうひょう）'
    'による統一的な報酬体系のもと、全国民が地理的条件によらず均質な医療へアクセスできることを'
    '前提としている。\u00b9\u00b7\u00b2 しかし、NDB（ナショナルデータベース）を用いた研究では、手術術式、処方パターン、'
    '検査オーダーなど多くの医療行為に顕著な地域差が報告されている。\u00b3\u2013\u2075'
)

doc.add_paragraph(
    '日本の制度の特徴的な点は、全レセプトに対する義務的な保険審査（しんさ）である。'
    '各都道府県の審査委員会がレセプトを審査し、不適切と判断された請求は減額または却下（査定）される。'
    '査定率は都道府県間で0.07%から0.28%まで変動する。\u2076 ここで根本的な問いが生じる：'
    'レセプトデータに観察される地域差は、臨床実践の真の差異を反映しているのか、'
    'それとも査定の差異による人為的産物なのか？'
)

doc.add_paragraph(
    'この問いは麻酔診療において特に重要である。麻酔方法の選択\u2014全身麻酔単独か、'
    '全身麻酔+区域麻酔の併用か\u2014は単なる好みの問題ではない。近年のエビデンスは、'
    '区域麻酔（特に硬膜外麻酔）と全身麻酔の併用が、がん手術における無再発生存期間と'
    '全生存期間を改善する可能性を示唆している。\u2077\u2013\u00b9\u00b2 '
    'Cochrane系統的レビューはがん再発に対する区域麻酔の限定的だが示唆的なエビデンスを報告し、\u2077 '
    'Sesslerらのランダム化比較試験は乳癌再発に有意差を認めなかったものの、\u2078 '
    '二次解析や観察研究のメタアナリシスは特定のがん種における利益の可能性を支持し続けている。\u2079\u2013\u00b9\u00b2 '
    '麻酔方法の地域差が査定の産物ではなく'
    '真の臨床的差異であれば、それは腫瘍学的転帰に影響しうる修正可能な健康格差を意味する。'
)

doc.add_paragraph(
    '令和8年（2026年）の診療報酬改定では、全身麻酔コードの名称が「マスク又は気管内挿管」から'
    '「声門上器具使用又は気管内挿管」に変更された。これは、声門上器具（SGA）の臨床普及に伴い、'
    '旧名称の「マスク」にSGAが含まれるか否かの解釈が審査委員間で分かれていた状況を反映している。'
    'このような制度的背景を踏まえ、地域差の本質を理解することは時宜を得ている。\u00b9\u00b3'
)

doc.add_paragraph(
    '本研究では、二次医療圏レベルの公開SCRデータを用いて：'
    '（1）335医療圏における麻酔診療の地域差を定量化し、'
    '（2）保険査定の差異が観察された地域差を説明するかを3つの独立した感度分析で検証し、'
    '（3）大学病院の近接性やペインクリニック活動を含む構造的決定因子を同定する。'
)

# ═══════════════════════════════════════════════════════════
# 方法
# ═══════════════════════════════════════════════════════════

add_heading_bmj('方法', level=1)

add_heading_bmj('研究デザインとデータソース', level=2)

doc.add_paragraph(
    '本研究は公開データを用いた横断的生態学的研究である。以下の3つのデータソースを使用した：'
)

doc.add_paragraph(
    '令和4年度（2022年度）の標準化レセプト出現比（SCR）：内閣府「地域差の見える化」'
    'イニシアティブとして公開。\u00b9\u2074 SCRは年齢・性別調整済みの期待値に対する出現比であり、'
    '全国平均を100とする。都道府県レベル（47都道府県）と二次医療圏レベル（335圏）のデータを使用した。',
    style='List Bullet'
)

doc.add_paragraph(
    '令和4年（2022年）医師・歯科医師・薬剤師統計（e-Stat）より、'
    '二次医療圏別の麻酔科医数。\u00b9\u2075',
    style='List Bullet'
)

doc.add_paragraph(
    '国土数値情報（A38-20）の二次医療圏境界データ、行政区域データ（N03）の都道府県境界'
    'および北方領土境界。\u00b9\u2076',
    style='List Bullet'
)

add_heading_bmj('分析対象の麻酔コード', level=2)

doc.add_paragraph('以下の診療報酬コードを分析対象とした：')

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
for i, text in enumerate(['コード', '術式名', '臨床的意義', '対象圏数']):
    hdr[i].paragraphs[0].add_run(text).bold = True
    hdr[i].paragraphs[0].runs[0].font.size = Pt(9)

codes_data = [
    ['L008', '閉鎖循環式全身麻酔', 'GA量の主要指標', '334'],
    ['L002', '硬膜外麻酔', '区域麻酔（単独又は併用）', '307'],
    ['L003', '硬膜外麻酔後における局所麻酔剤の持続的注入', 'GA+硬膜外併用の直接指標', '331'],
    ['L004', '脊椎麻酔', 'GA代替法', '334'],
    ['L009', '麻酔管理料（I）', '麻酔科専門医配置の代理指標', '314'],
    ['L100', '神経ブロック（入院）', 'ペインクリニック活動指標', '335'],
]
for row_data in codes_data:
    add_table_row(table, row_data)

doc.add_paragraph()
add_para('表1. 分析対象の麻酔関連診療行為コード。', italic=True)

add_heading_bmj('大学病院マッピング', level=2)

doc.add_paragraph(
    '81大学病院（国立44校、公立8校、私立29校）を所在地の市区町村に基づき64の二次医療圏に'
    'マッピングした。日本は「一県一医大」政策により全47都道府県に少なくとも1校の医学部があるため、'
    '同一県内での大学病院圏と非大学病院圏の比較が可能であり、これにより県レベルの査定方針を'
    '統制した分析が行える。'
)

add_heading_bmj('査定仮説に対する感度分析', level=2)

doc.add_paragraph(
    '「麻酔SCRの地域差は全て保険査定の差異で説明される」という帰無仮説を、'
    '3つの独立したアプローチで検証した：'
)

doc.add_paragraph(
    '検定1（県内分散）：査定方針が県内で統一的に適用されるならば、県内分散は無視できるはずである。'
    '一元配置分散分析により、全SCR分散を県間成分と県内成分に分解した。',
    style='List Bullet'
)

doc.add_paragraph(
    '検定2（コード間相関）：査定により全身麻酔が脊椎麻酔に再分類される（またはその逆の）場合、'
    'これらのコードは生態学的に負の相関を示しうる（ただし臨床的代替でも同様のパターンが生じる）。'
    '麻酔コード間のPearson相関を算出し、検定3の査定率の定量的推定と合わせて解釈した。',
    style='List Bullet'
)

doc.add_paragraph(
    '検定3（定量的インパクト）：公表されている審査統計（都道府県間の査定率差0.07-0.28%）を用いて、'
    '査定率差のSCRへの最大影響量を推定した。',
    style='List Bullet'
)

add_heading_bmj('統計解析', level=2)

doc.add_paragraph(
    '記述統計（平均、標準偏差、変動係数、パーセンタイル）、群間比較（Welchのt検定、Cohen\'s d）、\u00b9\u2077 '
    '分散分解（一元配置分散分析および階層的分解）、相関分析（Pearson r）を実施した。'
    'SCRの算出には間接法標準化を用いた。\u00b9\u2078 '
    '全ての解析はPython 3.12（pandas、scipy、geopandas）で実施した。'
    '3次元押出コロプレスマップはPlotly Mesh3dとDelaunay三角分割を用いて作成し、\u00b9\u2079 '
    '2次元地図はmatplotlibとgeopandasで作成した。'
)

add_heading_bmj('患者・市民参画', level=2)

doc.add_paragraph(
    '本研究は公開された集計データのみを使用しており、個人レベルの患者情報は含まれない。'
    '研究デザインおよび実施に患者の関与はない。'
)

# ═══════════════════════════════════════════════════════════
# 結果
# ═══════════════════════════════════════════════════════════

add_heading_bmj('結果', level=1)

add_heading_bmj('麻酔診療の地域差', level=2)

doc.add_paragraph(
    '全ての麻酔コードで顕著な地域差が認められた（表2）。全身麻酔（L008）のSCRは334圏で'
    '2.3から458.9の範囲にわたり（CV 54.6%）、硬膜外麻酔（L002、CV 83.2%）および'
    'GA+硬膜外併用の直接指標である持続硬膜外注入（L003、CV 64.9%）でさらに大きな変動を認めた。'
    '図1にL008 SCRの地理的分布を示す。西日本および大都市圏で高値を示す傾向が認められた。'
)

# Table 2
table2 = doc.add_table(rows=1, cols=7)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0].cells
for i, text in enumerate(['コード', '術式', 'P10', 'P50', 'P90', 'CV (%)', 'n']):
    hdr2[i].paragraphs[0].add_run(text).bold = True
    hdr2[i].paragraphs[0].runs[0].font.size = Pt(9)

table2_data = [
    ['L008', '全身麻酔', '32.3', '73.2', '131.5', '54.6', '334'],
    ['L002', '硬膜外麻酔', '7.6', '73.9', '184.0', '83.2', '307'],
    ['L003', '持続硬膜外注入', '19.0', '73.1', '150.0', '64.9', '331'],
    ['L004', '脊椎麻酔', '31.1', '84.3', '169.0', '56.3', '334'],
    ['L009', '麻酔管理料I', '32.5', '79.1', '149.1', '57.2', '314'],
    ['L100', '神経ブロック（入院）', '23.4', '72.2', '174.7', '72.1', '335'],
]
for row_data in table2_data:
    add_table_row(table2, row_data)

doc.add_paragraph()
add_para(
    '表2. 二次医療圏における麻酔関連診療行為の標準化レセプト出現比の分布（全国平均=100）。',
    italic=True
)

add_heading_bmj('感度分析1：県内分散', level=2)

doc.add_paragraph(
    '全身麻酔（L008）の全分散の85.8%は、査定方針が統一的に適用される同一県内で生じていた。'
    'この所見は他のコードでも一貫していた：硬膜外72.7%、脊椎麻酔58.0%、神経ブロック65.1%が'
    '県内変動であった。個別県内でのL008 SCRの変動は極めて大きく、東京都内（12圏）では'
    '61.4から458.9、熊本県内（10圏）では2.3から170.9であった。'
    '県レベルの査定方針ではL008の地域差の最大14.2%しか説明できない。'
)

add_heading_bmj('感度分析2：コード間相関', level=2)

doc.add_paragraph(
    '全身麻酔（L008）と脊椎麻酔（L004）は強い逆相関（r=\u22120.506、P<0.001）を示した。'
    'これは、GA量が多い地域では脊椎麻酔が少なく、その逆も成り立つという代替パターンを意味する。'
    'この逆相関は臨床的選択の差異と整合的である。原理的には査定による再分類（コード間の振替）や'
    '防衛的縮小請求（査定を予期した請求控え）でも逆相関は生じうる。しかし、都道府県間の最大査定率差が'
    '0.21ポイント（検定3）に過ぎないことを考慮すると、その規模の再分類では観察された相関を'
    '生じさせるには不十分である。最も簡潔な解釈は、この代替パターンが臨床的選好と地域の診療スタイルを'
    '反映しているというものである。'
)

doc.add_paragraph(
    '全身麻酔（L008）と硬膜外麻酔（L002）は正の相関（r=+0.319、P<0.001）を示した。'
    'これはGA量が多い地域で硬膜外も多いことを意味し、共通の供給要因（麻酔科の体制）の'
    '反映と解釈される。'
)

add_heading_bmj('感度分析3：査定の定量的インパクト', level=2)

doc.add_paragraph(
    '都道府県間の最大査定率差は0.21ポイント（0.07%\u20130.28%）であった。'
    'この差が全て単一の麻酔コードに集中したとしても、SCRシフトは約0.2ポイントに留まり、'
    'L008の四分位範囲62ポイントの0.3%未満、L002のレンジ198ポイントの0.1%未満である。'
    '査定率の変動は、観察された地域差を量的に説明する能力を持たない。'
)

add_heading_bmj('大学病院効果', level=2)

doc.add_paragraph(
    '大学病院の所在は麻酔診療パターンの最も強い単一予測因子であった。'
    '全身麻酔（L008）について、64の大学病院圏の平均SCRは130.2、270の非大学病院圏では67.7であった'
    '（Cohen\'s d=1.78、t=12.49）。特筆すべきことに、県内比較（査定方針を統制）において、'
    '大学病院圏のL008 SCRが非大学病院圏を上回ったのは47都道府県中47（100%）であり、'
    '県内平均差は+61.4ポイント（SD 33.7、t=12.49）であった。'
)

doc.add_paragraph(
    '階層的分散分解では、L008の分散の14.5%が県間差、38.5%が大学病院効果（県内）、'
    '47.0%が残差に帰属された。大学病院の有無という単一の二値変数だけで、'
    '日本全国の全身麻酔量の変動の約40%を説明した。'
)

doc.add_paragraph(
    '東京都では用量反応関係が認められた：SMA当たりの大学病院数0校（SCR 62\u2013110）→'
    '2校（83\u2013119）→3校（168.5）→5校（435.7）。'
)

add_heading_bmj('GA+硬膜外併用麻酔', level=2)

doc.add_paragraph(
    '持続硬膜外注入（L003）はCV=64.9%と主要麻酔コード中最大の変動を示した。'
    'L003 SCRは大学病院圏で1.73倍高く（126.4 vs 73.2、d=0.96、P<0.001）、'
    'L008と強い正の相関（r=0.753、P<0.001）を示した。'
    '硬膜外併用（L008+L002）に対する3つの査定仮説検定は全て査定仮説を棄却した：'
    '正のL008-L002相関（r=+0.319）、合算によるCV増大（50.2%→58.3%）、'
    '大学病院効果の持続（d=1.00）。'
)

add_heading_bmj('ペインクリニックのスピルオーバー効果', level=2)

doc.add_paragraph(
    'ペインクリニック活動（神経ブロックSCR、L100で代理）が手術麻酔での区域麻酔使用を'
    '予測するという仮説は、弱い支持のみを得た。神経ブロックSCRと区域麻酔総合指数の'
    '相関はr=0.153（P<0.01）であった。一方、神経ブロック活動はGA量とより強く相関し'
    '（r=0.307、P<0.001）、ペインクリニック活動は区域麻酔技法の特異的なスピルオーバーよりも、'
    '麻酔科全体の体制の反映と解釈される。'
)

# ═══════════════════════════════════════════════════════════
# 考察
# ═══════════════════════════════════════════════════════════

add_heading_bmj('考察', level=1)

add_heading_bmj('主要な知見', level=2)

doc.add_paragraph(
    '本研究は、日本の335二次医療圏にわたる麻酔診療の地域差が顕著かつ構造的であることを示した。'
    '3つの独立した感度分析は一致して、保険査定の差異が観察された地域差のごく一部しか説明できない'
    'ことを示す。支配的な決定因子は大学病院の近接性であり、それだけで全身麻酔量の全変動の約40%を'
    '説明する。この効果が査定方針の異なる47都道府県全てで認められたことは、供給側・施設側の'
    'メカニズムを強く示唆する。'
)

add_heading_bmj('既存文献との比較', level=2)

doc.add_paragraph(
    '本研究の所見は、医療行為の地域差に関する国際的な文献と整合的である。'
    '米国のDartmouth Atlasプロジェクトは、手術率の地域差が患者のニーズよりも医師供給と'
    '診療スタイルにより駆動されることを報告してきた。\u00b2\u2070\u00b7\u00b2\u00b9 同様のパターンは英国、\u00b2\u00b2 ドイツ、\u00b2\u00b3 '
    'オーストラリアでも報告されている。\u00b2\u2074 本研究はこの文献を2つの点で拡張する。'
    '第一に、日本固有の制度的構造（統一的な報酬体系、都道府県別の審査委員会）を利用して、'
    '管理的要因と臨床的要因を区別した。第二に、東アジアの皆保険制度における麻酔技法の'
    '地域差を、335圏という精細な地理的スケールで初めて体系的に分析した。'
)

doc.add_paragraph(
    '本研究で報告した大学病院効果（全身麻酔のd=1.78）は、医療行為の地域差研究で通常報告される'
    '効果量と比較して極めて大きい。これは、大学医局（いきょく）が地理的影響圏内の関連病院の'
    '臨床実践に大きな影響力を持つ日本の医学教育制度を反映していると考えられる。\u00b2\u2075\u00b7\u00b2\u2076'
)

add_heading_bmj('腫瘍学的転帰への含意', level=2)

doc.add_paragraph(
    'GA+硬膜外併用法（L003）のCV=64.9%という所見は、区域麻酔と全身麻酔の併用ががん手術における'
    '無再発生存期間を改善する可能性を示すエビデンスに照らして、\u2077\u2013\u00b9\u00b2 臨床的に重要である。'
    '大学病院圏と非大学病院圏のL003 SCRの1.73倍の差は、潜在的に有益なこの技法へのアクセスが'
    '患者の居住地に大きく依存することを意味し、皆保険制度の公平性の前提に疑問を投げかける。\u00b9\u00b7\u00b2'
)

add_heading_bmj('令和8年診療報酬改定', level=2)

doc.add_paragraph(
    '全身麻酔コードの「マスク又は気管内挿管」から「声門上器具使用又は気管内挿管」への名称変更は、'
    '旧名称のコーディング上の曖昧さを反映している。査定による地域差が量的に小さいという本研究の'
    '所見は、この改定がコーディングの明確化としては適切であるものの、我々が記録した地理的パターンを'
    '大きく変える可能性は低いことを示唆する。'
)

add_heading_bmj('強みと限界', level=2)

doc.add_paragraph(
    '強みとして、全国民をカバーする年齢・性別調整済み指標の使用、335圏という精細な地理的スケール、'
    '複数の独立した感度分析、査定方針を統制した県内比較が挙げられる。'
)

doc.add_paragraph(
    '限界として、地域レベル分析に固有の生態学的誤謬、\u00b2\u2077 個人レベルの意思決定と集計パターンの区別不能、'
    'SCRデータの査定後的性質（ただし感度分析で量的影響は小さいと示した）、因果推論を許容しない'
    '横断デザイン、大学病院効果と都市集中効果の不完全な分離、コード別査定率データの欠如が挙げられる。'
    'また、「防衛的縮小請求」\u2014査定を予期して臨床医が予防的に請求を控える行動\u2014は'
    'レセプトデータからは定量化できない。\u00b2\u2078'
)

add_heading_bmj('結論と政策的含意', level=2)

doc.add_paragraph(
    '日本における麻酔診療の地域差は真実であり大きく、主として施設的要因\u2014特に大学病院の近接性'
    '\u2014により駆動されており、保険査定の差異によるものではない。'
    'GA+硬膜外併用法の大学病院圏と非大学病院圏の1.73倍の差は、腫瘍学的利益のエビデンスを'
    '考慮すると特に懸念される。\u2077\u2013\u00b9\u00b2 政策的対応として、教育プログラムの拡充、大学センターからの'
    '専門家派遣、質指標としての麻酔技法パターンの監視が考えられる。\u00b2\u2079 '
    '本研究で開発した感度分析フレームワーク\u2014県内分散分解、コード間相関、査定の定量的'
    'インパクト推定\u2014は、日本の皆保険制度下のあらゆる診療行為の地域差研究に適用可能である。'
)

# ═══════════════════════════════════════════════════════════
# 文献
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_bmj('文献', level=1)

refs = [
    # Introduction: Japan's universal health insurance
    '1. Ikegami N, Yoo B-K, Hashimoto H, et al. Japanese universal health coverage: evolution, '
    'achievements, and challenges. Lancet 2011;378:1106\u201315.',
    '2. Ikegami N. Universal Health Coverage for Inclusive and Sustainable Development: Lessons from '
    'Japan. Washington, DC: World Bank, 2014.',
    # Introduction: NDB and regional variation
    '3. Matsuda S, Fujimori K. The claim database in Japan. Asian Pac J Dis Manag 2012;6:55\u201359.',
    '4. Suto M, Iba A, Sugiyama T, et al. Literature review of studies using the National Database '
    'of Health Insurance Claims of Japan (NDB): limitations and strategies in using the NDB for '
    'research. JMA J 2023;7:10\u201320.',
    '5. Hamada H, Sekimoto M, Imanaka Y. Effects of the per diem prospective payment system with '
    'DRG-like grouping system (DPC/PDPS) on resource usage and healthcare quality in Japan. Health '
    'Policy 2012;107:194\u2013201.',
    # Introduction: audit rates
    '6. \u793e\u4f1a\u4fdd\u967a\u8a3a\u7642\u5831\u916c\u652f\u6255\u57fa\u91d1. \u5be9\u67fb\u7d71\u8a08, \u4ee4\u548c4\u5e74\u5ea6. '
    'https://www.ssk.or.jp/tokeijoho/shinsatokei/index.html (2026\u5e743\u670815\u65e5\u30a2\u30af\u30bb\u30b9).',
    # Introduction: cancer and regional anaesthesia
    '7. Pei L, Tan G, Wang L, et al. Anaesthetic techniques for risk of malignant tumour recurrence. '
    'Cochrane Database Syst Rev 2019;(11):CD008877.',
    '8. Sessler DI, Pei L, Huang Y, et al. Recurrence of breast cancer after regional or general '
    'anaesthesia: a randomised controlled trial. Lancet 2019;394:1807\u201315.',
    '9. Exadaktylos AK, Buggy DJ, Moriarty DC, et al. Can anesthetic technique for primary breast '
    'cancer surgery affect recurrence or metastasis? Anesthesiology 2006;105:660\u20134.',
    '10. Sun Y, Li T, Gan TJ. The effects of perioperative regional anesthesia and analgesia on cancer '
    'recurrence and survival after oncology surgery: a systematic review and meta-analysis. Reg '
    'Anesth Pain Med 2015;40:589\u201398.',
    '11. Xie S, Li L, Meng F, et al. Regional anesthesia might reduce recurrence and metastasis rates '
    'in adult patients with cancers after surgery: a meta-analysis. BMC Anesthesiol 2024;24:19.',
    '12. Zhang D, Jiang J, Liu J, et al. Effects of perioperative epidural analgesia on cancer '
    'recurrence and survival. Front Oncol 2021;11:798435.',
    # Introduction: 2026 fee schedule revision
    '13. \u539a\u751f\u52b4\u50cd\u7701. \u4ee4\u548c8\u5e74\u5ea6\u8a3a\u7642\u5831\u916c\u6539\u5b9a. https://www.mhlw.go.jp (2026\u5e743\u670815\u65e5\u30a2\u30af\u30bb\u30b9).',
    # Methods: data sources
    '14. \u5185\u95a3\u5e9c. \u5730\u57df\u5dee\u306e\u898b\u3048\u308b\u5316. '
    'https://www5.cao.go.jp/keizai-shimon/kaigi/special/reform/mieruka/ (2026\u5e743\u670810\u65e5\u30a2\u30af\u30bb\u30b9).',
    '15. \u539a\u751f\u52b4\u50cd\u7701. \u4ee4\u548c4\u5e74\u533b\u5e2b\u30fb\u6b6f\u79d1\u533b\u5e2b\u30fb\u85ac\u5264\u5e2b\u7d71\u8a08. e-Stat. '
    'https://www.e-stat.go.jp (2026\u5e743\u670812\u65e5\u30a2\u30af\u30bb\u30b9).',
    '16. \u56fd\u571f\u4ea4\u901a\u7701. \u56fd\u571f\u6570\u5024\u60c5\u5831\u30c0\u30a6\u30f3\u30ed\u30fc\u30c9\u30b5\u30fc\u30d3\u30b9. '
    'https://nlftp.mlit.go.jp (2026\u5e743\u670810\u65e5\u30a2\u30af\u30bb\u30b9).',
    # Methods: statistical methods
    '17. Cohen J. Statistical Power Analysis for the Behavioral Sciences. 2nd ed. '
    'Hillsdale, NJ: Lawrence Erlbaum, 1988.',
    '18. Julious SA, Nicholl J, George S. Why do we continue to use standardized mortality ratios for '
    'small area comparisons? J Public Health Med 2001;23:40\u20136.',
    '19. Plotly Technologies Inc. Plotly: collaborative data science. Montreal, QC, 2015. '
    'https://plotly.com (2026\u5e743\u670815\u65e5\u30a2\u30af\u30bb\u30b9).',
    # Discussion: Dartmouth Atlas and international comparisons
    "20. Wennberg JE. Tracking Medicine: A Researcher's Quest to Understand Health Care. "
    'New York: Oxford University Press, 2010.',
    '21. Wennberg JE, Fisher ES, Skinner JS. Geography and the debate over Medicare reform. '
    'Health Aff (Millwood) 2002;Suppl Web Exclusives:W96\u2013114.',
    '22. NHS RightCare. The NHS Atlas of Variation in Healthcare, 2015. London: Public Health England.',
    '23. Wengler A, Nimptsch U, Mansky T. Hip and knee replacement surgery (arthroplasty) and '
    'geographic variations in Germany. Dtsch Arztebl Int 2014;111:407\u201316.',
    '24. Australian Commission on Safety and Quality in Health Care. Australian Atlas of Healthcare '
    'Variation. Sydney: ACSQHC, 2015.',
    # Discussion: ikyoku system
    '25. Otsuka T. The ikyoku system of university orthopedic surgery departments: an in-hospital '
    'organizational system unique to Japan. J Orthop Sci 2012;17:513\u201314.',
    '26. Onishi H. History of Japanese medical education. Korean J Med Educ 2018;30:283\u201394.',
    # Discussion: ecological fallacy
    '27. Greenland S. Ecologic versus individual-level sources of bias in ecologic estimates of '
    'contextual health effects. Int J Epidemiol 2001;30:1343\u201350.',
    # Discussion: defensive medicine / undercoding
    '28. Sekimoto M, Imanaka Y, Kobayashi H, et al. Impact of hospital characteristics on the cost '
    'and practice patterns of acute myocardial infarction patients in Japan. Health Policy '
    '2006;78:34\u201345.',
    # Conclusions: quality monitoring
    '29. Merry AF, Cooper JB, Soyannwo O, et al. International standards for a safe practice of '
    'anesthesia 2010. Can J Anaesth 2010;57:1027\u201334.',
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.name = 'Arial'

# ═══════════════════════════════════════════════════════════
# 図
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_bmj('図', level=1)

figure_list = [
    ('/home/ubuntu/map_L008_scr.png',
     '図1. 二次医療圏別の全身麻酔（L008）標準化レセプト出現比。'
     '実線は都道府県境界、破線は二次医療圏境界を示す。北方領土は白色（医療圏未設定）。'),
    ('/home/ubuntu/map_univ_presence.png',
     '図2. 二次医療圏別の大学病院所在状況。赤丸は1校以上の大学病院を含むSMA（335圏中64圏）。'),
    ('/home/ubuntu/map_L004_scr.png',
     '図3. 二次医療圏別の脊椎麻酔（L004）標準化レセプト出現比。'),
    ('/home/ubuntu/map_L008_L004_combined.png',
     '図4. 全身麻酔+脊椎麻酔合計（L008+L004）SCR。合算により査定による再分類の影響を中和。'),
    ('/home/ubuntu/map_L003_scr.png',
     '図5. 持続硬膜外注入（L003）SCR\u2014GA+硬膜外併用の直接指標。CV=64.9%。'),
    ('/home/ubuntu/map_L003_L008_ratio_corrected.png',
     '図6. L003/L008比（GA1件あたりの硬膜外併用率）。手術件数で補正済み。'),
]

for path, caption in figure_list:
    add_figure(path, caption, width=Inches(5.5))
    doc.add_paragraph()

threed_figures = [
    ('/home/ubuntu/3d_extruded/3D_ratio_by_anes_v2_jp.png',
     '図7. 3次元押出マップ：色=L003/L008比（硬膜外併用率）、高さ=麻酔科医数。'),
    ('/home/ubuntu/3d_extruded/3D_L008_by_anes_v2_jp.png',
     '図8. 3次元押出マップ：色=L008 SCR（全身麻酔量）、高さ=麻酔科医数。'
     '赤く高い領域は大都市圏の大学病院所在圏。'),
    ('/home/ubuntu/3d_extruded/3D_ratio_by_surgery_v2_jp.png',
     '図9. 3次元押出マップ：色=L003/L008比、高さ=手術あたりGA率。'
     '高さの変動は小さく、手術あたりGA率は比較的均一。色の変動が硬膜外併用率の差を表現。'),
]

for path, caption in threed_figures:
    add_figure(path, caption, width=Inches(5.5))
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# 宣言
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_bmj('宣言', level=1)

add_heading_bmj('倫理審査', level=2)
doc.add_paragraph(
    '本研究は公開された集計データのみを使用し、個人レベルのデータにはアクセスしていない。'
    '「人を対象とする生命科学・医学系研究に関する倫理指針」（2021年改正）に基づき、'
    '倫理委員会の承認は不要と判断した。'
)

add_heading_bmj('データ入手可能性', level=2)
doc.add_paragraph(
    '本研究で使用した全データは公開されている。'
    'SCRデータ：内閣府「地域差の見える化」、医師統計：e-Stat、'
    'GIS境界データ：国土数値情報。解析コードは責任著者に請求可能。'
)

add_heading_bmj('利益相反', level=2)
doc.add_paragraph('[著者が記入]')

add_heading_bmj('資金', level=2)
doc.add_paragraph('[著者が記入]')

add_heading_bmj('著者の貢献', level=2)
doc.add_paragraph('[著者が記入。BMJは詳細な貢献声明を要求する。]')

add_heading_bmj('透明性宣言', level=2)
doc.add_paragraph(
    '筆頭著者（本論文の保証人）は、本論文が報告された研究の正直、正確、透明な記述であること、'
    '研究の重要な側面が省略されていないこと、計画された研究からの逸脱があれば全て説明されている'
    'ことを確認する。'
)

# ═══════════════════════════════════════════════════════════
# 保存
# ═══════════════════════════════════════════════════════════

output_path = '/home/ubuntu/regional_anaesthesia_BMJ_JP.docx'
doc.save(output_path)
print(f"Saved: {output_path}")
