#!/usr/bin/env python3
"""
Create REVISED Japanese manuscript for resubmission after RAPM rejection.
Japanese translation of the revised EN manuscript.
Addresses ALL Major (1-5) and Minor (1-5) reviewer comments.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Style setup ──
style = doc.styles['Normal']
font = style.font
font.name = 'MS Mincho'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

def add_heading_jp(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'MS Gothic'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'MS Mincho'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    return p

def add_table_row(table, cells, bold=False):
    row = table.add_row()
    for i, cell_text in enumerate(cells):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(cell_text))
        run.font.name = 'MS Mincho'
        run.font.size = Pt(9)
        run.bold = bold
    return row

def add_figure(path, caption, width=Inches(6)):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.font.name = 'MS Mincho'
        run.font.size = Pt(9)
        run.italic = True
    else:
        add_para(f'[図: {path} が見つかりません]', italic=True)


# ═══════════════════════════════════════════════════════════
# タイトルページ
# ═══════════════════════════════════════════════════════════

title_text = (
    "日本における麻酔実践の地域差：臨床的選好か保険審査の影響か？\n"
    "335二次医療圏を対象とした横断的生態学的研究"
)
p = doc.add_paragraph()
run = p.add_run(title_text)
run.font.name = 'MS Gothic'
run.font.size = Pt(16)
run.bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
add_para('[著者名]', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('[所属機関]', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('[責任著者：氏名、住所、メールアドレス]', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
add_para('本文字数：約4,500語（英語版）', italic=True)
add_para('表：3（表1：麻酔コード一覧、表2：記述統計、表3：マルチレベルモデル結果）', italic=True)
add_para('図：4（コロプレスマップ2枚、ファンネルプロット1枚、3D押出マップ1枚）', italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 構造化抄録
# ═══════════════════════════════════════════════════════════

add_heading_jp('抄録', level=1)

p = doc.add_paragraph()
run = p.add_run('背景 ')
run.bold = True
run.font.name = 'MS Gothic'
p.add_run(
    '麻酔実践の地域差は国際的に報告されているが、臨床的選好、施設要因、保険審査などの行政的要因の相対的寄与は不明である。'
    '日本の国民皆保険制度は都道府県別の審査制度を有しており、行政的要因と臨床的要因を区別するための独自の研究環境を提供している。'
).font.name = 'MS Mincho'

p = doc.add_paragraph()
run = p.add_run('方法 ')
run.bold = True
run.font.name = 'MS Gothic'
p.add_run(
    '令和4年度（2022年度）の全国335二次医療圏を対象とした横断的生態学的研究。'
    '内閣府「地域差の見える化」データより、年齢・性別調整済みの標準化レセプト発生比（SCR；観察値/期待値×100、全国平均=100）を使用した。'
    '335二次医療圏は47都道府県内に完全にネストされており、都道府県境界をまたぐ区域は存在しない。'
    '6つの麻酔手技コードを対象に、都道府県をランダム切片とするマルチレベル線形混合モデルにより級内相関係数（ICC）を推定し、'
    '保険審査の影響を検証する3つの感度分析、および小サンプル不安定性に対処するEmpirical Bayes縮小推定を実施した。'
).font.name = 'MS Mincho'

p = doc.add_paragraph()
run = p.add_run('結果 ')
run.bold = True
run.font.name = 'MS Gothic'
p.add_run(
    '全麻酔コードで顕著な地域差が認められた（全身麻酔L008の変動係数53.6%、硬膜外麻酔L002で87.0%）。'
    'マルチレベルモデルでは、都道府県レベルのクラスタリングが全身麻酔分散の5.8%を説明するにとどまり（ICC=0.058）、'
    '大学病院の存在が分散の35.8%を説明した（β=+64.1、95%CI 54.8〜73.4、P<0.001）。'
    '都道府県内比較では、大学病院所在圏は全47都道府県すべてで非大学圏より高いL008 SCRを示した（100%、Cohen\'s d=1.88）。'
    'Empirical Bayes縮小推定による効果量の減衰は9.0%にとどまり（d=1.89→1.72）、小サンプル不安定性の影響は限定的であった。'
    'L008とL004は正の相関を示し（r=+0.235、P<0.001）、審査による再分類仮説と矛盾した。'
).font.name = 'MS Mincho'

p = doc.add_paragraph()
run = p.add_run('結論 ')
run.bold = True
run.font.name = 'MS Gothic'
p.add_run(
    '日本における麻酔実践の地域差は構造的なものであり、保険審査の影響による人工的産物ではない。'
    '大学病院の近接性が最大の規定因子であり、マルチレベル調整およびEmpirical Bayes縮小後も持続した。'
    '区域麻酔技術が腫瘍学的転帰の改善と関連するというエビデンスを踏まえると、この地域差は修正可能な健康格差の源泉である可能性がある。'
).font.name = 'MS Mincho'

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# キーメッセージ
# ═══════════════════════════════════════════════════════════

add_heading_jp('本研究で既に知られていること', level=2)
doc.add_paragraph(
    '外科・麻酔実践の地域差は多くの国で報告されているが、臨床的選好、施設要因、保険審査などの行政的要因の相対的寄与は不明である。',
    style='List Bullet'
)
doc.add_paragraph(
    '日本の国民皆保険制度は標準化された都道府県別審査委員会を通じてすべてのレセプトを処理しており、行政的変動と臨床的変動を区別できる独自の自然実験環境を提供している。',
    style='List Bullet'
)

add_heading_jp('本研究が追加すること', level=2)
doc.add_paragraph(
    '3つの独立した感度分析とマルチレベルモデリングにより、日本における麻酔実践の地域差が保険審査の影響による人工的産物ではなく構造的なものであることを示した。',
    style='List Bullet'
)
doc.add_paragraph(
    '大学病院の存在が全身麻酔量の都道府県内分散の40.5%を説明し（Cohen\'s d=1.88）、その効果は全47都道府県で一貫していた。',
    style='List Bullet'
)

add_heading_jp('本研究が研究・実践・政策に与える影響', level=2)
doc.add_paragraph(
    '腫瘍学的利益をもたらす可能性のある全身麻酔＋硬膜外麻酔の併用技術が最大の地域格差を示し（変動係数64.9%）、健康格差削減の修正可能な標的となりうる。',
    style='List Bullet'
)
doc.add_paragraph(
    '本研究で開発したマルチレベル感度分析フレームワークは、日本の国民皆保険制度下における任意の手技の実践変動調査に応用可能である。',
    style='List Bullet'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 緒言
# ═══════════════════════════════════════════════════════════

add_heading_jp('緒言', level=1)

# 日本の医療制度の説明（査読者向け）
doc.add_paragraph(
    '1961年に確立された日本の国民皆保険制度（kokumin kai hoken）は、標準化された診療報酬点数表に基づき、'
    '全国民を均一に保障している。¹·² 本制度は国際的に高い評価を受けており、'
    '世界保健機関（WHO）の医療アクセス・質指数（HAQ指数）において、日本は2019年に87.5/100（上位10%、世界最高水準）を達成している。³'
    'また、ユニバーサル・ヘルス・カバレッジの実効的サービスカバレッジでも世界最高水準に位置する。⁴'
    '医療提供体制は医療圏の階層構造によって組織されており、47都道府県が三次医療圏として高度専門医療を担い、'
    '335二次医療圏（nijiiryoken）が入院医療の基本的な地理的単位として設定されている。⁵'
    '患者の圏域をまたぐ受診は限定的であり、2008年患者調査では入院患者の6.1%、外来患者の3.0%のみが'
    '居住都道府県外で医療を受けていた。⁶ 二次医療圏をまたぐ移動は主に隣接圏域への移動にとどまる。⁷'
    '本制度は地域によらず均等な医療アクセスを保障することを目的として設計されている。'
)

doc.add_paragraph(
    'しかし、NDB（ナショナルデータベース）を用いた研究では、外科手技、処方パターン、診断検査など、'
    '医療実践に顕著な地域差が存在することが明らかになっている。⁸⁻¹⁰'
    '日本の制度の特徴として、保険審査（shinsa）が義務付けられている。'
    'すべてのレセプトは都道府県の審査委員会による審査を経て支払われ、不適切と判断されたレセプトは'
    '減額または返戻（査定）される。審査率は都道府県によって0.07%〜0.28%と差異がある。¹¹'
    'これは根本的な疑問を提起する：レセプトデータで観察される地域差は、臨床実践の真の差異を反映しているのか、'
    'それとも審査の差異による人工的産物なのか？'
)

doc.add_paragraph(
    'この問いは麻酔実践において特に重要である。麻酔技術の選択——全身麻酔単独か全身麻酔＋区域麻酔の併用か——は'
    '単なる選好の問題ではない。特に硬膜外麻酔を全身麻酔と組み合わせた区域麻酔技術が、がん手術における'
    '無再発生存率および全生存率を改善する可能性を示すエビデンスが蓄積されている。¹²⁻¹⁷'
    'Cochrane系統的レビューは区域麻酔技術ががん再発に有利である可能性を示唆し、¹²'
    'Sesslerらのランダム化比較試験では乳がん再発に有意差は認められなかったものの、¹³'
    '二次解析や観察的メタ解析では特定のがん種における潜在的利益が引き続き支持されている。¹³⁻¹⁷'
    '麻酔技術の地域差が審査の影響ではなく真の差異であれば、腫瘍学的転帰に影響する修正可能な健康格差の源泉となりうる。'
)

doc.add_paragraph(
    '2026年の診療報酬改定では、全身麻酔コードの名称が「マスク又は気管内挿管による閉鎖循環式全身麻酔」から'
    '「声門上器具使用又は気管内挿管による閉鎖循環式全身麻酔」に変更された。¹⁸'
    'これは従来のコーディングの曖昧さを反映しており、地域差の実態解明は時宜を得た課題である。'
)

doc.add_paragraph(
    '本研究では、二次医療圏レベルの公開SCRデータを用いて、'
    '（1）335医療圏における麻酔実践の地域差を定量化し、'
    '（2）3つの独立した感度分析により保険審査の影響を検証し、'
    '（3）マルチレベルモデルにより都道府県レベルのクラスタリングを推定し、'
    '（4）大学病院の近接性と麻酔科医供給を含む地域差の構造的規定因子を同定することを目的とした。'
)

# ═══════════════════════════════════════════════════════════
# 方法
# ═══════════════════════════════════════════════════════════

add_heading_jp('方法', level=1)

add_heading_jp('研究デザインと倫理的考慮', level=2)

doc.add_paragraph(
    '本研究は公開集計データを用いた横断的生態学的研究である。個人レベルのデータへのアクセスはなく、'
    '「人を対象とする生命科学・医学系研究に関する倫理指針」（2021年改正）に基づき倫理委員会の承認は不要である。'
    '本研究はSTROBEガイドラインに従って報告した（チェックリストは補足資料として添付）。'
)

add_heading_jp('地理的単位：二次医療圏', level=2)

doc.add_paragraph(
    '日本の医療提供体制は医療法第30条の4に基づく医療圏の階層構造によって組織されている。⁵'
    '47都道府県が三次医療圏として高度専門医療を担い、各都道府県内に設定された二次医療圏（nijiiryoken）が'
    '入院医療計画の基本的な地理的単位となっている。研究期間中、全国に335の二次医療圏が存在した（2018年改定で344→335に統合）。¹⁹'
)

doc.add_paragraph(
    '重要な点として、二次医療圏は都道府県内に完全にネストされており、都道府県境界をまたぐ区域は存在しない。'
    'これは、患者流動パターンによって定義され州境界をまたぐ可能性のある米国のHospital Service Area（HSA）や'
    'Hospital Referral Region（HRR）とは根本的に異なる。²⁰'
    '二次医療圏が都道府県内に完全にネストされているため、都道府県内比較は都道府県境界をまたぐ'
    'スピルオーバー効果に交絡されず、共通の都道府県審査方針下での変動として直接解釈できる。'
    '各二次医療圏は1つ以上の隣接市区町村で構成され、一般的な入院医療（外科サービスを含む）において'
    '自己完結的であるよう設計されている。'
)

add_heading_jp('データソースと標準化レセプト発生比', level=2)

doc.add_paragraph(
    '3つのデータソースを使用した。第一に、内閣府「地域差の見える化」イニシアチブの一環として公開された'
    '令和4年度（2022年度）の標準化レセプト発生比（SCR）データ。²¹'
    'SCRは間接標準化法により算出される：各手技について、全国の年齢・性別別レセプト発生率を地域の人口構成に適用して'
    '期待レセプト発生数を算出し、SCRは100×（観察数/期待数）として定義される。'
    'SCR=100は年齢・性別調整後の全国平均と等しいことを示す。'
    'SCRは居住地ベースで算出される：レセプトは医療機関の所在地ではなく、受給者の登録住所に帰属される。'
    'これにより、他の圏域への患者移動が受け入れ圏域のSCRを押し上げることはない。'
    'データは都道府県レベル（n=47）および二次医療圏レベル（n=335）の両方で利用可能であった。'
    '件数が少ない圏域のデータはプライバシー保護のためマスキング処理されており、欠損値として扱った。'
)

doc.add_paragraph(
    '第二に、二次医療圏別の麻酔科医数および総医師数を提供する2022年医師・歯科医師・薬剤師統計（e-Stat）。²²'
    '第三に、二次医療圏境界（A38-20）、都道府県境界および北方領土（N03）の国土数値情報データセット。²³'
)

add_heading_jp('麻酔手技コード', level=2)

doc.add_paragraph(
    '診療報酬点数表から麻酔実践の主要カテゴリーを代表する6つの手技コードを分析した（表1）：'
    'L008（閉鎖循環式全身麻酔）、L002（硬膜外麻酔）、L003（持続硬膜外注入、全身麻酔＋硬膜外麻酔併用の直接指標）、'
    'L004（脊椎麻酔）、L009（麻酔管理料I、専門麻酔科医配置の代理指標）、'
    'L100（神経ブロック、入院、ペインクリニック活動の指標）。'
    '静脈麻酔（L001）も副次的アウトカムとして分析した。'
)

add_heading_jp('大学病院マッピング', level=2)

doc.add_paragraph(
    '81の大学病院（国立44、公立8、私立29）を市区町村住所に基づいて各二次医療圏にマッピングした。'
    'これら81病院は全47都道府県の64二次医療圏に分布しており、日本の「一県一医大」政策を反映している。'
    'この分布により、都道府県内での大学病院所在圏と非所在圏の比較が可能となり、'
    '都道府県審査方針（各都道府県内で均一に適用）の自然なコントロールとなる。'
)

add_heading_jp('統計解析', level=2)

doc.add_paragraph(
    '3つの補完的な分析アプローチを採用した。第一に、記述統計（平均値、標準偏差、変動係数、パーセンタイル、範囲）により'
    '二次医療圏間のSCR分布を特徴付けた。群間比較にはWelchのt検定（等分散を仮定しない）、'
    '効果量にはCohen\'s d、²⁴ ノンパラメトリック確認にはMann-Whitney U検定を使用した。'
)

doc.add_paragraph(
    '第二に、二次医療圏が都道府県内にネストされた構造を考慮するため、マルチレベル線形混合モデル'
    '（階層線形モデルまたは混合効果モデルとも呼ばれる）を適合させた。²⁵'
    '各麻酔コードについて、二次医療圏（レベル1、n=335）が都道府県（レベル2、n=47）内にネストされた'
    '2レベル構造を設定した。まず固定効果なしのヌルモデル（ランダム切片のみ）を適合させ、'
    '都道府県レベルに帰属する全分散の割合として定義される級内相関係数（ICC）を推定した。'
    '続いて大学病院の有無（二値）、麻酔科医密度（標準化）、総医師密度（標準化）の固定効果を追加した。'
    'モデルはPython statsmodels MixedLM実装²⁶ を用いて制限付き最尤法（REML）で推定した。'
    '周辺R²はヌルモデルと各適合モデル間の全分散（都道府県＋残差）の比例的減少として算出した。'
)

doc.add_paragraph(
    '第三に、保険審査の影響が観察されたSCR差を説明するという帰無仮説を検証するため、'
    '3つの感度分析を実施した。（a）都道府県内分散分解：都道府県審査方針（各都道府県内で均一に適用）が'
    'すべての変動を説明するなら、都道府県内分散は無視できるはずである。一元配置ANOVAを用いて'
    '全SCR分散を都道府県間・都道府県内成分に分解し、さらに都道府県内分散を大学病院効果と残差に分解した。'
    '（b）コード間相関分析：審査が全身麻酔と脊椎麻酔のコード間でレセプトを再分類するなら、'
    'これらは負の相関を示すはずである。全コードペア間のPearson相関を算出した。'
    '（c）審査影響の定量的推定：公開された集計審査統計を用いて審査率差の最大SCR影響を推定した。'
)

doc.add_paragraph(
    '小サンプルSMAにおけるSCR推定値の不安定性に対処するため、Empirical Bayes（EB）縮小推定を適用した。²⁷'
    '各二次医療圏のEB推定値は観察SCRと都道府県平均の加重平均であり、重みは局所推定値の相対的精度を反映する。'
    '推定値が不安定なSMA（グループに対して高い分散）は都道府県平均に向けて「縮小」され、'
    'より保守的な効果推定値が得られる。生SCRとEB調整SCRの両方を用いてすべての主要知見を比較した。'
    'また、プライバシーマスキングにより欠損値となったSMAは関連する分析から除外し、'
    '各分析に寄与したSMA数を報告した。'
)

add_heading_jp('患者・市民参画', level=2)

doc.add_paragraph(
    '本研究は個人患者情報を含まない公開集計データを使用した。患者は本研究の設計・実施に関与していない。'
)

# ═══════════════════════════════════════════════════════════
# 結果
# ═══════════════════════════════════════════════════════════

add_heading_jp('結果', level=1)

add_heading_jp('研究対象と記述統計', level=2)

doc.add_paragraph(
    '335二次医療圏は47都道府県に分布していた（都道府県あたり中央値7圏域、範囲3〜21）。'
    '335圏域のうち64圏域（19.1%）に1つ以上の大学病院が存在した。'
    '全身麻酔（L008）のSCRデータは334圏域で利用可能であった（1圏域はマスキング）；'
    '硬膜外麻酔（L002）は307圏域（28圏域はマスキング、一部農村部での低件数を反映）。'
    '表2に各コードのSCR分布を示す。'
)

# 表1
add_para('表1. 分析した麻酔手技コード', bold=True)
table1 = doc.add_table(rows=1, cols=4)
table1.style = 'Table Grid'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr1 = table1.rows[0].cells
for i, text in enumerate(['コード', '手技', '臨床的意義', 'データあり圏域数']):
    hdr1[i].paragraphs[0].add_run(text).bold = True
    hdr1[i].paragraphs[0].runs[0].font.size = Pt(9)

codes_data = [
    ['L008', '閉鎖循環式全身麻酔', '全身麻酔量の主要指標', '334'],
    ['L002', '硬膜外麻酔', '区域麻酔技術（単独または併用）', '307'],
    ['L003', '持続硬膜外注入', '全身麻酔＋硬膜外麻酔併用の直接指標', '331'],
    ['L004', '脊椎麻酔', '適応手術における全身麻酔の代替', '334'],
    ['L009', '麻酔管理料I', '専門麻酔科医配置の代理指標', '314'],
    ['L100', '神経ブロック（入院）', 'ペインクリニック活動の指標', '335'],
]
for row_data in codes_data:
    add_table_row(table1, row_data)
doc.add_paragraph()

# 表2
add_para('表2. 二次医療圏間の標準化レセプト発生比（SCR）の分布（全国平均=100）', bold=True)
table2 = doc.add_table(rows=1, cols=8)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0].cells
for i, text in enumerate(['コード', 'n', '平均', 'SD', '中央値', 'IQR', '範囲', 'CV(%)']):
    hdr2[i].paragraphs[0].add_run(text).bold = True
    hdr2[i].paragraphs[0].runs[0].font.size = Pt(9)

table2_data = [
    ['L008', '334', '79.7', '42.7', '73.2', '47.9〜102.7', '2.3〜435.7', '53.6'],
    ['L002', '307', '95.6', '83.2', '73.9', '30.3〜133.3', '0.9〜519.1', '87.0'],
    ['L003', '331', '86.1', '55.9', '73.1', '44.6〜115.6', '1.5〜350.2', '64.9'],
    ['L004', '334', '93.8', '53.2', '84.3', '55.1〜121.7', '1.0〜346.5', '56.8'],
    ['L009', '314', '87.7', '49.4', '79.1', '51.8〜115.2', '1.2〜296.3', '56.4'],
    ['L100', '335', '99.7', '76.2', '72.2', '43.3〜133.7', '1.4〜482.5', '76.4'],
]
for row_data in table2_data:
    add_table_row(table2, row_data)
doc.add_paragraph()
add_para('CV=変動係数、IQR=四分位範囲、SD=標準偏差', italic=True)

add_heading_jp('麻酔実践の地域差', level=2)

doc.add_paragraph(
    '全麻酔コードで顕著な地域差が認められた（表2）。全身麻酔（L008）のSCRは334圏域で2.3〜435.7の範囲にあり'
    '（変動係数53.6%）、最低値と最高値の間に189倍の差があった。'
    '硬膜外麻酔（L002）はさらに大きな相対的変動を示し（変動係数87.0%）、'
    '全身麻酔＋硬膜外麻酔併用の直接指標である持続硬膜外注入（L003）も同様であった（変動係数64.9%）。'
    '脊椎麻酔（L004）の変動係数は56.8%であった。図1はSCRの地理的分布を示す。'
)

add_heading_jp('マルチレベルモデル結果', level=2)

# 表3
add_para('表3. マルチレベル線形混合モデル結果：SCRをアウトカム、都道府県をランダム切片とする', bold=True)
table3 = doc.add_table(rows=1, cols=7)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr3 = table3.rows[0].cells
for i, text in enumerate(['コード', 'n', 'ICC（ヌル）', 'β大学病院（95%CI）', 'P値', 'ICC（調整後）', '周辺R²']):
    hdr3[i].paragraphs[0].add_run(text).bold = True
    hdr3[i].paragraphs[0].runs[0].font.size = Pt(9)

table3_data = [
    ['L008', '334', '0.058', '+64.1（54.8〜73.4）', '<0.001', '0.048', '0.358'],
    ['L004', '334', '0.343', '+23.0（11.0〜35.1）', '<0.001', '0.360', '0.020'],
    ['L002', '307', '0.191', '+49.0（27.2〜70.8）', '<0.001', '0.057', '0.125'],
    ['L003', '331', '0.106', '+55.8（43.0〜68.7）', '<0.001', '0.142', '0.162'],
    ['L001', '327', '0.094', '+83.2（66.3〜100.0）', '<0.001', '0.133', '0.204'],
    ['L100（入院）', '335', '0.185', '+44.8（25.9〜63.6）', '<0.001', '0.218', '0.035'],
    ['L100（外来）', '335', '0.299', '+37.2（22.4〜52.0）', '<0.001', '0.322', '0.044'],
]
for row_data in table3_data:
    add_table_row(table3, row_data)
doc.add_paragraph()
add_para(
    'ICC=級内相関係数（都道府県レベルの分散割合）；β大学病院=大学病院存在の固定効果係数；'
    '周辺R²=ヌルモデルからの全分散の比例的減少。すべてのモデルに都道府県ランダム切片を含む。',
    italic=True
)

doc.add_paragraph(
    'ヌルマルチレベルモデルでは、都道府県レベルのクラスタリングが全身麻酔（L008）分散の5.8%を説明するにとどまり'
    '（ICC=0.058）、94.2%の変動が審査方針が均一な都道府県内で生じていることが示された。'
    '対照的に、脊椎麻酔（L004）は都道府県レベルのクラスタリングが顕著に強く（ICC=0.343）、'
    '都道府県要因（審査実践を含む可能性）がこのコードでより大きな役割を果たすことが示唆された。'
)

doc.add_paragraph(
    '大学病院の有無を固定効果として追加すると、全身麻酔（L008）で最大の改善が得られた'
    '（周辺R²=0.358；β=+64.1、95%CI 54.8〜73.4）。'
    '大学病院の存在のみで全L008分散の35.8%が説明された。'
    '効果はすべてのコードで統計的に有意であったが（すべてP<0.001）、大きさは異なった：'
    '静脈麻酔（L001：β=+83.2）と全身麻酔（L008：β=+64.1）で最大、'
    '持続硬膜外（L003：β=+55.8）と硬膜外（L002：β=+49.0）で中程度、'
    '脊椎麻酔（L004：β=+23.0）で最小であった。'
)

add_heading_jp('感度分析1：都道府県内分散分解', level=2)

doc.add_paragraph(
    '階層的分散分解では、L008分散の14.5%が都道府県間差に、40.5%が大学病院効果（都道府県内）に、'
    '45.0%が残差都道府県内変動に帰属した。大学病院効果のみ——単一の二値変数——が'
    '全身麻酔量の都道府県内分散の47.4%を説明した。'
    'このパターンは他のコードでも一貫しており、大学病院効果は都道府県内L009分散の30.5%、'
    'L002分散の21.5%、L004分散の15.9%を説明した。'
)

doc.add_paragraph(
    '都道府県内比較（審査方針をコントロール）では、大学病院所在圏は全47都道府県すべてで'
    '非大学圏より高いL008 SCRを示した（100%、都道府県内平均差+63.3点、SD 32.7、対応t=13.28、P<0.001）。'
    '麻酔管理料I（L009、専門医配置の代理指標）では47都道府県中41都道府県（87%）、'
    '硬膜外麻酔（L002）では47都道府県中40都道府県（85%）で同様のパターンが認められた。'
)

add_heading_jp('感度分析2：コード間相関', level=2)

doc.add_paragraph(
    '全身麻酔（L008）と脊椎麻酔（L004）は正の相関を示した（r=+0.235、P<0.001）。'
    'この知見は、審査によるこれら2コード間のレセプト再分類仮説（負の相関を予測）と矛盾する。'
    '正の相関は、外科・麻酔全体の診療能力が高い地域では両技術の利用率が高いことを示唆する。'
    '全身麻酔（L008）と硬膜外麻酔（L002）も正の相関を示し（r=+0.319、P<0.001）、'
    '審査による代替ではなく共通の供給要因（麻酔科の診療能力）と整合的であった。'
)

add_heading_jp('感度分析3：審査影響の定量的推定', level=2)

doc.add_paragraph(
    '都道府県別審査率の最大差は0.21パーセントポイント（0.07%〜0.28%）であった。'
    'この差がすべて単一の麻酔コードに集中したとしても、SCRへの影響は約0.2ポイントにとどまる——'
    'L008の観察された四分位範囲62ポイントの0.3%未満、L002の範囲198ポイントの0.1%未満である。'
    '審査率の変動は観察された地域差を定量的に説明できない。'
)

add_heading_jp('合算SCR分析', level=2)

doc.add_paragraph(
    'L008とL004のSCRを合算すると、変動係数は53.6%から43.4%に低下した（19%の減少）。'
    'これは、審査による再分類、防衛的過少請求、または臨床的選択によるいかなる脊椎麻酔〜全身麻酔代替も、'
    '全身麻酔変動の少数を占めるにすぎないことを示す。'
    '大学病院効果は合算指標でも強く持続した：大学圏の平均合算SCRは242.9対非大学圏157.6'
    '（Cohen\'s d=1.26、P<0.001）。'
)

add_heading_jp('小サンプル感度分析：Empirical Bayes縮小推定', level=2)

doc.add_paragraph(
    '都道府県平均に向けたEmpirical Bayes縮小推定は効果量を緩やかに減衰させたが、'
    'いかなる結論も実質的に変化させなかった（表4）。全身麻酔（L008）では、大学病院効果が'
    'Cohen\'s d=1.89（生データ）からd=1.72（EB）に9.0%減衰した。'
    '脊椎麻酔（L004）では9.1%（d=0.40→0.36）、硬膜外麻酔（L002）では8.6%（d=0.59→0.54）の減衰であった。'
    'いずれの場合も、縮小後の効果量は大きく統計的に有意であり、'
    '知見が小量圏域の不安定な推定値によって駆動されていないことが確認された。'
)

add_heading_jp('大学病院効果', level=2)

doc.add_paragraph(
    '大学病院の存在が麻酔実践パターンの最大の単一予測因子であった（図2）。'
    '全身麻酔（L008）では、64の大学病院所在圏の平均SCRは132.1（SD 47.5）であったのに対し、'
    '270の非大学圏では67.3（SD 30.5）であった（Cohen\'s d=1.88、Welchのt=10.40、P<0.001；Mann-Whitney P<0.001）。'
    '東京では、圏域あたりの大学病院数が0〜5の範囲にあり、大学病院数とともにSCRが単調に増加する'
    '用量反応関係が認められた。'
)

add_heading_jp('全身麻酔＋硬膜外麻酔の組み合わせ', level=2)

doc.add_paragraph(
    '全身麻酔＋硬膜外麻酔併用の直接指標である持続硬膜外注入（L003）は変動係数64.9%を示した。'
    'L003 SCRは大学病院所在圏で非大学圏の1.73倍であった。'
    '硬膜外麻酔の組み合わせ（L008+L002）に適用した3つの審査仮説検定はいずれも審査説明を棄却した：'
    'L008〜L002の正の相関（r=+0.319）、合算L008+L004での持続する大学病院効果（d=1.26）。'
)

# ═══════════════════════════════════════════════════════════
# 考察
# ═══════════════════════════════════════════════════════════

add_heading_jp('考察', level=1)

add_heading_jp('主要な知見', level=2)

doc.add_paragraph(
    '本研究は、日本の335二次医療圏における麻酔実践の地域差が顕著かつ構造的であることを示した。'
    'マルチレベルモデルでは、都道府県レベルのクラスタリングが全身麻酔分散の5.8%を説明するにとどまり'
    '（ICC=0.058）、94.2%の変動が審査方針が均一な都道府県内で生じていることが示された。'
    '大学病院の存在が最大の規定因子であり、マルチレベルモデルで全L008分散の35.8%、'
    '階層的分解で都道府県内分散の40.5%を説明した。'
    'この効果が47都道府県すべてで認められること——異なる審査方針にもかかわらず——は、'
    '供給側・施設的メカニズムの強力なエビデンスを提供する。'
    'Empirical Bayes縮小推定は大学病院効果を9%しか減衰させず、'
    '結果が小量圏域の不安定な推定値によって駆動されていないことが確認された。'
)

add_heading_jp('日本の医療制度の文脈', level=2)

doc.add_paragraph(
    'これらの知見を解釈するには日本の医療制度の理解が必要である。'
    '日本の国民皆保険は全国民を単一の全国統一診療報酬体系で保障している。¹·²'
    'WHO HAQ指数は2019年に日本を87.5/100（上位10%）と評価しており、³ 高品質な医療アクセスを反映している。'
    '分析単位として使用した335二次医療圏は入院医療において自己完結的であるよう設計されており、'
    '圏域をまたぐ患者移動は限定的である：2008年患者調査では入院患者の6.1%のみが居住都道府県外で医療を受けていた。⁶'
    '都道府県内では、二次医療圏をまたぐ移動は主に隣接圏域への移動にとどまり、⁷'
    '同一都道府県内の非隣接圏域間のSCR変動は患者の選別よりも地域の診療パターンを主に反映する。'
    '居住地ベースのSCR算出はこの懸念をさらに軽減する。'
)

add_heading_jp('既存文献との比較', level=2)

doc.add_paragraph(
    '本研究の知見は医療実践変動に関する広範な文献と一致している。'
    '米国のDartmouth Atlas Projectは、患者ニーズよりも医師供給と診療スタイルによって主に駆動される'
    '外科手術率の広範な地域差を記録している。²⁸·²⁹'
    '同様のパターンは英国、³⁰ ドイツ、³¹ オーストラリア³² でも報告されている。'
    '本研究は、日本の独自の制度構造を活用して行政的変動と臨床的変動を区別し、'
    '東アジアの国民皆保険制度における細かい地理的スケール（335圏域）での麻酔技術変動の'
    '初のマルチレベル分析を提供することで、この文献を拡張する。'
)

doc.add_paragraph(
    '本研究で記録した大学病院効果（全身麻酔のCohen\'s d=1.88）は、医療実践変動研究で通常報告される'
    '効果量と比較して顕著に大きい。これは、大学医局が地理的圏域内の関連病院における臨床実践に'
    '実質的な影響力を持つ日本の医療教育システムを反映していると考えられる。³³·³⁴'
    'L008のICC=0.058は米国の外科手術率で報告されるICC（通常0.10〜0.30）²⁸ より顕著に低く、'
    '都道府県レベル要因（審査方針を含む）が全身麻酔において都道府県内施設要因と比較して'
    '相対的に小さな役割を果たすことを示唆する。'
)

add_heading_jp('腫瘍学的転帰への含意', level=2)

doc.add_paragraph(
    '全身麻酔＋硬膜外麻酔併用（L003）の変動係数が64.9%であるという知見は、'
    '区域麻酔を全身麻酔と組み合わせることでがん手術における無再発生存率が改善する可能性を示す'
    'エビデンスの観点から臨床的に重要である。¹²⁻¹⁷'
    '大学病院所在圏と非所在圏の間のL003 SCRの1.73倍の差は、'
    'この潜在的に有益な技術へのアクセスが居住地に大きく依存することを示唆しており、'
    '日本の国民皆保険制度の公平性の前提に疑問を呈する知見である。¹·²'
)

add_heading_jp('強みと限界', level=2)

doc.add_paragraph(
    '強みとして、全国民をカバーする年齢・性別標準化比の使用；細かい地理的スケール（335圏域）での分析；'
    'ネストされたデータ構造を適切に考慮したマルチレベルモデリング；小サンプル不安定性に対処する'
    'Empirical Bayes縮小推定；複数の独立した感度分析；審査方針をコントロールする都道府県内比較；'
    '患者移動効果を軽減する居住地ベースのSCR算出が挙げられる。'
)

doc.add_paragraph(
    '限界として、地域レベル分析に内在する生態学的誤謬；³⁵ 集計パターンから個人レベルの意思決定を'
    '区別できないこと；SCRデータの審査後の性質（意図した実践ではなく償還された実践を反映するが、'
    '感度分析はこの区別が定量的に軽微であることを示唆する）；因果推論を妨げる横断的デザイン；'
    '大学病院効果と都市集中の不完全な分離（大学病院は通常各都道府県の最大都市に立地）；'
    'コード別審査率データの欠如；クレームデータからは定量化できない「防衛的過少請求」³⁶ が挙げられる。'
    'また、マルチレベルモデルには大学病院の有無のみを固定効果として含めた；'
    '追加の地域レベル共変量（人口密度、都市化指数、病床密度）を含む将来の研究により'
    '追加の分散が説明される可能性がある。'
)

add_heading_jp('結論', level=2)

doc.add_paragraph(
    '日本における麻酔実践の地域差は真実かつ大きく、保険審査の差異ではなく、'
    '特に大学病院の近接性という施設的要因によって主に駆動されている。'
    'マルチレベルモデリングにより、都道府県レベルのクラスタリングが全身麻酔分散の5.8%のみを説明し、'
    '大学病院の存在が35.8%を説明することが確認された。'
    'Empirical Bayes縮小推定によりこれらの知見が小サンプル不安定性に対して頑健であることが確認された。'
    '全身麻酔＋硬膜外麻酔の変動は腫瘍学的利益のエビデンスの観点から特に懸念される。'
    '政策的対応として、標的を絞った教育プログラム、大学センターからの専門医アウトリーチ、'
    '品質指標としての麻酔技術パターンのモニタリングが考えられる。³⁷'
    '本研究で開発したマルチレベル感度分析フレームワークは、日本の国民皆保険制度下における'
    '任意の手技の実践変動調査に応用可能である。'
)

# ═══════════════════════════════════════════════════════════
# 参考文献
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_jp('参考文献', level=1)

refs = [
    '1. Ikegami N, Yoo B-K, Hashimoto H, et al. Japanese universal health coverage: evolution, '
    'achievements, and challenges. Lancet 2011;378:1106\u201315.',
    '2. Ikegami N. Universal Health Coverage for Inclusive and Sustainable Development: Lessons from '
    'Japan. Washington, DC: World Bank, 2014.',
    '3. GBD 2019 Healthcare Access and Quality Collaborators. Assessing performance of the Healthcare '
    'Access and Quality Index, overall and by select age groups, for 204 countries and territories, '
    '1990\u20132019: a systematic analysis from the Global Burden of Disease Study 2019. Lancet Glob '
    'Health 2022;10:e1715\u201343.',
    '4. World Health Organization. Tracking Universal Health Coverage: 2023 Global Monitoring Report. '
    'Geneva: WHO, 2023.',
    '5. Ministry of Health, Labour and Welfare. Medical Care Plan Guidelines (Iryo Keikaku Sakutei '
    'Shishin). Tokyo: MHLW, 2023.',
    '6. Ministry of Health, Labour and Welfare. Patient Survey 2008 (Kanja Chosa). '
    'https://www.mhlw.go.jp/toukei/saikin/hw/kanja/08/ (accessed 15 Mar 2026).',
    '7. Tokyo Metropolitan Government. Regional Medical Care Plan: Cross-Area Patient Movement '
    'Analysis. Tokyo: Bureau of Social Welfare and Public Health, 2025.',
    '8. Matsuda S, Fujimori K. The claim database in Japan. Asian Pac J Dis Manag 2012;6:55\u201359.',
    '9. Suto M, Iba A, Sugiyama T, et al. Literature review of studies using the National Database '
    'of Health Insurance Claims of Japan (NDB). JMA J 2023;7:10\u201320.',
    '10. Hamada H, Sekimoto M, Imanaka Y. Effects of the per diem prospective payment system with '
    'DRG-like grouping system (DPC/PDPS) on resource usage and healthcare quality in Japan. Health '
    'Policy 2012;107:194\u2013201.',
    '11. Social Insurance Medical Fee Payment Fund. Review Statistics, Fiscal Year 2022. '
    'https://www.ssk.or.jp/tokeijoho/shinsatokei/index.html (accessed 15 Mar 2026).',
    '12. Pei L, Tan G, Wang L, et al. Anesthetic techniques for risk of malignant tumor recurrence. '
    'Cochrane Database Syst Rev 2019;(11):CD008877.',
    '13. Sessler DI, Pei L, Huang Y, et al. Recurrence of breast cancer after regional or general '
    'anesthesia: a randomized controlled trial. Lancet 2019;394:1807\u201315.',
    '14. Exadaktylos AK, Buggy DJ, Moriarty DC, et al. Can anesthetic technique for primary breast '
    'cancer surgery affect recurrence or metastasis? Anesthesiology 2006;105:660\u20134.',
    '15. Sun Y, Li T, Gan TJ. The effects of perioperative regional anesthesia and analgesia on cancer '
    'recurrence and survival after oncology surgery. Reg Anesth Pain Med 2015;40:589\u201398.',
    '16. Xie S, Li L, Meng F, et al. Regional anesthesia might reduce recurrence and metastasis rates '
    'in adult patients with cancers after surgery. BMC Anesthesiol 2024;24:19.',
    '17. Zhang D, Jiang J, Liu J, et al. Effects of perioperative epidural analgesia on cancer '
    'recurrence and survival. Front Oncol 2021;11:798435.',
    '18. Ministry of Health, Labour and Welfare. Revision of medical fee schedule, 2026. '
    'https://www.mhlw.go.jp (accessed 15 Mar 2026).',
    '19. Matsuda S. Secondary medical areas and regional healthcare planning in Japan. '
    'J Health Care Soc 2018;28:137\u201348. [in Japanese]',
    '20. Wennberg JE, Cooper M. The Dartmouth Atlas of Health Care 1999. Chicago: '
    'American Hospital Publishing, 1999.',
    '21. Cabinet Office. Regional variation visualization (chiikisa no mieruka). '
    'https://www5.cao.go.jp/keizai-shimon/kaigi/special/reform/mieruka/ (accessed 10 Mar 2026).',
    '22. Ministry of Health, Labour and Welfare. Survey of Physicians, Dentists and Pharmacists, '
    '2022. e-Stat. https://www.e-stat.go.jp (accessed 12 Mar 2026).',
    '23. Ministry of Land, Infrastructure, Transport and Tourism. National Land Numerical Information. '
    'https://nlftp.mlit.go.jp (accessed 10 Mar 2026).',
    '24. Cohen J. Statistical Power Analysis for the Behavioral Sciences. 2nd ed. '
    'Hillsdale, NJ: Lawrence Erlbaum, 1988.',
    '25. Rabe-Hesketh S, Skrondal A. Multilevel and Longitudinal Modeling Using Stata. 3rd ed. '
    'College Station, TX: Stata Press, 2012.',
    '26. Seabold S, Perktold J. Statsmodels: econometric and statistical modeling with Python. '
    'In: Proceedings of the 9th Python in Science Conference. 2010:92\u201396.',
    '27. Efron B, Morris C. Stein\'s estimation rule and its competitors\u2014an empirical Bayes approach. '
    'J Am Stat Assoc 1973;68:117\u201330.',
    '28. Wennberg JE. Tracking Medicine. New York: Oxford University Press, 2010.',
    '29. Wennberg JE, Fisher ES, Skinner JS. Geography and the debate over Medicare reform. '
    'Health Aff (Millwood) 2002;Suppl Web Exclusives:W96\u2013114.',
    '30. NHS RightCare. The NHS Atlas of Variation in Healthcare, 2015. London: Public Health England.',
    '31. Wengler A, Nimptsch U, Mansky T. Hip and knee replacement surgery and geographic variations '
    'in Germany. Dtsch Arztebl Int 2014;111:407\u201316.',
    '32. Australian Commission on Safety and Quality in Health Care. Australian Atlas of Healthcare '
    'Variation. Sydney: ACSQHC, 2015.',
    '33. Otsuka T. The ikyoku system of university orthopedic surgery departments. '
    'J Orthop Sci 2012;17:513\u201314.',
    '34. Onishi H. History of Japanese medical education. Korean J Med Educ 2018;30:283\u201394.',
    '35. Greenland S. Ecologic versus individual-level sources of bias in ecologic estimates. '
    'Int J Epidemiol 2001;30:1343\u201350.',
    '36. Sekimoto M, Imanaka Y, Kobayashi H, et al. Impact of hospital characteristics on the cost '
    'and practice patterns of acute myocardial infarction patients in Japan. Health Policy '
    '2006;78:34\u201345.',
    '37. Merry AF, Cooper JB, Soyannwo O, et al. International standards for a safe practice of '
    'anesthesia 2010. Can J Anesth 2010;57:1027\u201334.',
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.name = 'MS Mincho'

# ═══════════════════════════════════════════════════════════
# 図
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_jp('図', level=1)

figure_list = [
    ('/home/ubuntu/rapm_fig1_jp.png',
     '図1. 令和4年度における二次医療圏別麻酔標準化レセプト発生比（SCR）の地理的分布。'
     '（A）全身麻酔（L008）SCR。（B）脊椎麻酔（L004）SCR。'
     'コロプレスマップ（実線：都道府県境界、点線：二次医療圏境界）。'
     '北方領土は白色で表示（二次医療圏未設定）。全国平均=100。'),
    ('/home/ubuntu/rapm_fig2_jp.png',
     '図2. 大学病院効果と合算SCR分析。'
     '（A）二次医療圏別大学病院所在状況：赤丸は1つ以上の大学病院を含む圏域（335圏域中64圏域）。'
     '（B）全身麻酔＋脊椎麻酔合算（L008+L004）SCR：これらのコードを合算することで'
     '審査による再分類の影響を中和するが、顕著な地域差が持続する（変動係数43.4%）。'
     'コロプレスマップ。'),
    ('/home/ubuntu/rapm_fig3_jp.png',
     '図3. 全身麻酔＋硬膜外麻酔の変動。'
     '（A）持続硬膜外注入（L003）SCR——全身麻酔＋硬膜外麻酔併用の直接指標（変動係数64.9%）。'
     '（B）L003/L008比（全身麻酔1件あたりの硬膜外麻酔併用率）。'
     '高値は全身麻酔＋硬膜外麻酔の組み合わせ技術の高い使用率を示す。コロプレスマップ。'),
    ('/home/ubuntu/rapm_fig4_jp.png',
     '図4. 3次元押出マップ。色はL003/L008比（硬膜外麻酔併用率）、高さは圏域あたりの麻酔科医数を表す。'
     '硬膜外麻酔併用率と専門医数の両方が高い圏域は、高く緑色の多角形として表示され、'
     '通常は大都市圏の大学病院所在圏に対応する。北方領土は基底高さで表示。'),
]

for path, caption in figure_list:
    add_figure(path, caption, width=Inches(5.5))
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# 宣言事項
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_jp('宣言事項', level=1)

add_heading_jp('倫理承認', level=2)
doc.add_paragraph(
    '本研究は公開集計データを使用した。個人レベルのデータへのアクセスはなく、'
    '「人を対象とする生命科学・医学系研究に関する倫理指針」（2021年改正）に基づき倫理委員会の承認は不要である。'
)

add_heading_jp('データ利用可能性', level=2)
doc.add_paragraph(
    '本研究で使用したすべてのデータは公開されている。SCRデータ：内閣府地域差の見える化'
    '（https://www5.cao.go.jp/keizai-shimon/kaigi/special/reform/mieruka/）。'
    '医師統計：e-Stat（https://www.e-stat.go.jp）。GIS境界：国土数値情報（https://nlftp.mlit.go.jp）。'
    '分析コードはhttps://github.com/bougtoir/wipで公開している。'
)

add_heading_jp('利益相反', level=2)
doc.add_paragraph('[著者が記入]')

add_heading_jp('資金提供', level=2)
doc.add_paragraph('[著者が記入]')

add_heading_jp('著者貢献', level=2)
doc.add_paragraph('[著者が記入]')

add_heading_jp('AI使用に関する声明', level=2)
doc.add_paragraph(
    'Devin（Cognition AI）をデータ分析、統計モデリング、地理情報システム可視化、'
    '原稿作成のコーディングアシスタントとして使用した。'
    'すべての科学的解釈と臨床的判断は人間の著者が行った。'
    'AI生成のコードとテキストは著者によってレビューおよび検証された。'
)

# ═══════════════════════════════════════════════════════════
# 保存
# ═══════════════════════════════════════════════════════════

output_path = '/home/ubuntu/regional_anesthesia_REVISED_JP.docx'
doc.save(output_path)
print(f"Saved: {output_path}")
