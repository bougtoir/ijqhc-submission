#!/usr/bin/env python3
"""Create cover letter for revised manuscript submission (BMJ Open target)."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Date
p = doc.add_paragraph(datetime.date.today().strftime('%B %d, %Y'))
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph('')

# Addressee
doc.add_paragraph('The Editor')
doc.add_paragraph('BMJ Open')

doc.add_paragraph('')

p = doc.add_paragraph('Dear Editor,')
doc.add_paragraph('')

# Body
p = doc.add_paragraph()
p.add_run('Re: ').bold = True
p.add_run('Submission of original research article — ')
p.add_run('"Regional variation in anaesthesia practice in Japan: clinical preference or claims audit artefact? '
          'A cross-sectional ecological study of 335 secondary medical areas"').italic = True

doc.add_paragraph('')

doc.add_paragraph(
    'We are pleased to submit the above manuscript for consideration for publication in BMJ Open. '
    'This study examines geographic variation in anaesthesia practice across Japan\'s 335 secondary '
    'medical areas (SMAs) using national claims data, and investigates whether observed variation '
    'reflects true differences in clinical practice or artefacts of regional variation in claims '
    'audit stringency.'
)

doc.add_paragraph(
    'An earlier version of this work was submitted to BMJ (manuscript ID: BMJ-2026-090550) and '
    'subsequently transferred to Regional Anesthesia & Pain Medicine (RAPM), where it was reviewed '
    'but not accepted. We have substantially revised the manuscript in response to the constructive '
    'reviewer feedback received at RAPM, incorporating the following major improvements:'
)

# Bullet list of improvements
improvements = [
    ('Multilevel modelling: ', 'We implemented hierarchical linear mixed models (SMAs nested within '
     '47 prefectures) with intraclass correlation coefficients (ICCs) to quantify the proportion of '
     'variance attributable to prefecture-level versus SMA-level factors. This directly addresses '
     'the clustered data structure.'),
    ('Small-sample sensitivity analysis: ', 'We applied empirical Bayes shrinkage estimation to '
     'address potential instability in SCR estimates for SMAs with small claim volumes. The university '
     'hospital effect was attenuated by only 8.6–9.1%, confirming robustness.'),
    ('Healthcare system context: ', 'We added a comprehensive explanation of Japan\'s universal health '
     'coverage system and medical area structure for international readers, supported by WHO Healthcare '
     'Access and Quality Index data (Japan: 87.5/100, top decile globally) and national patient '
     'movement statistics.'),
    ('Residence-based SCR clarification: ', 'We explicitly clarified that Standardised Claim-occurrence '
     'Ratios are calculated on a residence basis (attributed to patients\' registered addresses, not '
     'provider locations), which mitigates concerns about patient movement bias.'),
    ('Methods restructuring: ', 'The Methods section has been rewritten from bullet-point format into '
     'flowing paragraph style with clearly defined subsections, following STROBE guidelines for '
     'cross-sectional ecological studies.'),
    ('Data corrections: ', 'We corrected two university hospital geographic mappings (Asahikawa Medical '
     'University and Hirosaki University) and re-ran all analyses with verified data.'),
]

for bold_part, text_part in improvements:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(bold_part).bold = True
    p.add_run(text_part)

doc.add_paragraph('')

doc.add_paragraph(
    'Key findings include: (1) substantial geographic variation in general anaesthesia claims '
    '(L008 SCR range: 11.2–463.6; coefficient of variation 43.7%); (2) a strong inverse correlation '
    'between general anaesthesia (L008) and spinal anaesthesia (L004) SCRs (Spearman ρ = −0.50, '
    'p < 0.001), suggesting claims audit–driven reclassification rather than true clinical difference; '
    '(3) university hospital presence explains 40.5% of within-prefecture variance in general '
    'anaesthesia SCR (Cohen\'s d = 1.88, p < 0.001); and (4) when combining L008 and L004 SCRs '
    'to neutralise audit substitution effects, the coefficient of variation decreases from 43.7% '
    'to 26.1%, indicating that a substantial proportion of apparent variation is an audit artefact.'
)

doc.add_paragraph(
    'We believe this study is well suited to BMJ Open\'s scope and readership for several reasons. '
    'First, it addresses a health policy question — whether administrative claims data accurately '
    'reflect clinical reality — that has direct implications for evidence-based healthcare planning '
    'in any system using claims-based metrics. Second, it provides a methodological framework '
    '(sensitivity analyses combining related procedure codes to detect audit effects) that is '
    'transferable to other countries and clinical domains. Third, it offers an accessible '
    'introduction to Japan\'s healthcare system structure, contributing to the global literature '
    'on healthcare variation research.'
)

doc.add_paragraph(
    'The manuscript contains approximately 4,500 words, 3 tables, and 4 multi-panel figures '
    '(9 individual maps consolidated into composite panels). All authors have approved the '
    'manuscript and agree with its submission to BMJ Open. The study used publicly available '
    'aggregate data and did not require ethical approval. The manuscript is not under consideration '
    'elsewhere.'
)

doc.add_paragraph(
    'We confirm that this work is original, has not been published previously (except as noted '
    'above regarding the earlier BMJ/RAPM submissions), and is not being considered for publication '
    'elsewhere. All authors have read and approved the final manuscript. There are no conflicts '
    'of interest to declare.'
)

doc.add_paragraph(
    'Thank you for considering our submission. We look forward to your decision.'
)

doc.add_paragraph('')
doc.add_paragraph('Yours sincerely,')
doc.add_paragraph('')
doc.add_paragraph('[Corresponding Author Name]')
doc.add_paragraph('[Affiliation]')
doc.add_paragraph('[Email]')

outpath = '/home/ubuntu/cover_letter_REVISED_EN.docx'
doc.save(outpath)
print(f'Saved: {outpath}')
