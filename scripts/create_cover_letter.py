#!/usr/bin/env python3
"""
Create BMJ cover letter (EN + JP) as docx files.
Follows BMJ submission guidelines for cover letters.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_cover_letter_en():
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    def add_para(text, bold=False, italic=False, align=None, space_after=Pt(6)):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.bold = bold
        run.italic = italic
        if align:
            p.alignment = align
        p.paragraph_format.space_after = space_after
        return p

    # Date
    add_para('[Date]', space_after=Pt(12))

    # Addressee
    add_para('The Editor', space_after=Pt(2))
    add_para('The BMJ', space_after=Pt(2))
    add_para('BMA House, Tavistock Square', space_after=Pt(2))
    add_para('London WC1H 9JR', space_after=Pt(2))
    add_para('United Kingdom', space_after=Pt(12))

    # Salutation
    add_para('Dear Editor,', space_after=Pt(12))

    # Re: line
    add_para(
        'Re: Submission of original research article \u2014 '
        '"Regional variation in anaesthesia practice across 335 secondary medical areas in Japan: '
        'a cross-sectional analysis of national claims data"',
        bold=True, space_after=Pt(12)
    )

    # Body paragraphs
    add_para(
        'We submit the above manuscript for consideration as an original research article in The BMJ. '
        'This work has not been published previously, is not under consideration elsewhere, and all '
        'authors have approved the submitted version.'
    )

    add_para(
        'Japan\u2019s universal health insurance system was designed to ensure equitable access to '
        'healthcare regardless of geography. Yet our analysis of publicly available standardised claim '
        'ratio data across all 335 secondary medical areas reveals striking regional variation in '
        'anaesthesia practice: coefficients of variation of 54.6% for general anaesthesia, 83.2% for '
        'epidural anaesthesia, and 64.9% for combined general\u2013epidural technique.'
    )

    add_para(
        'A distinctive feature of the Japanese system is mandatory prefectural insurance auditing, '
        'which raises the question of whether observed variation is genuine or an administrative '
        'artefact. We address this through three independent sensitivity analyses\u2014within-prefecture '
        'variance decomposition, cross-code correlation analysis, and quantitative audit impact '
        'estimation\u2014all of which converge on the conclusion that audit variation explains, at most, '
        'a trivial fraction of observed differences. The dominant determinant is university hospital '
        'proximity, which alone explains 38.5% of all variation in general anaesthesia volume '
        '(Cohen\u2019s d=1.78)\u2014a remarkably large effect for a single binary predictor.'
    )

    add_para(
        'We believe this study is of particular interest to The BMJ\u2019s readership for three reasons:'
    )

    # Numbered reasons
    add_para(
        '1. It provides a novel sensitivity analysis framework that distinguishes administrative '
        'from clinical sources of practice variation\u2014a methodological contribution applicable '
        'beyond anaesthesia and beyond Japan.'
    )

    add_para(
        '2. The finding that combined general\u2013epidural anaesthesia (which may confer oncological '
        'benefit) shows 1.73-fold variation between university and non-university hospital areas '
        'identifies a potentially modifiable source of health inequality within a universal '
        'insurance system.'
    )

    add_para(
        '3. The study exploits Japan\u2019s unique institutional structure\u2014uniform fee schedule with '
        'prefectural audit committees\u2014as a natural experiment, offering insights relevant to health '
        'systems globally that grapple with distinguishing genuine practice variation from '
        'administrative noise in claims data.'
    )

    add_para(
        'The manuscript adheres to STROBE reporting guidelines for cross-sectional studies. All data '
        'used are publicly available, and the study uses aggregate data only, with no individual-level '
        'patient information. Ethics committee approval was not required under Japan\u2019s Ethical '
        'Guidelines for Medical and Biological Research Involving Human Subjects (2021 revision).'
    )

    add_para(
        'All authors meet the ICMJE criteria for authorship and have no conflicts of interest to '
        'declare. The manuscript contains approximately 4,000 words, 2 tables, and 9 figures '
        '(6 two-dimensional choropleth maps and 3 three-dimensional extruded maps).'
    )

    add_para(
        'We confirm that this manuscript has not been published elsewhere and is not under '
        'consideration by any other journal. All authors have read and approved the final manuscript.'
    )

    add_para(
        'Thank you for considering this submission. We look forward to hearing from you.',
        space_after=Pt(18)
    )

    # Sign-off
    add_para('Yours sincerely,', space_after=Pt(18))

    add_para('[Corresponding author name]', space_after=Pt(2))
    add_para('[Title, Department]', space_after=Pt(2))
    add_para('[Institution]', space_after=Pt(2))
    add_para('[Address]', space_after=Pt(2))
    add_para('[Email]', space_after=Pt(2))
    add_para('[Telephone]', space_after=Pt(2))

    # On behalf of
    add_para('On behalf of all co-authors', italic=True, space_after=Pt(2))

    output_path = '/home/ubuntu/cover_letter_BMJ_EN.docx'
    doc.save(output_path)
    print(f"Saved: {output_path}")


def create_cover_letter_jp():
    doc = Document()

    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    def add_para(text, bold=False, italic=False, align=None, space_after=Pt(6)):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.bold = bold
        run.italic = italic
        if align:
            p.alignment = align
        p.paragraph_format.space_after = space_after
        return p

    # Date
    add_para('[日付]', space_after=Pt(12))

    # Addressee
    add_para('The Editor', space_after=Pt(2))
    add_para('The BMJ', space_after=Pt(2))
    add_para('BMA House, Tavistock Square', space_after=Pt(2))
    add_para('London WC1H 9JR', space_after=Pt(2))
    add_para('United Kingdom', space_after=Pt(12))

    add_para('Dear Editor,', space_after=Pt(12))

    add_para(
        '件名：原著論文の投稿 \u2014 '
        '「日本の335二次医療圏における麻酔診療の地域差：全国レセプトデータを用いた横断的分析」',
        bold=True, space_after=Pt(12)
    )

    add_para(
        '上記の論文をThe BMJへの原著論文として投稿いたします。本論文は他誌に掲載されておらず、'
        '他誌で査読中でもありません。全著者が投稿版を承認しています。'
    )

    add_para(
        '日本の国民皆保険制度は、地理的条件によらず均質な医療アクセスを保障することを'
        '目的として設計されました。しかし、本研究で全335二次医療圏の公開SCR（標準化レセプト出現比）'
        'データを分析した結果、麻酔診療に顕著な地域差が認められました：'
        '変動係数は全身麻酔54.6%、硬膜外麻酔83.2%、GA+硬膜外併用法64.9%でした。'
    )

    add_para(
        '日本の制度には、都道府県単位の保険審査（査定）という特徴的な仕組みがあり、'
        '観察された地域差が真の臨床的差異なのか行政的な人為的産物なのかという問いが生じます。'
        '本研究では3つの独立した感度分析\u2014県内分散分解、コード間相関分析、査定インパクトの'
        '定量的推定\u2014を用いてこの問いに取り組み、いずれの分析も査定による変動では観察された'
        '地域差のごく一部しか説明できないという結論に収束しました。'
        '最も強い決定因子は大学病院の近接性であり、単一の二値変数で全身麻酔量の全分散の'
        '38.5%を説明しました（Cohen\u2019s d=1.78）。'
    )

    add_para(
        '本論文がThe BMJの読者にとって特に有意義である理由は以下の3点です：'
    )

    add_para(
        '1. 診療行為の地域差における行政的要因と臨床的要因を区別する新規の感度分析'
        'フレームワークを提示しており、麻酔科領域や日本に限らず方法論的に応用可能である。'
    )

    add_para(
        '2. 腫瘍学的利益が示唆されるGA+硬膜外併用法が、大学病院圏と非大学病院圏で'
        '1.73倍の差を示し、皆保険制度内における修正可能な健康格差の源泉を同定した。'
    )

    add_para(
        '3. 日本固有の制度的構造（統一的診療報酬点数表と都道府県審査委員会）を自然実験として'
        '活用しており、レセプトデータにおける真の診療変動と行政的ノイズの区別に取り組む'
        '世界各国の医療制度にとって示唆を提供する。'
    )

    add_para(
        '本論文はSTROBE報告ガイドライン（横断研究）に準拠しています。使用した全データは'
        '公開データであり、個人レベルの患者情報は含まれていません。「人を対象とする生命科学・'
        '医学系研究に関する倫理指針」（2021年改正）に基づき、倫理委員会の承認は不要と判断しました。'
    )

    add_para(
        '全著者がICMJE基準による著者資格を満たしており、申告すべき利益相反はありません。'
        '本文は約4,000語、表2つ、図9枚（2次元コロプレスマップ6枚、3次元押出マップ3枚）で'
        '構成されています。'
    )

    add_para(
        '本論文が他誌に掲載済みでなく、他誌で査読中でないことを確認いたします。'
        '全著者が最終版を読み、承認しています。'
    )

    add_para(
        'ご査読のほど、何卒よろしくお願い申し上げます。',
        space_after=Pt(18)
    )

    add_para('敬具', space_after=Pt(18))

    add_para('[責任著者氏名]', space_after=Pt(2))
    add_para('[役職、診療科/部門]', space_after=Pt(2))
    add_para('[所属機関]', space_after=Pt(2))
    add_para('[住所]', space_after=Pt(2))
    add_para('[メールアドレス]', space_after=Pt(2))
    add_para('[電話番号]', space_after=Pt(2))

    add_para('共著者全員を代表して', italic=True, space_after=Pt(2))

    output_path = '/home/ubuntu/cover_letter_BMJ_JP.docx'
    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == '__main__':
    create_cover_letter_en()
    create_cover_letter_jp()
