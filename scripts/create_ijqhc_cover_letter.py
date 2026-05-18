#!/usr/bin/env python3
"""Create English cover letter for IJQHC submission."""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'documents', 'IJQHC')
os.makedirs(OUTPUT_DIR, exist_ok=True)

doc = Document()
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)


def add_para(text, italic=False, bold=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = italic
    run.bold = bold
    return p


# Header block
add_para("[Corresponding author name]", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para("[Affiliation]", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para("[Postal address]", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para("[Email]  |  [Telephone]", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para("[Date]", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
doc.add_paragraph()

# Recipient
add_para("The Editor-in-Chief")
add_para("International Journal for Quality in Health Care")
add_para("Oxford University Press")
doc.add_paragraph()

# Salutation
add_para("Dear Editor,")
doc.add_paragraph()

# Opening
add_para(
    "I am pleased to submit for your consideration the enclosed original "
    "research article entitled \"Regional variation in anaesthesia practice "
    "in Japan: structural determinant or claims-audit artefact?\" for "
    "possible publication in the International Journal for Quality in "
    "Health Care as an Original Research Article.")

# Why this journal
add_para(
    "We believe that our manuscript is well aligned with IJQHC's international "
    "focus on quality, safety and equity of health care, for three reasons. "
    "First, we use Japan's universal insurance system as a natural experiment "
    "to address a question that has been raised - but rarely settled - in "
    "most other health systems: when claims-based variation is observed, is "
    "it a signal about clinical practice or an artefact of differential "
    "auditing? Second, our within-prefecture decomposition and multilevel "
    "sensitivity framework offer a transferable analytic approach that any "
    "system with a uniform fee schedule and regional auditing could adopt. "
    "Third, the variation we identify in combined general-epidural "
    "anaesthesia is clinically meaningful in the context of emerging "
    "evidence on oncological outcomes and speaks to a modifiable equity "
    "issue under universal coverage - directly relevant to IJQHC's "
    "international readership in quality improvement and policy.")

# Key findings
add_para(
    "Key findings: We studied all 335 secondary medical areas of Japan. "
    "Coefficients of variation across areas ranged from 53% (general "
    "anaesthesia) to 87% (epidural anaesthesia). Multilevel modelling "
    "showed that only 5.8% of general anaesthesia variance lay between "
    "prefectures - where audit policy differs - while 94% occurred within "
    "prefectures where audit policy is uniform. University hospital "
    "presence alone explained 35.8% of total variance and was positive in "
    "all 47 of 47 prefectures, with a large effect size (Cohen's d 1.88). "
    "Three pre-specified sensitivity analyses converged in rejecting "
    "differential auditing as a plausible explanation, and empirical Bayes "
    "shrinkage confirmed that the findings are robust to low-volume "
    "instability. We conclude that the observed variation is predominantly "
    "structural and institutional, not administrative, and identify the "
    "variation in combined general-epidural anaesthesia as a priority "
    "equity concern.")

# Compliance declarations
add_para(
    "The manuscript is original, has not been previously published and is not "
    "under consideration for publication elsewhere. All authors have read and "
    "approved the submitted manuscript and have agreed to its submission to "
    "IJQHC. The study used publicly available aggregate data only; ethics "
    "committee approval was not required under the Japanese Ethical "
    "Guidelines for Medical and Biological Research Involving Human Subjects "
    "(2021 revision). The reporting follows the STROBE checklist for "
    "cross-sectional studies, which is uploaded as supplementary material.")

add_para(
    "Statistical analysis: the lead author, who has training in clinical "
    "epidemiology and statistics, performed the analyses and takes full "
    "responsibility for them. No professional statistician outside the "
    "author list was consulted. The contribution of artificial intelligence "
    "tooling (use of Devin / Cognition AI as a coding assistant for data "
    "processing, modelling and visualisation, with all scientific "
    "interpretation performed by the human authors) is declared in the "
    "Contributorship section of the End-Matter file.")

add_para(
    "The manuscript is within the IJQHC length requirements for an Original "
    "Research Article (main text approximately 3000 words; structured "
    "abstract within 400 words; 3 tables, 2 figures; 30 references). We "
    "have uploaded the title page, anonymised main manuscript, End-Matter "
    "(Contributorship, Ethics, Funding, Conflict of interests, "
    "Acknowledgments, Data Availability) and STROBE checklist as separate "
    "files in accordance with the journal's submission requirements.")

add_para(
    "We believe our work will be of substantive interest to IJQHC's "
    "international audience and we look forward to the reviewers' comments. "
    "Please do not hesitate to contact me if any further information is "
    "required.")

doc.add_paragraph()
add_para("Yours sincerely,")
doc.add_paragraph()
add_para("[Signature]")
add_para("[Corresponding author name], on behalf of all authors", italic=True)

out = os.path.join(OUTPUT_DIR, 'cover_letter_IJQHC_EN.docx')
doc.save(out)
print(f"Saved: {out}")
