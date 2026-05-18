#!/usr/bin/env python3
"""Create IJQHC-formatted Japanese (mirror) version of the manuscript.

The English version is the official submission. This file is a Japanese
translation provided to the corresponding author for internal review.
The structure mirrors the English IJQHC version (same headings, same
tables, same number of figures, same references).
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'documents', 'IJQHC')
os.makedirs(OUTPUT_DIR, exist_ok=True)
REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FIG_DIR = os.path.join(REPO_ROOT, 'output')

doc = Document()

for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'MS Mincho'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 2.0
style.paragraph_format.space_after = Pt(0)


def add_run_with_refs(paragraph, text, italic=False, bold=False):
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.name = 'MS Mincho'
        run.font.size = Pt(11)
        if italic:
            run.italic = True
        if bold:
            run.bold = True
        if part.startswith('{') and part.endswith('}'):
            run.text = part[1:-1]
            run.font.superscript = True


def add_para(text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    add_run_with_refs(p, text, italic=italic, bold=bold)
    return p


def add_heading(text, level=1, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'MS Gothic'
    run.font.size = Pt(13) if level == 1 else Pt(12)
    run.bold = True
    return p


def add_subheading(text):
    return add_heading(text, level=2, space_before=8, space_after=4)


def add_blank():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)


# ============================================================
# 表紙
# ============================================================
title_jp = ("日本における麻酔診療の地域差：構造的決定要因か、レセプト審査に由来するアーチファクトか")
title_en = ("Regional variation in anaesthesia practice in Japan: "
            "structural determinant or claims-audit artefact?")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(title_jp)
run.font.name = 'MS Gothic'
run.font.size = Pt(14)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"(英文題名: {title_en})")
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
run.italic = True
add_blank()

add_para("ランニングヘッダー：日本の麻酔診療の地域差", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_blank()

add_para("【著者および所属：別途記載】", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("【責任著者：氏名・住所・電話・メール】", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_blank()

add_para("論文種別：原著論文（Original research article）", italic=True)
add_para("本文語数（要旨・文献・表・図を除く）：英文版で約3,000語", italic=True)
add_para("表：3 / 図：2 / 引用文献：30", italic=True)
add_para("報告ガイドライン：STROBE（横断研究用）チェックリストを補足資料として提出", italic=True)

doc.add_page_break()

# ============================================================
# 構造化抄録
# ============================================================
add_para(title_jp, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_blank()
add_heading("要旨", level=1, space_before=0)

p = doc.add_paragraph()
add_run_with_refs(p, "背景　", bold=True)
add_run_with_refs(
    p,
    "麻酔診療の地域差は多くの高所得国で報告されてきたが、臨床上の選好、施設要因、"
    "事務的なコーディング慣行のいずれが主たる要因かは明らかでない。日本は単一の"
    "全国一律診療報酬体系と都道府県別の保険レセプト審査を併用する独自の自然実験"
    "環境を提供する。我々は、日本の335二次医療圏における麻酔診療の地域差を定量化"
    "し、観察された差が構造的要因によるものか、それとも審査の差に由来する人工"
    "的アーチファクトかを検証することを目的とした。")

p = doc.add_paragraph()
add_run_with_refs(p, "方法　", bold=True)
add_run_with_refs(
    p,
    "横断的生態学的研究。公開されている2022年度の標準化レセプト出現比（年齢・"
    "性別調整した観察値／期待値、居住地ベース、全国平均=100）を用い、47都道府県"
    "に入れ子状に含まれる335二次医療圏で6つの麻酔関連診療報酬コードを解析した。"
    "都道府県をランダム切片とするマルチレベル線形混合モデルにより級内相関係数を"
    "推定し、大学病院所在の効果を評価した。事前指定した3つの感度分析（圏内分散"
    "分解・コード間相関・査定率差の最大影響推定）により審査仮説を検定した。低"
    "症例数医療圏での推定不安定性に対処するため経験ベイズ縮小推定を併用した。")

p = doc.add_paragraph()
add_run_with_refs(p, "結果　", bold=True)
add_run_with_refs(
    p,
    "全コードで顕著な地域差が認められた。変動係数は全身麻酔53.6%、硬膜外麻酔"
    "87.0%、持続硬膜外注入64.9%、脊椎麻酔56.8%であった。マルチレベルモデル"
    "では、全身麻酔の分散のうち都道府県レベルに帰属するのは5.8%（級内相関係数"
    "0.058）に過ぎず、94.2%は審査ポリシーが一定の都道府県内で発生していた。"
    "大学病院所在は全身麻酔の総分散の35.8%、圏内分散の40.5%を説明し（Cohenの"
    "d 1.88）、47都道府県すべてで大学病院所在の二次医療圏のレセプト出現比が"
    "高かった。3つの感度分析はいずれも審査仮説を支持しなかった。最大査定率差"
    "（0.21パーセントポイント）は全身麻酔の四分位範囲の0.3%未満、硬膜外麻酔の"
    "0.1%未満しか説明できなかった。経験ベイズ縮小により効果量は8.6〜9.1%減衰"
    "するに留まった。")

p = doc.add_paragraph()
add_run_with_refs(p, "結論　", bold=True)
add_run_with_refs(
    p,
    "日本の麻酔診療の地域差は大きく、その大部分は保険レセプト審査の差によるもの"
    "ではなく構造的なものである。大学病院の所在が支配的な決定要因であった。"
    "腫瘍学的アウトカムとの関連が示唆されている全身麻酔・硬膜外麻酔併用の"
    "顕著な地域差は、皆保険を前提とする日本において、本来必要とされる医療技術"
    "へのアクセス公平性を損ないうる修正可能な不平等であることを示唆する。")

add_blank()
p = doc.add_paragraph()
add_run_with_refs(p, "キーワード：", bold=True)
add_run_with_refs(p, "麻酔；小地域差；マルチレベル分析；医療サービス研究；医療の質；日本")

doc.add_page_break()

# ============================================================
# 緒言
# ============================================================
add_heading("緒言", level=1, space_before=0)

add_para(
    "1961年に確立された日本の国民皆保険制度は、全人口を対象に全国一律の診療報酬"
    "体系のもとで医療を提供してきた。{1}日本は実効的な医療カバレッジで世界最高"
    "水準の達成度を示しており{2}、医療提供体系は47の三次医療圏（都道府県）と335"
    "の二次医療圏（入院医療を地域内で完結させることを目標とする計画単位）からなる"
    "階層構造として組織化されている。{3}都道府県を超えた患者移動は限定的であり、"
    "2008年患者調査では入院患者のうち県外で受療したのは6.1%に過ぎない。{4}"
    "本制度は地理的条件によらず公平な医療アクセスを提供することを意図して設計"
    "された。")

add_para(
    "しかし、レセプト情報・特定健診等情報データベース（NDB）を用いた研究は、"
    "外科手術・処方・診断検査などにおいて、依然として大きな地域差が残存している"
    "ことを示している。{5}日本の制度の特徴の一つは、都道府県別の保険レセプト"
    "審査（社会保険診療報酬支払基金等による事前審査と査定）であり、査定率は"
    "0.07%から0.28%まで都道府県間で変動する。{6}このことは、医療の質と安全"
    "の研究にとって本質的な問いを提起する。すなわち、レセプトデータで観察"
    "される地域差は臨床上の真の差を反映しているのか、それとも審査の差による"
    "人工的アーチファクトなのか。両者の区別は、医療の質改善・支払い制度・"
    "医療公平性のいずれにおいても異なる政策含意をもつ。")

add_para(
    "麻酔診療はこの問いに適した対象である。全身麻酔単独と全身・区域麻酔併用との"
    "選択は臨床的に重要であり、特に硬膜外鎮痛と全身麻酔の併用は一部のがん手術"
    "における無再発生存・全生存を改善しうるとする研究もあるが、{7,8}ランダム化"
    "比較試験の結果は一致していない。{9}2026年度の診療報酬改定では、長年指摘"
    "されてきた気道確保のあいまいさを反映して全身麻酔コードの表記が改訂され"
    "た。{10}観察された地域差が構造的なものか審査由来のものかを明らかにする"
    "ことは、皆保険下の医療公平性を評価するうえで時宜に適った課題である。")

add_para(
    "我々は、二次医療圏単位の標準化レセプト出現比を用いて、以下の研究課題に"
    "取り組んだ。(i) 335二次医療圏における麻酔診療の地域差はどれほど大きいか。"
    "(ii) 観察された差は都道府県別の審査差によって説明されるか。(iii) 都道府県"
    "間と都道府県内のいずれに、より大きな分散があるか。(iv) 大学病院所在を"
    "含むどの構造要因が観察された差と関連するか。")

# ============================================================
# 方法
# ============================================================
add_heading("方法", level=1)

add_subheading("研究デザインと倫理事項")
add_para(
    "本研究は公開集計データのみを用いた横断的生態学的研究である。個人レベル"
    "データへのアクセスは行っていない。「人を対象とする生命科学・医学系研究"
    "に関する倫理指針」（2021年改訂）の対象外であり、倫理委員会の承認は要しな"
    "かった。本研究はSTROBE（横断研究用）チェックリストに準拠して報告した。"
    "{30}チェックリストは補足資料として添付している。")

add_subheading("地理的単位")
add_para(
    "日本の医療提供体系は医療法第30条の4に基づき階層構造で組織化されている。"
    "{3}47都道府県が三次医療圏に対応し、各都道府県内で二次医療圏（全国計335）"
    "が入院医療計画の基本単位として指定されている。重要な点として、二次医療圏"
    "は都道府県内に完全に入れ子状に収まり、都道府県境を越えることはない。"
    "これは米国のHospital Service AreaやHospital Referral Region（州境を越え"
    "うる）とは異なる構造であり、{11}審査ポリシーが一定の都道府県内での比較"
    "が、流出効果の交絡を受けずに解釈できることを意味する。")

add_subheading("データソースと標準化レセプト出現比")
add_para(
    "3つのデータソースを使用した。第一に、内閣府「経済財政諮問会議　地域差の"
    "「見える化」」が公開する2022年度の標準化レセプト出現比。{12}標準化レセプト"
    "出現比は間接標準化により算出される。各診療コードについて、全国の年齢・"
    "性別ごとのレセプト出現率を地域の人口構造に適用して期待値を算出し、観察値"
    "／期待値×100として求める。これらの比は居住地ベースで算出されるため、"
    "受療のための他地域への移動は受療地域の比を上昇させない。低症例数の医療圏"
    "は提供元によりマスクされて欠損値となる。データは都道府県および二次医療圏"
    "の双方の単位で利用可能である。第二に、2022年医師・歯科医師・薬剤師統計"
    "（e-Stat）から二次医療圏別の麻酔科医・総医師数を取得した。{13}第三に、"
    "国土数値情報から地理境界データを取得した。{14}")

add_subheading("麻酔関連診療報酬コードと大学病院マッピング")
add_para(
    "診療報酬点数表から6つの麻酔関連コードを解析対象とした（表1）：L008（"
    "閉鎖循環式全身麻酔）、L002（硬膜外麻酔）、L003（持続硬膜外注入。全身・"
    "硬膜外併用の直接指標）、L004（脊椎麻酔）、L009（麻酔管理料I。麻酔科専門医"
    "配置の代理指標）、L100（神経ブロック・入院）。81大学病院（国立44、公立8、"
    "私立29）を所在市区町村に基づき二次医療圏にマッピングした。81病院は335二次"
    "医療圏のうち64医療圏（19.1%）に分布し、47都道府県すべてに存在する。これに"
    "より、共通の審査ポリシー下にある都道府県内で大学病院所在医療圏と非所在"
    "医療圏を比較することが可能となる。")

add_subheading("統計解析")
add_para(
    "3つの相補的アプローチを用いた。第一に、変動係数を含む記述統計を算出した。"
    "群間比較にはWelchのt検定、Cohenのdによる効果量、{15}補助的にMann–Whitney"
    "のU検定を用いた。第二に、二次医療圏（レベル1、n=335）が都道府県（レベル2、"
    "n=47）に入れ子状に含まれる構造に対応するため、マルチレベル線形混合モデルを"
    "適合した。{16}まずランダム切片のみのヌルモデルにより級内相関係数を推定し、"
    "次に大学病院所在（二値）、麻酔科医密度（標準化）、総医師密度（標準化）を"
    "固定効果として加えたモデルを推定した。推定はPython statsmodelsのMixedLM"
    "実装による制限付き最尤法を用いた。{17}周辺R²はヌルモデルからの総分散減少"
    "比として算出した。低症例数医療圏での比の不安定性に対処するため経験ベイズ"
    "縮小推定{18}を併用し、生の比と縮小後の比の双方で主結果を比較した。")

add_para(
    "第三に、保険レセプト審査が観察された差を説明するという帰無仮説を検定する"
    "ため、3つの事前指定感度分析を実施した。(a) 圏内分散分解：都道府県内で"
    "審査が一律であれば、都道府県内分散は小さいはず。一元配置分散分析により"
    "総分散を都道府県間と都道府県内に分解し、さらに都道府県内分散を大学病院"
    "効果と残差に階層的に分解した。(b) コード間相関：審査による全身麻酔と脊椎"
    "麻酔のコード入れ替えが起きていれば、両者は負相関を示すはず。(c) 査定率差"
    "の最大影響推定：観察された都道府県間査定率差で説明可能な比のシフトの上限"
    "を算出した。代替可能なコード間で審査による入れ替えが地域差の主因であれば、"
    "両者の合算は変動係数を大きく低下させるはずであり、L008+L004、L008+L002の"
    "合算比も検討した。なお、本研究の計画・実施・報告には患者・市民の参画はない。")

# ============================================================
# 結果
# ============================================================
add_heading("結果", level=1)

add_subheading("研究対象と麻酔診療の地域差")
add_para(
    "335二次医療圏は47都道府県に分布した（都道府県あたりの中央値7、範囲3〜21）。"
    "64医療圏（19.1%）に少なくとも1つの大学病院が所在した。L008の標準化レセプト"
    "出現比は335医療圏中334で利用可能、L002は307医療圏で利用可能（28医療圏は"
    "低症例によりマスク）。全コードで顕著な地域差が観察された（表2）。L008は"
    "2.3〜435.7の範囲（変動係数53.6%、最小〜最大で189倍）、L002の変動係数は"
    "87.0%と最大、L003は64.9%、L004は56.8%であった。地理的分布を図1に示す。")

add_subheading("マルチレベルモデルと大学病院効果")
add_para(
    "ヌルマルチレベルモデルでは、L008の総分散のうち都道府県レベルに帰属するのは"
    "わずか5.8%（級内相関係数0.058）であり、94.2%は審査ポリシーが一定の都道府県"
    "内で発生していた。L004は都道府県レベルクラスタリングが強かった（級内相関"
    "係数0.343）。大学病院所在を固定効果として追加するとL008で最大の改善が"
    "得られた（周辺R² 0.358、β=+64.1、95%信頼区間54.8〜73.4）。すべての"
    "コードで統計的に有意（すべてP<0.001）であった（表3）。47都道府県すべて"
    "において、大学病院所在医療圏のL008出現比が非所在医療圏より高く、都道府県"
    "内の平均差は+63.3点（対応のあるt=13.28、P<0.001）であった（図2）。"
    "大学病院効果のCohenのdはL008で1.88、L003で1.30、L002で0.59、L004で0.40"
    "であった。")

add_subheading("審査仮説に対する感度分析")
add_para(
    "3つの事前指定感度分析はいずれも審査仮説を支持しなかった。第一に、L008の"
    "総分散のうち都道府県間差は14.5%、大学病院効果（都道府県内）は40.5%、"
    "残差は45.0%を占めた。大学病院効果のみで都道府県内分散の47.4%を説明した。"
    "第二に、L008とL004は負ではなく正の相関を示し（r=+0.235、P<0.001）、"
    "L008とL002も正の相関（r=+0.319、P<0.001）であった。これらは審査による"
    "コード入れ替えと矛盾し、共通の供給要因（麻酔科の人的資源・設備等）を"
    "示唆する。第三に、観察された最大査定率差0.21パーセントポイントが説明"
    "しうる比のシフトは約0.2点であり、L008の四分位範囲の0.3%未満、L002の"
    "0.1%未満であった。L008+L004の合算は変動係数を53.6%から43.4%に減らした"
    "（19%減少）にとどまり、合算指標でも大学病院効果は依然大きかった"
    "（Cohenのd 1.26）。")

add_subheading("頑健性：経験ベイズ縮小と外れ値")
add_para(
    "都道府県平均への経験ベイズ縮小により効果量はわずかに減衰したが、いずれの"
    "結論も変わらなかった。L008の大学病院効果はCohenのd=1.89（生）から1.72"
    "（縮小）へ9.0%減衰、L004は9.1%減衰（d 0.40→0.36）、L002は8.6%減衰"
    "（d 0.59→0.54）であった。すべての効果量は依然大きく、統計的に有意で"
    "あった。L008平均から3標準偏差以上離れた外れ値として千葉県の安房医療圏"
    "（出現比216.7）と東京都区中央部医療圏（出現比435.7、大学病院5施設）が"
    "同定されたが、これらを除外しても結論の方向性・有意性は変化しなかった。")

# ============================================================
# 考察
# ============================================================
add_heading("考察", level=1)

add_subheading("主要所見")
add_para(
    "335二次医療圏における麻酔診療の地域差は実質的であり、変動係数は全身麻酔の"
    "53%から硬膜外麻酔の87%まで及ぶ。マルチレベルモデルでは、L008分散のうち"
    "都道府県レベル由来はわずか5.8%にとどまり、94.2%は審査ポリシーが一定の"
    "都道府県内で発生していた。大学病院所在は総分散の35.8%、都道府県内分散の"
    "40.5%を説明し、しかも審査運用が異なる47都道府県すべてで一貫して認められ"
    "た。3つの独立した感度分析はいずれも審査差を地域差の主因として支持しな"
    "かった。したがって、観察された地域差は事務的なコーディングではなく、"
    "施設能力と臨床・組織要因に駆動される構造的なものと結論される。")

add_subheading("強みと限界")
add_para(
    "強みとして、年齢・性別調整された全国規模指標の使用、335医療圏という細かな"
    "地理単位での解析、入れ子データ構造を適切に扱うマルチレベルモデリング、"
    "低症例不安定性に対処する経験ベイズ縮小、複数の事前指定感度分析、審査"
    "ポリシーを一定とする都道府県内比較、患者移動を緩和する居住地ベース比の"
    "採用が挙げられる。限界として、地域単位解析に伴う生態学的誤謬、{19}横断"
    "デザインゆえの因果推論の限界、レセプトが審査後（実施意図ではなく算定後）"
    "の状態を反映すること、大学病院所在と都市集中の部分的重複、コード別査定率"
    "データの欠如、レセプト単独からは「ディフェンシブな低算定」を定量化できない"
    "ことが挙げられる。マルチレベルモデルには大学病院所在のみを構造的固定"
    "効果として含めた。今後は都市化指標、病床密度、症例構成など追加の地域共変量"
    "を取り入れる必要がある。")

add_subheading("先行研究との位置づけ")
add_para(
    "本研究の所見は医療の地域差研究の先行知見と整合的である。米国のDartmouth "
    "Atlasは医師供給と診療スタイルが手術率の地域差を主に駆動することを示し"
    "た。{21}英国NHS、{22}ドイツ、{23}オーストラリア{24}でも同様の知見が"
    "報告されている。本研究は、全国一律の診療報酬体系と都道府県別の審査運用が"
    "並存する日本特有の構造を活用して、事務的要因と臨床的要因を分離した点、"
    "および皆保険下の東アジアにおいて麻酔技術の小地域差をマルチレベルで解析した"
    "初の研究である点で先行研究を拡張する。本研究で示された大学病院効果（"
    "Cohenのd=1.88）は、この分野で報告される効果量と比べて大きい。これは関連"
    "病院の臨床に対する大学医局制度の影響を反映している可能性がある。{25}"
    "またL008の級内相関係数0.058は、米国の手術率で報告される値と比べて低く、"
    "{21}全身麻酔については都道府県内の施設要因が都道府県レベル要因を凌駕する"
    "ことを示している。")

add_para(
    "本研究の知見と枠組みは日本国外にも4つの点で意義を持つ。第一に、本研究で"
    "用いた圏内分散分解は、一律の診療報酬体系と地域単位の請求審査を併せもつ"
    "他の皆保険制度—台湾NHI、{20}韓国NHIS、ドイツ・フランスの社会保険、"
    "英国NHS{22,23}—にそのまま移植可能である。第二に、構造的決定因としての"
    "高次医療機関（大学・ティーチング病院）集中の優越は、米国・英国の"
    "ティーチング病院における麻酔手技構成の差異{21,22}、ならびにOECD各国の"
    "他の高次医療集中型診療科{26}でも報告されており、47/47都道府県で一貫して"
    "観察された本研究の大学病院効果は日本固有ではなく、高所得国の医療制度に"
    "普遍的な特徴である可能性を示唆する。第三に、全身・硬膜外麻酔併用の"
    "地域差は国際的なequity agendaに関わる：区域麻酔が腫瘍学的・機能的"
    "アウトカムに関連しうるという萌芽的エビデンス{7-9}に照らせば、麻酔技術"
    "分布のモニタリングは、技術集約的医療への公平なアクセスを標榜する"
    "あらゆる皆保険制度にとってシステムレベル質指標となりうる。{27,29}"
    "第四に、低中所得国への外挿には慎重を要する。診療報酬体系や審査運用が"
    "不均一であり、レセプトデータが不完全なことも多く、分散分解の直接的な"
    "適用は限定的である。とはいえ、「見かけの小地域差を政策対応の前に"
    "事務的要因か構造的要因かで吟味する」という本研究の一般原則は、皆保険"
    "改革の中で診療報酬体系の標準化が進む国々にも移植可能である。{2}")

add_subheading("政策・実践・研究へのインプリケーション")
add_para(
    "L003（持続硬膜外注入）の変動係数64.9%という大きな地域差は、全身・区域"
    "麻酔併用が一部のがん手術における腫瘍学的アウトカムに関連しうるという"
    "知見{7-9}に照らし、臨床的に重要である。L003出現比が大学病院所在医療圏で"
    "1.73倍であることは、潜在的に有益な技術へのアクセスが居住地に依存する"
    "ことを意味し、皆保険の公平性の前提を損ないうる。質改善の取り組みとして、"
    "ターゲットを絞った教育、大学病院からの専門医アウトリーチ、麻酔技術分布の"
    "システムレベル質指標としてのモニタリング{27,29}が、麻酔の安全な実践の"
    "国際基準{28}と整合した形で考えうる。本研究で開発した圏内分散分解と"
    "マルチレベル感度分析の枠組みは、皆保険下の他の診療領域にも適用可能で"
    "あり、事務的要因と構造的要因を区別するための汎用的ツールを提供する。")

# ============================================================
# 結論
# ============================================================
add_heading("結論", level=1)
add_para(
    "日本の麻酔診療の地域差は大きく、その大部分は保険レセプト審査の差ではなく"
    "構造的なものである。大学病院所在は支配的な決定要因であり、L008の総分散の"
    "三分の一以上を説明する。腫瘍学的アウトカムとの関連が示唆される全身・"
    "硬膜外併用麻酔における顕著な差は特に懸念される。質改善および公平性政策"
    "は、地域差をレセプト処理の不可避な特徴として扱うのではなく、麻酔診療の"
    "供給・組織要因に取り組むべきである。")

# ============================================================
# 引用文献（英文と同じ番号順）
# ============================================================
add_heading("引用文献", level=1)

REFERENCES = [
    "Ikegami N, Yoo BK, Hashimoto H, et al. Japanese universal health coverage: evolution, achievements, and challenges. Lancet 2011;378:1106-15.",
    "GBD 2019 Universal Health Coverage Collaborators. Measuring universal health coverage based on an index of effective coverage of health services in 204 countries and territories, 1990-2019. Lancet 2020;396:1250-84.",
    "Matsuda S. Health policy in Japan: current situation and future challenges. JMA J 2019;2:1-10.",
    "Ministry of Health, Labour and Welfare. Patient Survey 2008. Tokyo: MHLW, 2009.",
    "Hashimoto H, Ikegami N, Shibuya K, et al. Cost containment and quality of care in Japan: is there a trade-off? Lancet 2011;378:1174-82.",
    "Cabinet Office. Indicators of regional variation in healthcare and long-term care: insurance auditing rates by prefecture. Tokyo: Cabinet Office, 2023.",
    "Sessler DI, Pei L, Huang Y, et al. Recurrence of breast cancer after regional or general anaesthesia: a randomised controlled trial. Lancet 2019;394:1807-15.",
    "Chen WK, Miao CH. The effect of anesthetic technique on survival in human cancers: a meta-analysis of retrospective and prospective studies. PLoS One 2013;8:e56540.",
    "Weng M, Chen W, Hou W, et al. The effect of neuraxial anaesthesia on cancer recurrence and survival after cancer surgery: an updated meta-analysis. Oncotarget 2016;7:15262-73.",
    "Ministry of Health, Labour and Welfare. 2026 revision of the medical fee schedule: notice on anaesthesia code revisions. Tokyo: MHLW, 2025.",
    "Wennberg JE, Cooper MM, eds. The Dartmouth Atlas of Health Care. Chicago: American Hospital Publishing, 1999.",
    "Cabinet Office. Regional variation visualisation (chiikisa no mieruka). https://www5.cao.go.jp/keizai-shimon/kaigi/special/reform/mieruka/ (accessed 10 Mar 2026).",
    "Ministry of Health, Labour and Welfare. Survey of Physicians, Dentists and Pharmacists 2022 (e-Stat). https://www.e-stat.go.jp (accessed 12 Mar 2026).",
    "Ministry of Land, Infrastructure, Transport and Tourism. National Land Numerical Information download service. https://nlftp.mlit.go.jp (accessed 10 Mar 2026).",
    "Cohen J. Statistical Power Analysis for the Behavioral Sciences. 2nd ed. Hillsdale, NJ: Lawrence Erlbaum, 1988.",
    "Rabe-Hesketh S, Skrondal A. Multilevel and Longitudinal Modeling Using Stata. 3rd ed. College Station, TX: Stata Press, 2012.",
    "Seabold S, Perktold J. Statsmodels: econometric and statistical modeling with Python. In: Proceedings of the 9th Python in Science Conference. 2010:92-6.",
    "Efron B, Morris C. Stein's estimation rule and its competitors: an empirical Bayes approach. J Am Stat Assoc 1973;68:117-30.",
    "Greenland S. Ecologic versus individual-level sources of bias in ecologic estimates of contextual health effects. Int J Epidemiol 2001;30:1343-50.",
    "Cheng T-M. Taiwan's new national health insurance program: genesis and experience so far. Health Aff (Millwood) 2003;22:61-76.",
    "Wennberg JE, Fisher ES, Skinner JS. Geography and the debate over Medicare reform. Health Aff (Millwood) 2002;Suppl Web Exclusives:W96-114.",
    "NHS RightCare. The NHS Atlas of Variation in Healthcare. London: Public Health England, 2015.",
    "Wengler A, Nimptsch U, Mansky T. Hip and knee replacement surgery (arthroplasty) and geographic variations in Germany. Dtsch Arztebl Int 2014;111:407-16.",
    "Australian Commission on Safety and Quality in Health Care. Australian Atlas of Healthcare Variation. Sydney: ACSQHC, 2015.",
    "Otsuka T. The ikyoku system of university orthopaedic surgery departments: an in-hospital organisational system unique to Japan. J Orthop Sci 2012;17:513-14.",
    "OECD. Geographic Variations in Health Care: What Do We Know and What Can Be Done to Improve Health System Performance? Paris: OECD Publishing, 2014.",
    "Donabedian A. The quality of care: how can it be assessed? JAMA 1988;260:1743-8.",
    "Merry AF, Cooper JB, Soyannwo O, et al. International standards for a safe practice of anesthesia 2010. Can J Anesth 2010;57:1027-34.",
    "Mainz J. Defining and classifying clinical indicators for quality improvement. Int J Qual Health Care 2003;15:523-30.",
    "von Elm E, Altman DG, Egger M, et al. The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement: guidelines for reporting observational studies. Lancet 2007;370:1453-7.",
]
for i, ref in enumerate(REFERENCES, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{i}. {ref}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

# ============================================================
# 表（引用文献の後）
# ============================================================
doc.add_page_break()
add_heading("表", level=1, space_before=0)

add_para("表1．解析対象とした麻酔関連診療報酬コード", bold=True)
table1 = doc.add_table(rows=1, cols=3)
table1.style = 'Table Grid'
hdr = table1.rows[0].cells
hdr[0].text = "コード"
hdr[1].text = "名称"
hdr[2].text = "備考"
for cell in hdr:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'MS Gothic'
table1_rows = [
    ("L008", "閉鎖循環式全身麻酔", "全身麻酔の総量を表す主要指標"),
    ("L002", "硬膜外麻酔", "区域麻酔技術の指標"),
    ("L003", "持続硬膜外注入", "全身・硬膜外併用麻酔の直接指標"),
    ("L004", "脊椎麻酔", "代替的な区域麻酔技術"),
    ("L009", "麻酔管理料I", "麻酔科専門医配置の代理指標"),
    ("L100", "神経ブロック・入院", "ペインクリニック活動の指標"),
]
for code, name, note in table1_rows:
    row = table1.add_row().cells
    row[0].text = code
    row[1].text = name
    row[2].text = note
    for cell in row:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'MS Mincho'
                run.font.size = Pt(10)
add_blank()
add_para("出典：診療報酬点数表（2022年度）。各コードについて全国平均=100となるよう"
         "年齢・性別を間接標準化した標準化レセプト出現比を算出した。", italic=True)
doc.add_page_break()

add_para("表2．335二次医療圏における標準化レセプト出現比の分布"
         "（全国平均=100、2022年度）", bold=True)
table2 = doc.add_table(rows=1, cols=8)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
for i, h in enumerate(["コード", "n", "平均（標準偏差）", "中央値（四分位範囲）",
                        "最小", "最大", "変動係数(%)", "最大／最小"]):
    hdr2[i].text = h
    for p in hdr2[i].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = 'MS Gothic'
table2_rows = [
    ("L008", "334", "85.7 (45.9)", "76.5 (53.0–110.5)", "2.3", "435.7", "53.6", "189"),
    ("L002", "307", "61.4 (53.4)", "47.7 (24.5–83.4)", "0", "519.1", "87.0", "—"),
    ("L003", "307", "61.4 (39.8)", "53.5 (32.0–82.0)", "0", "224.7", "64.9", "—"),
    ("L004", "320", "82.5 (46.8)", "76.5 (49.5–108.5)", "0", "318.0", "56.8", "—"),
    ("L009", "330", "78.3 (43.0)", "71.0 (47.5–101.5)", "0", "230.0", "55.0", "—"),
    ("L100", "324", "60.5 (35.7)", "53.5 (35.5–78.0)", "0", "247.0", "59.0", "—"),
]
for r in table2_rows:
    row = table2.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val
        for p in row[i].paragraphs:
            for run in p.runs:
                run.font.name = 'MS Mincho'
                run.font.size = Pt(9)
add_blank()
add_para("最大／最小は最小値が0でない場合のみ示した。低症例数医療圏は提供元によりマスク"
         "され欠損となるため、各コードのn欄は非欠損医療圏数を反映している。",
         italic=True)
doc.add_page_break()

add_para("表3．マルチレベル線形混合モデル：標準化レセプト出現比をアウトカム、"
         "都道府県をランダム切片", bold=True)
table3 = doc.add_table(rows=1, cols=5)
table3.style = 'Table Grid'
hdr3 = table3.rows[0].cells
for i, h in enumerate(["コード", "ヌルモデルICC",
                        "β大学（95%CI）", "P値", "周辺R²"]):
    hdr3[i].text = h
    for p in hdr3[i].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'MS Gothic'
table3_rows = [
    ("L008", "0.058", "+64.1（54.8〜73.4）", "<0.001", "0.358"),
    ("L002", "0.082", "+49.0（38.4〜59.6）", "<0.001", "0.205"),
    ("L003", "0.071", "+55.8（45.2〜66.4）", "<0.001", "0.245"),
    ("L004", "0.343", "+23.0（12.4〜33.6）", "<0.001", "0.045"),
    ("L009", "0.205", "+47.6（36.4〜58.8）", "<0.001", "0.180"),
    ("L100", "0.299", "+18.4（8.6〜28.2）", "<0.001", "0.038"),
]
for r in table3_rows:
    row = table3.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val
        for p in row[i].paragraphs:
            for run in p.runs:
                run.font.name = 'MS Mincho'
                run.font.size = Pt(10)
add_blank()
add_para("β大学：大学病院所在（0/1）の固定効果係数；CI：信頼区間；ICC：級内相関係数"
         "（ヌルランダム切片モデルにおける都道府県レベルへの分散帰属比）；周辺R²：大学病院"
         "固定効果追加によるヌルモデルからの総分散減少比。推定は制限付き最尤法による。",
         italic=True)
doc.add_page_break()

# ============================================================
# 図凡例
# ============================================================
add_heading("図凡例", level=1, space_before=0)
add_para("図1．335二次医療圏における麻酔の標準化レセプト出現比の地理分布（2022年度）。"
         "(A) 閉鎖循環式全身麻酔（L008）。(B) 脊椎麻酔（L004）。標準化レセプト出現比"
         "（全国平均=100）の五分位で陰影付け。低症例によりマスクされた医療圏は灰色で示す。")
add_blank()
add_para("図2．大学病院所在と全身麻酔標準化レセプト出現比。(A) 少なくとも1つの大学病院"
         "を含む二次医療圏（335医療圏中64）の分布。(B) 47都道府県すべてにおける大学病院"
         "所在医療圏と非所在医療圏の平均L008出現比の都道府県内比較。47/47都道府県で"
         "大学病院効果が正であった（都道府県内平均差+63.3点、対応のあるt=13.28、P<0.001）。")

for fig_path, caption in [
    (os.path.join(FIG_DIR, 'rapm_fig1_jp.png'), '図1．'),
    (os.path.join(FIG_DIR, 'rapm_fig2_jp.png'), '図2．'),
]:
    if os.path.exists(fig_path):
        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.bold = True
        run.font.name = 'MS Gothic'
        run.font.size = Pt(12)
        doc.add_picture(fig_path, width=Inches(6.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

out = os.path.join(OUTPUT_DIR, 'regional_anaesthesia_IJQHC_JP.docx')
doc.save(out)
print(f"Saved: {out}")
