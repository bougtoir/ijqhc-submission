#!/usr/bin/env python3
"""
Create BJA (British Journal of Anaesthesia) cover letter (EN) as docx.
Tailored to BJA submission guidelines:
- Clinical investigations: max 3000 words, 40 refs, 6 tables+figures
- Our paper exceeds 6 tables+figures (2 tables + 9 figures = 11),
  so we justify this in the cover letter and propose supplementary material.
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_cover_letter_bja():
    doc = Document()

    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    def add_para(text, bold=False, italic=False, align=None, space_after=Pt(6)):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.bold = bold
        run.italic = italic
        if align:
            p.alignment = align
        p.paragraph_format.space_after = space_after
        return p

    # Date
    add_para('[Date]', space_after=Pt(12))

    # Addressee
    add_para('Professor Hugh Hemmings', space_after=Pt(2))
    add_para('Editor-in-Chief', space_after=Pt(2))
    add_para('British Journal of Anaesthesia', space_after=Pt(2))
    add_para('Weill Cornell Medical College', space_after=Pt(2))
    add_para('New York, NY, USA', space_after=Pt(2))
    add_para('Email: bja@med.cornell.edu', space_after=Pt(12))

    # Salutation
    add_para('Dear Professor Hemmings,', space_after=Pt(12))

    # Re: line
    add_para(
        'Re: Submission of clinical investigation \u2014 '
        '"Regional variation in anaesthesia practice across 335 secondary medical areas in Japan: '
        'a cross-sectional analysis of national claims data"',
        bold=True, space_after=Pt(12)
    )

    # Body
    add_para(
        'We submit the above manuscript for consideration as a clinical investigation in the '
        'British Journal of Anaesthesia. This work has not been published previously, is not under '
        'consideration elsewhere, and all authors have approved the submitted version.'
    )

    add_para(
        'Using publicly available standardised claim ratio (SCR) data covering the entire Japanese '
        'population, we quantify regional variation in anaesthesia practice across all 335 secondary '
        'medical areas. We report coefficients of variation of 54.6% for general anaesthesia, 83.2% '
        'for epidural anaesthesia, and 64.9% for the combined general\u2013epidural technique\u2014variation '
        'that is substantially greater than previously recognised in any national healthcare system.'
    )

    add_para(
        'A unique feature of this study is the exploitation of Japan\u2019s institutional structure\u2014a '
        'uniform national fee schedule combined with prefectural insurance audit committees\u2014to '
        'address a question that is fundamental to all claims-based practice variation research: '
        'does observed variation reflect genuine differences in clinical practice, or is it an '
        'artefact of differential administrative auditing? Three independent sensitivity analyses '
        '(within-prefecture variance decomposition, cross-code correlation, and quantitative audit '
        'impact estimation) converge on the conclusion that audit variation explains less than 1% of '
        'observed differences.'
    )

    add_para(
        'The dominant determinant is university hospital proximity. In Japan\u2019s medical education '
        'system, university departments (ikyoku) exert substantial influence over clinical practice '
        'in affiliated hospitals. We show that university hospital presence explains 38.5% of all '
        'variance in general anaesthesia volume (Cohen\u2019s d=1.78)\u2014a remarkably large effect that '
        'is consistent across all 47 prefectures. The combined GA\u2013epidural technique is 1.73 times '
        'more common in university hospital areas.'
    )

    add_para(
        'We believe this study is of particular relevance to the BJA\u2019s readership for three reasons:'
    )

    add_para(
        '1. It documents the largest geographic variation in anaesthesia technique reported to date '
        'at a fine spatial scale (335 areas), with direct implications for anaesthesia workforce '
        'planning and training policy.'
    )

    add_para(
        '2. Given the growing body of evidence suggesting oncological benefit from combined '
        'regional\u2013general anaesthesia, the 1.73-fold variation in this technique represents a '
        'potentially modifiable source of outcome inequality that is directly relevant to '
        'perioperative medicine.'
    )

    add_para(
        '3. The sensitivity analysis framework we develop\u2014distinguishing administrative from '
        'clinical sources of variation using within-jurisdiction variance decomposition\u2014is a '
        'methodological contribution applicable to anaesthesia practice variation research in any '
        'country with claims-based data.'
    )

    # Address the table/figure excess
    add_para(
        'We note that the manuscript includes 2 tables and 9 figures, which exceeds the BJA\u2019s '
        'guideline of 6 tables and figures combined for clinical investigations. We believe this is '
        'justified because the geographic nature of the analysis requires spatial visualisation: '
        'Figures 1\u20136 are two-dimensional choropleth maps of individual anaesthesia codes across 335 '
        'areas, and Figures 7\u20139 are three-dimensional extruded maps encoding two variables '
        'simultaneously (colour and height). We would be happy to move selected figures to '
        'supplementary material if the Editor prefers, or to consolidate figures into multi-panel '
        'composites.'
    )

    add_para(
        'The study adheres to STROBE reporting guidelines for cross-sectional studies and uses '
        'RECORD (REporting of studies Conducted using Observational Routinely collected Data) '
        'extensions where applicable. All data used are publicly available aggregate statistics; '
        'no individual-level patient data were accessed. Ethics committee approval was not required '
        'under Japan\u2019s Ethical Guidelines for Medical and Biological Research Involving Human '
        'Subjects (2021 revision).'
    )

    add_para(
        'All authors meet the ICMJE criteria for authorship. We have no conflicts of interest to '
        'declare and no external funding to report for this study.'
    )

    add_para(
        'Thank you for considering this submission. We look forward to your response.',
        space_after=Pt(18)
    )

    # Sign-off
    add_para('Yours sincerely,', space_after=Pt(18))

    add_para('[Corresponding author name]', space_after=Pt(2))
    add_para('[Title, Department]', space_after=Pt(2))
    add_para('[Institution]', space_after=Pt(2))
    add_para('[Address]', space_after=Pt(2))
    add_para('[Email]', space_after=Pt(2))
    add_para('[ORCID]', space_after=Pt(2))

    add_para('On behalf of all co-authors', italic=True, space_after=Pt(2))

    output_path = '/home/ubuntu/cover_letter_BJA_EN.docx'
    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == '__main__':
    create_cover_letter_bja()
