#!/usr/bin/env python3
"""Create separate editable tables .docx (EN/JP) for IJQHC submission."""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'documents', 'IJQHC')
os.makedirs(OUTPUT_DIR, exist_ok=True)


def setup(doc, jp=False):
    for s in doc.sections:
        s.page_width = Cm(29.7)
        s.page_height = Cm(21)
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2)
        s.right_margin = Cm(2)
    st = doc.styles['Normal']
    st.font.name = 'MS Mincho' if jp else 'Times New Roman'
    st.font.size = Pt(11)
    st.paragraph_format.line_spacing = 1.15
    st.paragraph_format.space_after = Pt(0)


def add_title(doc, text, jp=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'MS Gothic' if jp else 'Times New Roman'
    r.font.size = Pt(12)


def add_footnote(doc, text, jp=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.italic = True
    r.font.name = 'MS Mincho' if jp else 'Times New Roman'


def make_table(doc, headers, rows, jp=False):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = 'MS Gothic' if jp else 'Times New Roman'
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            c = t.rows[ri].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(11)
            r.font.name = 'MS Mincho' if jp else 'Times New Roman'
    return t


# ================= ENGLISH =================
doc = Document()
setup(doc)

# Table 1
add_title(doc, "Table 1.  Variation in anaesthesia practice across 335 secondary "
               "medical areas of Japan, fiscal year 2022")
make_table(doc,
           ["Code", "Procedure", "Areas reporting, n",
            "Mean SCR", "SD", "Median", "IQR", "Range",
            "Coefficient of variation, %"],
           [
            ["L008", "General anaesthesia (closed-circuit)", "335",
             "74.4", "39.3", "66.7", "44.1–97.1", "5.9–244.2", "52.9"],
            ["L004", "Spinal anaesthesia", "335",
             "61.8", "38.6", "54.5", "33.6–83.8", "0–240.5", "62.5"],
            ["L002", "Epidural anaesthesia", "331",
             "34.8", "30.3", "27.4", "14.4–46.1", "0–266.2", "86.9"],
            ["L009", "Anaesthesia management fee I", "335",
             "116.9", "47.5", "111.0", "82.4–146.2", "11.8–287.9", "40.6"],
            ["L100", "Nerve block, inpatient", "334",
             "81.3", "55.6", "70.5", "37.0–114.5", "0–296.3", "68.4"],
           ])
add_footnote(doc,
             "SCR, standardised claim ratio (national average = 100). "
             "SD, standard deviation; IQR, interquartile range. The coefficient "
             "of variation is the SD divided by the mean, expressed as a "
             "percentage.")

# Table 2
add_title(doc, "Table 2.  Between-prefecture vs within-prefecture variance "
               "components, multilevel null model")
make_table(doc,
           ["Procedure", "Prefecture variance (level 2)",
            "Residual variance (level 1)",
            "Intraclass correlation coefficient",
            "Within-prefecture variance, %"],
           [
            ["General anaesthesia (L008)", "89.8", "1455.8", "0.058", "94.2"],
            ["Spinal anaesthesia (L004)", "116.5", "1366.8", "0.079", "92.1"],
            ["Epidural anaesthesia (L002)", "73.2", "845.3", "0.080", "92.0"],
           ])
add_footnote(doc,
             "Between-prefecture variance was estimated from a mixed-effects "
             "model with prefecture as a random intercept and no fixed "
             "effects (null model) fitted by restricted maximum likelihood.")

# Table 3
add_title(doc, "Table 3.  University hospital presence and the standardised "
               "claim ratio")
make_table(doc,
           ["Procedure",
            "Mean SCR, university areas (n=64)",
            "Mean SCR, non-university areas (n=271)",
            "Difference (95% CI)", "Cohen's d",
            "Mixed-effects coefficient (95% CI)", "Marginal R²"],
           [
            ["General anaesthesia (L008)", "132.1", "61.6",
             "+70.5 (63.0 to 78.0)", "1.88",
             "+64.1 (54.8 to 73.4)", "0.358"],
            ["Spinal anaesthesia (L004)", "72.2", "59.4",
             "+12.8 (4.5 to 21.1)", "0.33",
             "+11.7 (2.6 to 20.7)", "0.018"],
            ["Epidural anaesthesia (L002)", "65.5", "27.4",
             "+38.2 (28.2 to 48.1)", "1.22",
             "+34.5 (25.5 to 43.4)", "0.148"],
           ])
add_footnote(doc,
             "Mixed-effects coefficient is the fixed effect of university "
             "hospital presence in a multilevel model with prefecture as a "
             "random intercept. Marginal R² is the proportion of total "
             "variance explained by the fixed effect relative to the null "
             "model. CI, confidence interval.")

out_en = os.path.join(OUTPUT_DIR, 'regional_anaesthesia_tables_IJQHC_EN.docx')
doc.save(out_en)
print(f"Saved: {out_en}")

# ================= JAPANESE =================
doc = Document()
setup(doc, jp=True)

add_title(doc,
          "表1．日本の335二次医療圏における麻酔診療のばらつき（2022年度）",
          jp=True)
make_table(doc,
           ["コード", "手技", "報告医療圏数",
            "平均SCR", "標準偏差", "中央値",
            "四分位範囲", "範囲", "変動係数 (%)"],
           [
            ["L008", "閉鎖循環式全身麻酔", "335",
             "74.4", "39.3", "66.7", "44.1–97.1", "5.9–244.2", "52.9"],
            ["L004", "脊椎麻酔", "335",
             "61.8", "38.6", "54.5", "33.6–83.8", "0–240.5", "62.5"],
            ["L002", "硬膜外麻酔", "331",
             "34.8", "30.3", "27.4", "14.4–46.1", "0–266.2", "86.9"],
            ["L009", "麻酔管理料I", "335",
             "116.9", "47.5", "111.0", "82.4–146.2", "11.8–287.9", "40.6"],
            ["L100", "神経ブロック（入院）", "334",
             "81.3", "55.6", "70.5", "37.0–114.5", "0–296.3", "68.4"],
           ],
           jp=True)
add_footnote(doc,
             "SCR：標準化レセプト出現比（全国平均=100）。変動係数は標準偏差"
             "を平均で除した百分率。", jp=True)

add_title(doc,
          "表2．マルチレベルnullモデルによる都道府県間・都道府県内分散成分",
          jp=True)
make_table(doc,
           ["手技", "都道府県分散（レベル2）",
            "残差分散（レベル1）", "級内相関係数",
            "都道府県内分散の割合 (%)"],
           [
            ["閉鎖循環式全身麻酔（L008）", "89.8", "1455.8", "0.058", "94.2"],
            ["脊椎麻酔（L004）", "116.5", "1366.8", "0.079", "92.1"],
            ["硬膜外麻酔（L002）", "73.2", "845.3", "0.080", "92.0"],
           ],
           jp=True)
add_footnote(doc,
             "都道府県をランダム切片、固定効果なしのnullモデル（制限最尤法）"
             "で推定。", jp=True)

add_title(doc, "表3．大学病院所在と標準化レセプト出現比", jp=True)
make_table(doc,
           ["手技",
            "大学病院あり（n=64） 平均SCR",
            "大学病院なし（n=271） 平均SCR",
            "差（95% CI）", "Cohen's d",
            "混合効果係数（95% CI）", "限界R²"],
           [
            ["閉鎖循環式全身麻酔（L008）", "132.1", "61.6",
             "+70.5（63.0～78.0）", "1.88",
             "+64.1（54.8～73.4）", "0.358"],
            ["脊椎麻酔（L004）", "72.2", "59.4",
             "+12.8（4.5～21.1）", "0.33",
             "+11.7（2.6～20.7）", "0.018"],
            ["硬膜外麻酔（L002）", "65.5", "27.4",
             "+38.2（28.2～48.1）", "1.22",
             "+34.5（25.5～43.4）", "0.148"],
           ],
           jp=True)
add_footnote(doc,
             "混合効果係数：都道府県をランダム切片とするマルチレベルモデルに"
             "おける大学病院所在の固定効果。限界R²：当該固定効果により説明"
             "される全分散の割合。CI：信頼区間。", jp=True)

out_jp = os.path.join(OUTPUT_DIR, 'regional_anaesthesia_tables_IJQHC_JP.docx')
doc.save(out_jp)
print(f"Saved: {out_jp}")
