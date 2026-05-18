#!/usr/bin/env python3
"""
Create BMJ-format docx for regional variation in anaesthesia practice in Japan.
- STROBE compliant
- British English
- Structured abstract
- Colour figures embedded
- BMJ formatting (Vancouver references, structured headings)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Style setup ──
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# Helper functions
def add_heading_bmj(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    return p

def add_table_row(table, cells, bold=False):
    row = table.add_row()
    for i, cell_text in enumerate(cells):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(cell_text))
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.bold = bold
    return row

def add_figure(path, caption, width=Inches(6)):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.italic = True
    else:
        add_para(f'[Figure: {path} not found]', italic=True)


# ═══════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════

title_text = (
    "Regional variation in anaesthesia practice across 335 secondary medical areas in Japan: "
    "a cross-sectional analysis of national claims data"
)
p = doc.add_paragraph()
run = p.add_run(title_text)
run.font.name = 'Arial'
run.font.size = Pt(16)
run.bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # blank line

add_para('[Author names]', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('[Affiliations]', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('[Corresponding author: name, address, email]', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
add_para('Word count: approximately 4,000', italic=True)
add_para('Tables: 2', italic=True)
add_para('Figures: 9 (6 two-dimensional choropleth maps + 3 three-dimensional extruded maps)', italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# STRUCTURED ABSTRACT (BMJ format: ≤250 words)
# ═══════════════════════════════════════════════════════════

add_heading_bmj('ABSTRACT', level=1)

p = doc.add_paragraph()
run = p.add_run('Objective ')
run.bold = True
run.font.name = 'Arial'
p.add_run(
    'To quantify regional variation in anaesthesia practice across Japanese secondary medical areas '
    'and to determine whether observed differences reflect insurance audit artefacts or genuine '
    'variation in clinical practice.'
).font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run('Design ')
run.bold = True
run.font.name = 'Arial'
p.add_run(
    'Cross-sectional ecological study using publicly available national claims data.'
).font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run('Setting ')
run.bold = True
run.font.name = 'Arial'
p.add_run(
    '335 secondary medical areas (SMAs) across all 47 prefectures of Japan.'
).font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run('Main outcome measures ')
run.bold = True
run.font.name = 'Arial'
p.add_run(
    'Standardised claim ratios (SCRs), age- and sex-adjusted, for general anaesthesia (L008), '
    'epidural anaesthesia (L002), continuous epidural infusion (L003), and spinal anaesthesia (L004). '
    'Three sensitivity analyses tested whether insurance audit variation explained observed differences.'
).font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run('Results ')
run.bold = True
run.font.name = 'Arial'
p.add_run(
    'Substantial regional variation was observed: coefficients of variation were 54.6% for general '
    'anaesthesia (L008), 83.2% for epidural anaesthesia (L002), and 64.9% for continuous epidural '
    'infusion (L003). Three sensitivity analyses indicated that audit variation is insufficient to '
    'explain observed differences: (1) 85.8% of L008 variance occurred within prefectures, where '
    'audit policy is uniform; (2) general and spinal anaesthesia showed inverse correlation '
    '(r=\u22120.506), consistent with clinical substitution (although also compatible with audit-driven '
    'reclassification, the magnitude of audit rate differences makes this unlikely); '
    '(3) audit rate differences (0.2 percentage points) could explain less than 1% of observed SCR '
    'variation. University hospital presence was the strongest single predictor, explaining 38.5% of '
    'total L008 variance (Cohen\'s d=1.78). The combined GA\u2013epidural technique (L003) was 1.73 times '
    'higher in university hospital areas.'
).font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run('Conclusions ')
run.bold = True
run.font.name = 'Arial'
p.add_run(
    'Regional variation in anaesthesia practice in Japan is structural, not an artefact of insurance '
    'auditing. University hospital proximity is the dominant determinant. Given evidence linking '
    'regional anaesthesia techniques to improved oncological outcomes, this variation may represent '
    'a modifiable source of health inequality.'
).font.name = 'Arial'

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# WHAT IS ALREADY KNOWN / WHAT THIS STUDY ADDS (BMJ box)
# ═══════════════════════════════════════════════════════════

add_heading_bmj('What is already known on this topic', level=2)
doc.add_paragraph(
    'Regional variation in surgical and anaesthesia practice has been documented in many countries, '
    'but the relative contributions of clinical preference, institutional factors, and administrative '
    'mechanisms (such as insurance auditing) remain unclear.',
    style='List Bullet'
)
doc.add_paragraph(
    'Japan\'s universal insurance system processes all claims through standardised audit, creating '
    'a unique natural experiment in which administrative variation can be distinguished from clinical variation.',
    style='List Bullet'
)

add_heading_bmj('What this study adds', level=2)
doc.add_paragraph(
    'Three independent sensitivity analyses demonstrate that regional variation in anaesthesia practice '
    'in Japan is structural rather than an artefact of insurance auditing.',
    style='List Bullet'
)
doc.add_paragraph(
    'University hospital presence explains 38.5% of variance in general anaesthesia volume\u2014a '
    'remarkably large effect for a single binary predictor\u2014and the effect is consistent across '
    'all 47 prefectures.',
    style='List Bullet'
)
doc.add_paragraph(
    'Combined general\u2013epidural anaesthesia, which may confer oncological benefit, shows the greatest '
    'regional inequality (CV=64.9%), suggesting a modifiable target for reducing health disparities.',
    style='List Bullet'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# INTRODUCTION
# ═══════════════════════════════════════════════════════════

add_heading_bmj('INTRODUCTION', level=1)

doc.add_paragraph(
    'Japan\'s universal health insurance system, established in 1961, covers the entire population '
    'under a standardised fee schedule (shinryo houshu tensuhyo) that specifies reimbursement for '
    'each medical procedure.\u00b9\u00b7\u00b2 This system was designed to ensure equitable access to healthcare '
    'regardless of geography. Yet studies using the National Database of Health Insurance Claims '
    '(NDB) have revealed substantial regional variation in medical practice, including surgical '
    'procedures, prescription patterns, and diagnostic testing.\u00b3\u2013\u2075'
)

doc.add_paragraph(
    'A distinctive feature of the Japanese system is mandatory insurance auditing (shinsa). All '
    'claims are reviewed by prefectural audit committees before reimbursement, and claims deemed '
    'inappropriate are reduced or rejected (satei). Audit rates vary across prefectures, ranging '
    'from 0.07% to 0.28%.\u2076 This raises a fundamental question: does observed regional variation '
    'in claims data reflect genuine differences in clinical practice, or is it an artefact of '
    'differential auditing?'
)

doc.add_paragraph(
    'This question has particular relevance for anaesthesia practice. The choice of anaesthesia '
    'technique\u2014general anaesthesia alone versus combined general\u2013regional techniques\u2014is '
    'not merely a matter of preference. A growing body of evidence suggests that regional anaesthesia '
    'techniques, particularly epidural analgesia combined with general anaesthesia, may improve '
    'recurrence-free survival and overall survival in cancer surgery.\u2077\u2013\u00b9\u00b2 '
    'A Cochrane systematic review found limited but suggestive evidence favouring regional techniques '
    'for cancer recurrence,\u2077 and the landmark randomised controlled trial by Sessler et al. '
    'reported no significant difference in breast cancer recurrence with regional anaesthesia, '
    'although secondary analyses and observational meta-analyses continue to support a potential '
    'benefit for certain cancer types.\u2078\u2013\u00b9\u00b2 '
    'If regional variation in anaesthesia technique is genuine rather than an audit artefact, it may '
    'represent a modifiable source of health inequality with implications for oncological outcomes.'
)

doc.add_paragraph(
    'The 2026 revision of the Japanese fee schedule renamed the general anaesthesia code from '
    '"mask or endotracheal intubation" to "supraglottic airway device or endotracheal intubation," '
    'reflecting longstanding ambiguity in coding that may have contributed to differential auditing '
    'of general anaesthesia claims.\u00b9\u00b3 Understanding the true nature of regional variation is therefore '
    'timely.'
)

doc.add_paragraph(
    'We used publicly available standardised claim ratio (SCR) data at the secondary medical area '
    'level to: (1) quantify regional variation in anaesthesia practice across 335 areas; '
    '(2) test whether insurance audit variation explains observed differences through three '
    'independent sensitivity analyses; and (3) identify structural determinants of variation, '
    'including university hospital proximity and pain clinic activity.'
)

# ═══════════════════════════════════════════════════════════
# METHODS
# ═══════════════════════════════════════════════════════════

add_heading_bmj('METHODS', level=1)

add_heading_bmj('Study design and data sources', level=2)

doc.add_paragraph(
    'This was a cross-sectional ecological study using publicly available data. We used three data sources:'
)

doc.add_paragraph(
    'Standardised claim ratios (SCRs) for fiscal year 2022 (Reiwa 4), published by the Cabinet Office '
    'as part of the "Regional Variation Visualisation" initiative.\u00b9\u2074 SCRs are age- and sex-adjusted '
    'ratios of observed to expected claim frequencies, where the national average equals 100. '
    'Data were available at both the prefectural level (47 prefectures) and the secondary medical area '
    '(SMA) level (335 areas).',
    style='List Bullet'
)

doc.add_paragraph(
    'Physician statistics from the 2022 Survey of Physicians, Dentists, and Pharmacists (e-Stat), '
    'providing the number of anaesthesiologists by secondary medical area.\u00b9\u2075',
    style='List Bullet'
)

doc.add_paragraph(
    'Geographic data from the National Land Numerical Information dataset (A38-20) for SMA boundaries, '
    'and administrative boundary data (N03) for prefectural borders and the Northern Territories.\u00b9\u2076',
    style='List Bullet'
)

add_heading_bmj('Anaesthesia codes analysed', level=2)

doc.add_paragraph(
    'We analysed the following procedure codes from the Japanese fee schedule:'
)

# Table of codes
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
for i, text in enumerate(['Code', 'Procedure', 'Clinical significance', 'SMAs with data']):
    hdr[i].paragraphs[0].add_run(text).bold = True
    hdr[i].paragraphs[0].runs[0].font.size = Pt(9)

codes_data = [
    ['L008', 'Closed-circuit general anaesthesia', 'Primary indicator of GA volume', '334'],
    ['L002', 'Epidural anaesthesia', 'Regional technique (alone or combined)', '307'],
    ['L003', 'Continuous epidural infusion', 'Direct indicator of GA+epidural combination', '331'],
    ['L004', 'Spinal anaesthesia', 'Alternative to GA for suitable procedures', '334'],
    ['L009', 'Anaesthesia management fee I', 'Proxy for specialist anaesthesiologist staffing', '314'],
    ['L100', 'Nerve block (inpatient)', 'Indicator of pain clinic activity', '335'],
]
for row_data in codes_data:
    add_table_row(table, row_data)

doc.add_paragraph()
add_para('Table 1. Anaesthesia procedure codes analysed.', italic=True)

add_heading_bmj('University hospital mapping', level=2)

doc.add_paragraph(
    'We mapped 81 university hospitals (44 national, 8 public, 29 private) to 64 SMAs based on '
    'municipal address. Under Japan\'s "one medical school per prefecture" policy, all 47 prefectures '
    'have at least one university hospital, enabling within-prefecture comparisons that control for '
    'prefectural audit policy.'
)

add_heading_bmj('Sensitivity analyses for the audit hypothesis', level=2)

doc.add_paragraph(
    'We tested the null hypothesis that "all regional variation in anaesthesia SCRs is explained by '
    'differential insurance auditing" using three independent approaches:'
)

doc.add_paragraph(
    'Test 1 (within-prefecture variance): If audit policy (applied uniformly within each prefecture) '
    'explains all variation, then within-prefecture variance should be negligible. We decomposed total '
    'SCR variance into between-prefecture and within-prefecture components using one-way ANOVA.',
    style='List Bullet'
)

doc.add_paragraph(
    'Test 2 (cross-code correlation): If auditing reclassifies general anaesthesia claims as spinal '
    'anaesthesia (or vice versa), these codes may show negative correlation at the ecological level '
    '(though clinical substitution would also produce this pattern). We calculated Pearson correlations '
    'between anaesthesia code pairs and interpreted them in conjunction with the quantitative audit '
    'impact estimate (Test 3).',
    style='List Bullet'
)

doc.add_paragraph(
    'Test 3 (quantitative impact): We estimated the maximum SCR impact of audit rate differences '
    'using published aggregate audit statistics (range 0.07\u20130.28% across prefectures).\u2076',
    style='List Bullet'
)

add_heading_bmj('Combined SCR analysis', level=2)

doc.add_paragraph(
    'To further test the audit reclassification hypothesis, we computed combined SCRs (L008+L004 for '
    'GA\u2013spinal substitution; L008+L002 for GA\u2013epidural substitution). If audit-driven '
    'reclassification is the primary source of variation, combining the reclassified codes should '
    'substantially reduce the coefficient of variation.'
)

add_heading_bmj('Statistical analysis', level=2)

doc.add_paragraph(
    'We report descriptive statistics (mean, standard deviation, coefficient of variation, percentiles) '
    'for SCRs. Between-group comparisons used Welch\'s t-test and Cohen\'s d for effect size.\u00b9\u2077 Variance '
    'decomposition used one-way ANOVA (between-prefecture) and hierarchical decomposition '
    '(between-prefecture, university hospital effect, residual). Correlations are Pearson r. '
    'Standardised claim ratios were computed using indirect standardisation, a standard method in '
    'small-area variation research.\u00b9\u2078 '
    'All analyses were conducted in Python 3.12 using pandas, scipy, and geopandas. Three-dimensional '
    'extruded choropleth maps were created using Plotly Mesh3d with Delaunay triangulation of polygon '
    'boundaries.\u00b9\u2079 Two-dimensional maps were created using matplotlib and geopandas.'
)

add_heading_bmj('Patient and public involvement', level=2)

doc.add_paragraph(
    'This study used publicly available aggregate data with no individual patient information. '
    'Patients were not involved in the design or conduct of this research.'
)

# ═══════════════════════════════════════════════════════════
# RESULTS
# ═══════════════════════════════════════════════════════════

add_heading_bmj('RESULTS', level=1)

add_heading_bmj('Regional variation in anaesthesia practice', level=2)

doc.add_paragraph(
    'Substantial variation was observed across all anaesthesia codes (Table 2). General anaesthesia '
    '(L008) SCRs ranged from 2.3 to 458.9 across 334 SMAs (coefficient of variation [CV] 54.6%). '
    'Epidural anaesthesia (L002) showed even greater variation (CV 83.2%), as did continuous epidural '
    'infusion (L003, CV 64.9%), the direct indicator of combined GA\u2013epidural technique. '
    'Spinal anaesthesia (L004) had a CV of 56.3%. Figure 1 shows the geographic distribution of '
    'L008 SCRs, with higher values concentrated in western Japan and metropolitan areas.'
)

# Table 2
table2 = doc.add_table(rows=1, cols=7)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0].cells
for i, text in enumerate(['Code', 'Procedure', 'P10', 'P50', 'P90', 'CV (%)', 'n']):
    hdr2[i].paragraphs[0].add_run(text).bold = True
    hdr2[i].paragraphs[0].runs[0].font.size = Pt(9)

table2_data = [
    ['L008', 'General anaesthesia', '32.3', '73.2', '131.5', '54.6', '334'],
    ['L002', 'Epidural anaesthesia', '7.6', '73.9', '184.0', '83.2', '307'],
    ['L003', 'Continuous epidural infusion', '19.0', '73.1', '150.0', '64.9', '331'],
    ['L004', 'Spinal anaesthesia', '31.1', '84.3', '169.0', '56.3', '334'],
    ['L009', 'Anaesthesia management fee I', '32.5', '79.1', '149.1', '57.2', '314'],
    ['L100', 'Nerve block (inpatient)', '23.4', '72.2', '174.7', '72.1', '335'],
]
for row_data in table2_data:
    add_table_row(table2, row_data)

doc.add_paragraph()
add_para(
    'Table 2. Distribution of standardised claim ratios for anaesthesia procedures across '
    'secondary medical areas (national average = 100).',
    italic=True
)

add_heading_bmj('Sensitivity analysis 1: within-prefecture variance', level=2)

doc.add_paragraph(
    'For general anaesthesia (L008), 85.8% of total variance occurred within prefectures, where '
    'audit policy is uniform (Table 3). This finding was consistent across codes: within-prefecture '
    'variance accounted for 72.7% (epidural), 58.0% (spinal), and 65.1% (nerve block) of total '
    'variation. Within individual prefectures, L008 SCRs varied enormously: from 61.4 to 458.9 '
    'within Tokyo (12 SMAs), and from 2.3 to 170.9 within Kumamoto (10 SMAs). Audit policy, '
    'which is applied uniformly within each prefecture, can explain at most 14.2% of L008 variation.'
)

add_heading_bmj('Sensitivity analysis 2: cross-code correlation', level=2)

doc.add_paragraph(
    'General anaesthesia (L008) and spinal anaesthesia (L004) showed a strong inverse correlation '
    '(r=\u22120.506, P<0.001), indicating a substitution pattern: areas with high GA volumes tend to '
    'have low spinal volumes and vice versa. This inverse correlation is consistent with genuine '
    'differences in clinical technique selection. It is also compatible in principle with audit-driven '
    'reclassification (where cases are shifted between the two codes) or defensive undercoding '
    '(where clinicians anticipate audit outcomes). However, given that the maximum audit rate '
    'difference across prefectures is only 0.21 percentage points (Test 3), reclassification of this '
    'magnitude is insufficient to produce the observed correlation. The most parsimonious '
    'interpretation is that the substitution pattern reflects clinical preference and regional '
    'practice style.'
)

doc.add_paragraph(
    'General anaesthesia (L008) and epidural anaesthesia (L002) showed a positive correlation '
    '(r=+0.319, P<0.001), indicating that areas with more GA also perform more epidural anaesthesia. '
    'This is consistent with a common supply factor (anaesthesia department capacity) rather than '
    'audit-driven substitution.'
)

add_heading_bmj('Sensitivity analysis 3: quantitative audit impact', level=2)

doc.add_paragraph(
    'The maximum difference in prefectural audit rates was 0.21 percentage points (0.07% to 0.28%). '
    'Even if this entire difference were concentrated on a single anaesthesia code, it would produce '
    'an SCR shift of approximately 0.2 points\u2014less than 0.3% of the observed 62-point interquartile '
    'range for L008 and less than 0.1% of the 198-point range for L002. Audit rate variation is '
    'quantitatively incapable of explaining the observed regional differences.'
)

add_heading_bmj('Combined SCR analysis', level=2)

doc.add_paragraph(
    'Combining L008 and L004 SCRs reduced the CV from 54.6% to 43.4%\u2014a modest 19% reduction, '
    'indicating that spinal\u2013GA substitution accounts for a minority of GA variation. Combining '
    'L008 and L002 actually increased the CV from 50.2% to 58.3%, as these codes are positively '
    'correlated. In both cases, the university hospital effect remained strong: combined L008+L004 '
    'showed d=1.16 (P<0.001) and combined L008+L002 showed d=1.00 (P<0.001) for the university '
    'hospital comparison.'
)

add_heading_bmj('University hospital effect', level=2)

doc.add_paragraph(
    'University hospital presence was the strongest single predictor of anaesthesia practice patterns. '
    'For general anaesthesia (L008), the 64 university hospital SMAs had a mean SCR of 130.2 compared '
    'with 67.7 for the 270 non-university SMAs (Cohen\'s d=1.78, t=12.49). Remarkably, in within-'
    'prefecture comparisons (which control for audit policy), university hospital SMAs had higher '
    'L008 SCRs in all 47 of 47 prefectures (100%), with a mean within-prefecture difference of '
    '+61.4 points (SD 33.7, t=12.49).'
)

doc.add_paragraph(
    'Hierarchical variance decomposition attributed 14.5% of L008 variance to between-prefecture '
    'differences, 38.5% to the university hospital effect (within prefecture), and 47.0% to residual '
    'within-prefecture variation. The university hospital effect alone\u2014a single binary variable\u2014'
    'explained nearly 40% of all variation in general anaesthesia volume across Japan.'
)

doc.add_paragraph(
    'A dose\u2013response relationship was apparent in Tokyo, where the number of university hospitals '
    'per SMA ranged from 0 to 5: SCRs increased from 62\u2013110 (0 hospitals) to 83\u2013119 (2 hospitals) '
    'to 168.5 (3 hospitals) to 435.7 (5 hospitals).'
)

add_heading_bmj('Combined general\u2013epidural anaesthesia', level=2)

doc.add_paragraph(
    'Continuous epidural infusion (L003), the direct indicator of combined GA\u2013epidural technique, '
    'showed a CV of 64.9%\u2014the largest variation among major anaesthesia codes. L003 SCR was 1.73 '
    'times higher in university hospital SMAs (126.4 vs 73.2, d=0.96, P<0.001) and correlated '
    'strongly with L008 (r=0.753, P<0.001). Three audit-hypothesis tests applied to the epidural '
    'combination (L008+L002) all rejected the audit explanation: positive L008\u2013L002 correlation '
    '(r=+0.319), increased CV on combination (50.2%\u219258.3%), and persistent university hospital '
    'effect (d=1.00).'
)

add_heading_bmj('Pain clinic spillover', level=2)

doc.add_paragraph(
    'The hypothesis that pain clinic activity (proxied by nerve block SCR, L100) would predict '
    'regional anaesthesia use in surgery was only weakly supported. The correlation between nerve '
    'block SCR and a composite regional anaesthesia index was r=0.153 (P<0.01). However, nerve block '
    'activity correlated more strongly with GA volume (r=0.307, P<0.001), suggesting that pain clinic '
    'activity reflects overall anaesthesia department capacity rather than a specific spillover of '
    'regional anaesthesia skills.'
)

# ═══════════════════════════════════════════════════════════
# DISCUSSION
# ═══════════════════════════════════════════════════════════

add_heading_bmj('DISCUSSION', level=1)

add_heading_bmj('Principal findings', level=2)

doc.add_paragraph(
    'This study demonstrates that regional variation in anaesthesia practice across Japan\'s 335 '
    'secondary medical areas is substantial and structural. Three independent sensitivity analyses '
    'converge on the conclusion that insurance audit variation explains, at most, a trivial fraction '
    'of observed differences. The dominant determinant is university hospital proximity, which alone '
    'explains nearly 40% of all variation in general anaesthesia volume. The finding that this effect '
    'is present in all 47 prefectures\u2014despite differing audit policies\u2014provides compelling '
    'evidence for a supply-side, institutional mechanism.'
)

add_heading_bmj('Comparison with existing literature', level=2)

doc.add_paragraph(
    'Our findings are consistent with the broader literature on medical practice variation. '
    'The Dartmouth Atlas project in the United States has documented extensive regional variation in '
    'surgical rates that is driven primarily by physician supply and practice style rather than '
    'patient need.\u00b2\u2070\u00b7\u00b2\u00b9 Similar patterns have been reported in the United Kingdom,\u00b2\u00b2 '
    'Germany,\u00b2\u00b3 and Australia.\u00b2\u2074 Our study extends this literature in two ways. '
    'First, we exploit Japan\'s unique institutional structure\u2014uniform fee schedule, prefectural '
    'audit committees\u2014to distinguish administrative from clinical sources of variation. Second, '
    'we provide the first systematic analysis of anaesthesia technique variation at a fine geographic '
    'scale (335 areas) in an East Asian universal insurance system.'
)

doc.add_paragraph(
    'The university hospital effect we document (d=1.78 for general anaesthesia) is remarkably large '
    'compared with effect sizes typically reported in medical practice variation research. This likely '
    'reflects the Japanese medical education system, in which university departments (ikyoku) exert '
    'substantial influence over clinical practice in affiliated hospitals within their geographic '
    'sphere.\u00b2\u2075\u00b7\u00b2\u2076'
)

add_heading_bmj('Implications for oncological outcomes', level=2)

doc.add_paragraph(
    'The finding that combined GA\u2013epidural technique (L003) shows a CV of 64.9% is clinically '
    'important in light of evidence suggesting that regional anaesthesia combined with general '
    'anaesthesia may improve recurrence-free survival in cancer surgery.\u2077\u2013\u00b9\u00b2 The 1.73-fold '
    'difference in L003 SCR between university and non-university hospital areas suggests that '
    'patients\' access to this potentially beneficial technique depends substantially on where they '
    'live\u2014a finding that challenges the equity premise of Japan\'s universal insurance system.\u00b9\u00b7\u00b2'
)

add_heading_bmj('The 2026 fee schedule revision', level=2)

doc.add_paragraph(
    'The renaming of the general anaesthesia code in 2026 (from "mask or endotracheal intubation" '
    'to "supraglottic airway device or endotracheal intubation") reflects acknowledged ambiguity '
    'in the previous coding, which may have led to differential audit decisions regarding supraglottic '
    'airway use. Our finding that audit-driven variation is quantitatively small suggests that this '
    'revision, while appropriate for coding clarity, is unlikely to substantially alter the geographic '
    'patterns we have documented.'
)

add_heading_bmj('Strengths and limitations', level=2)

doc.add_paragraph(
    'Strengths include the use of age- and sex-standardised ratios that cover the entire national '
    'population, analysis at a fine geographic scale (335 areas), multiple independent sensitivity '
    'analyses, and within-prefecture comparisons that control for audit policy. The ecological design '
    'avoids individual consent requirements and allows comprehensive geographic coverage.'
)

doc.add_paragraph(
    'Limitations include the ecological fallacy inherent in area-level analysis;\u00b2\u2077 the inability to '
    'distinguish individual patient-level decisions from aggregate patterns; the post-audit nature '
    'of SCR data (which reflects reimbursed rather than intended practice, though our sensitivity '
    'analyses suggest this distinction is quantitatively minor); the cross-sectional design, which '
    'precludes causal inference; incomplete separation of the university hospital effect from urban '
    'concentration (university hospitals are typically in the largest city of each prefecture); and '
    'the absence of code-specific audit rate data. Additionally, "defensive undercoding"\u2014where '
    'clinicians pre-emptively avoid claiming procedures they expect to be audited\u2014cannot be '
    'quantified from claims data.\u00b2\u2078'
)

add_heading_bmj('Conclusions and policy implications', level=2)

doc.add_paragraph(
    'Regional variation in anaesthesia practice in Japan is genuine and large, driven primarily by '
    'institutional factors\u2014particularly university hospital proximity\u2014rather than by '
    'differential insurance auditing. The 1.73-fold variation in combined GA\u2013epidural anaesthesia '
    'between university and non-university hospital areas is of particular concern given emerging '
    'evidence of oncological benefit.\u2077\u2013\u00b9\u00b2 Policy responses could include targeted education programmes, '
    'specialist outreach from university centres, and monitoring of anaesthesia technique patterns '
    'as a quality indicator.\u00b2\u2079 The sensitivity analysis framework developed here\u2014using within-'
    'prefecture variance decomposition, cross-code correlation, and quantitative audit impact '
    'estimation\u2014is applicable to investigating practice variation in any procedure under Japan\'s '
    'universal insurance system.'
)

# ═══════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_bmj('REFERENCES', level=1)

refs = [
    # Introduction: Japan's universal health insurance
    '1. Ikegami N, Yoo B-K, Hashimoto H, et al. Japanese universal health coverage: evolution, '
    'achievements, and challenges. Lancet 2011;378:1106\u201315.',
    '2. Ikegami N. Universal Health Coverage for Inclusive and Sustainable Development: Lessons from '
    'Japan. Washington, DC: World Bank, 2014.',
    # Introduction: NDB and regional variation
    '3. Matsuda S, Fujimori K. The claim database in Japan. Asian Pac J Dis Manag 2012;6:55\u201359.',
    '4. Suto M, Iba A, Sugiyama T, et al. Literature review of studies using the National Database '
    'of Health Insurance Claims of Japan (NDB): limitations and strategies in using the NDB for '
    'research. JMA J 2023;7:10\u201320.',
    '5. Hamada H, Sekimoto M, Imanaka Y. Effects of the per diem prospective payment system with '
    'DRG-like grouping system (DPC/PDPS) on resource usage and healthcare quality in Japan. Health '
    'Policy 2012;107:194\u2013201.',
    # Introduction: audit rates
    '6. Social Insurance Medical Fee Payment Fund (Shiharai Kikin). Review Statistics, Fiscal Year '
    '2022. https://www.ssk.or.jp/tokeijoho/shinsatokei/index.html (accessed 15 Mar 2026).',
    # Introduction: cancer and regional anaesthesia
    '7. Pei L, Tan G, Wang L, et al. Anaesthetic techniques for risk of malignant tumour recurrence. '
    'Cochrane Database Syst Rev 2019;(11):CD008877.',
    '8. Sessler DI, Pei L, Huang Y, et al. Recurrence of breast cancer after regional or general '
    'anaesthesia: a randomised controlled trial. Lancet 2019;394:1807\u201315.',
    '9. Exadaktylos AK, Buggy DJ, Moriarty DC, et al. Can anesthetic technique for primary breast '
    'cancer surgery affect recurrence or metastasis? Anesthesiology 2006;105:660\u20134.',
    '10. Sun Y, Li T, Gan TJ. The effects of perioperative regional anesthesia and analgesia on cancer '
    'recurrence and survival after oncology surgery: a systematic review and meta-analysis. Reg '
    'Anesth Pain Med 2015;40:589\u201398.',
    '11. Xie S, Li L, Meng F, et al. Regional anesthesia might reduce recurrence and metastasis rates '
    'in adult patients with cancers after surgery: a meta-analysis. BMC Anesthesiol 2024;24:19.',
    '12. Zhang D, Jiang J, Liu J, et al. Effects of perioperative epidural analgesia on cancer '
    'recurrence and survival. Front Oncol 2021;11:798435.',
    # Introduction: 2026 fee schedule revision
    '13. Ministry of Health, Labour and Welfare. Revision of medical fee schedule, 2026 '
    '(Reiwa 8 nendo shinryo houshu kaitei). https://www.mhlw.go.jp (accessed 15 Mar 2026).',
    # Methods: data sources
    '14. Cabinet Office. Regional variation visualisation (chiikisa no mieruka). '
    'https://www5.cao.go.jp/keizai-shimon/kaigi/special/reform/mieruka/ (accessed 10 Mar 2026).',
    '15. Ministry of Health, Labour and Welfare. Survey of Physicians, Dentists and Pharmacists, '
    '2022. e-Stat. https://www.e-stat.go.jp (accessed 12 Mar 2026).',
    '16. Ministry of Land, Infrastructure, Transport and Tourism. National Land Numerical Information '
    'download service. https://nlftp.mlit.go.jp (accessed 10 Mar 2026).',
    # Methods: statistical methods
    '17. Cohen J. Statistical Power Analysis for the Behavioral Sciences. 2nd ed. '
    'Hillsdale, NJ: Lawrence Erlbaum, 1988.',
    '18. Julious SA, Nicholl J, George S. Why do we continue to use standardized mortality ratios for '
    'small area comparisons? J Public Health Med 2001;23:40\u20136.',
    '19. Plotly Technologies Inc. Plotly: collaborative data science. Montreal, QC, 2015. '
    'https://plotly.com (accessed 15 Mar 2026).',
    # Discussion: Dartmouth Atlas and international comparisons
    '20. Wennberg JE. Tracking Medicine: A Researcher\'s Quest to Understand Health Care. '
    'New York: Oxford University Press, 2010.',
    '21. Wennberg JE, Fisher ES, Skinner JS. Geography and the debate over Medicare reform. '
    'Health Aff (Millwood) 2002;Suppl Web Exclusives:W96\u2013114.',
    '22. NHS RightCare. The NHS Atlas of Variation in Healthcare, 2015. London: Public Health England.',
    '23. Wengler A, Nimptsch U, Mansky T. Hip and knee replacement surgery (arthroplasty) and '
    'geographic variations in Germany. Dtsch Arztebl Int 2014;111:407\u201316.',
    '24. Australian Commission on Safety and Quality in Health Care. Australian Atlas of Healthcare '
    'Variation. Sydney: ACSQHC, 2015.',
    # Discussion: ikyoku system
    '25. Otsuka T. The ikyoku system of university orthopedic surgery departments: an in-hospital '
    'organizational system unique to Japan. J Orthop Sci 2012;17:513\u201314.',
    '26. Onishi H. History of Japanese medical education. Korean J Med Educ 2018;30:283\u201394.',
    # Discussion: ecological fallacy
    '27. Greenland S. Ecologic versus individual-level sources of bias in ecologic estimates of '
    'contextual health effects. Int J Epidemiol 2001;30:1343\u201350.',
    # Discussion: defensive medicine / undercoding
    '28. Sekimoto M, Imanaka Y, Kobayashi H, et al. Impact of hospital characteristics on the cost '
    'and practice patterns of acute myocardial infarction patients in Japan. Health Policy '
    '2006;78:34\u201345.',
    # Conclusions: quality monitoring
    '29. Merry AF, Cooper JB, Soyannwo O, et al. International standards for a safe practice of '
    'anesthesia 2010. Can J Anaesth 2010;57:1027\u201334.',
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.name = 'Arial'

# ═══════════════════════════════════════════════════════════
# FIGURES
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_bmj('FIGURES', level=1)

# 2D Maps
figure_list = [
    ('/home/ubuntu/en_map_L008_scr.png',
     'Figure 1. General anaesthesia (L008) standardised claim ratio by secondary medical area. '
     'Solid lines indicate prefectural boundaries; dashed lines indicate SMA boundaries. '
     'Northern Territories shown in white (no SMA designated).'),
    ('/home/ubuntu/en_map_univ_presence.png',
     'Figure 2. University hospital presence by secondary medical area. Red circles indicate '
     'SMAs containing one or more university hospitals (64 of 335 SMAs).'),
    ('/home/ubuntu/en_map_L004_scr.png',
     'Figure 3. Spinal anaesthesia (L004) standardised claim ratio by secondary medical area.'),
    ('/home/ubuntu/en_map_L008_L004_combined.png',
     'Figure 4. Combined general + spinal anaesthesia (L008+L004) standardised claim ratio. '
     'Combining these codes neutralises audit-driven reclassification between GA and spinal techniques.'),
    ('/home/ubuntu/en_map_L003_scr.png',
     'Figure 5. Continuous epidural infusion (L003) standardised claim ratio\u2014direct indicator '
     'of combined GA\u2013epidural technique. CV=64.9%.'),
    ('/home/ubuntu/en_map_L003_L008_ratio_corrected.png',
     'Figure 6. L003/L008 ratio (combined epidural rate per GA case), adjusted for surgery volume. '
     'Higher values indicate greater use of combined GA\u2013epidural technique.'),
]

for path, caption in figure_list:
    add_figure(path, caption, width=Inches(5.5))
    doc.add_paragraph()

# 3D Maps
threed_figures = [
    ('/home/ubuntu/3d_extruded/3D_ratio_by_anes_v2.png',
     'Figure 7. Three-dimensional extruded map: colour represents L003/L008 ratio (combined epidural rate), '
     'height represents anaesthesiologist count. Areas with both high ratio and high specialist count '
     'appear as tall, green-coloured polygons.'),
    ('/home/ubuntu/3d_extruded/3D_L008_by_anes_v2.png',
     'Figure 8. Three-dimensional extruded map: colour represents L008 SCR (general anaesthesia volume), '
     'height represents anaesthesiologist count. Red, tall areas indicate high GA volume with many '
     'anaesthesiologists (typically metropolitan university hospital areas).'),
    ('/home/ubuntu/3d_extruded/3D_ratio_by_surgery_v2.png',
     'Figure 9. Three-dimensional extruded map: colour represents L003/L008 ratio, height represents '
     'GA per surgery ratio (L008/K-chapter SCR). Height variation is modest, indicating that GA per '
     'surgery is relatively uniform; colour variation shows the epidural combination rate differs substantially.'),
]

for path, caption in threed_figures:
    add_figure(path, caption, width=Inches(5.5))
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# STROBE CHECKLIST (appendix)
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_bmj('STROBE Checklist (Cross-sectional study)', level=1)

strobe_items = [
    ('1a', 'Title and abstract', 'Title page and structured abstract'),
    ('1b', 'Informative abstract', 'Structured abstract with all required elements'),
    ('2', 'Background/rationale', 'Introduction paragraphs 1\u20133'),
    ('3', 'Objectives', 'Introduction paragraph 5'),
    ('4', 'Study design', 'Methods: Study design and data sources'),
    ('5', 'Setting', 'Methods: 335 SMAs, FY2022 data'),
    ('6a', 'Participants/eligibility', 'N/A (ecological study using aggregate data)'),
    ('7', 'Variables', 'Methods: Anaesthesia codes analysed (Table 1)'),
    ('8', 'Data sources', 'Methods: Study design and data sources'),
    ('9', 'Bias', 'Methods: Sensitivity analyses; Discussion: Limitations'),
    ('10', 'Study size', 'Methods: 335 SMAs, 47 prefectures'),
    ('11', 'Quantitative variables', 'Methods: Statistical analysis'),
    ('12', 'Statistical methods', 'Methods: Statistical analysis'),
    ('13', 'Participants', 'Results: 334 SMAs with L008 data'),
    ('14', 'Descriptive data', 'Results: Table 2'),
    ('15', 'Outcome data', 'Results: Regional variation section'),
    ('16', 'Main results', 'Results: Sensitivity analyses 1\u20133, University hospital effect'),
    ('17', 'Other analyses', 'Results: Combined SCR, Pain clinic spillover'),
    ('18', 'Key results', 'Discussion: Principal findings'),
    ('19', 'Limitations', 'Discussion: Strengths and limitations'),
    ('20', 'Interpretation', 'Discussion: Comparison with existing literature'),
    ('21', 'Generalisability', 'Discussion: Conclusions and policy implications'),
    ('22', 'Funding', '[To be completed]'),
]

strobe_table = doc.add_table(rows=1, cols=3)
strobe_table.style = 'Table Grid'
strobe_hdr = strobe_table.rows[0].cells
for i, text in enumerate(['Item', 'STROBE recommendation', 'Location in manuscript']):
    strobe_hdr[i].paragraphs[0].add_run(text).bold = True
    strobe_hdr[i].paragraphs[0].runs[0].font.size = Pt(9)

for item in strobe_items:
    add_table_row(strobe_table, item)

# ═══════════════════════════════════════════════════════════
# DECLARATIONS
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_bmj('DECLARATIONS', level=1)

add_heading_bmj('Ethics approval', level=2)
doc.add_paragraph(
    'This study used publicly available aggregate data. No individual-level data were accessed. '
    'Ethics committee approval was not required under the Ethical Guidelines for Medical and '
    'Biological Research Involving Human Subjects (Japan, 2021 revision).'
)

add_heading_bmj('Data availability', level=2)
doc.add_paragraph(
    'All data used in this study are publicly available. SCR data: Cabinet Office regional variation '
    'visualisation (https://www5.cao.go.jp/keizai-shimon/kaigi/special/reform/mieruka/). '
    'Physician statistics: e-Stat (https://www.e-stat.go.jp). GIS boundaries: National Land '
    'Numerical Information (https://nlftp.mlit.go.jp). Analysis code is available from the '
    'corresponding author on request.'
)

add_heading_bmj('Competing interests', level=2)
doc.add_paragraph('[To be completed by authors.]')

add_heading_bmj('Funding', level=2)
doc.add_paragraph('[To be completed by authors.]')

add_heading_bmj('Contributors', level=2)
doc.add_paragraph('[To be completed by authors. BMJ requires detailed contribution statements.]')

add_heading_bmj('Transparency declaration', level=2)
doc.add_paragraph(
    'The lead author (the manuscript\'s guarantor) affirms that the manuscript is an honest, '
    'accurate, and transparent account of the study being reported; that no important aspects '
    'of the study have been omitted; and that any discrepancies from the study as planned have '
    'been explained.'
)

# ═══════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════

output_path = '/home/ubuntu/regional_anaesthesia_BMJ_EN.docx'
doc.save(output_path)
print(f"Saved: {output_path}")
