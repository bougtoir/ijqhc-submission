#!/usr/bin/env python3
"""Create IJQHC End-Matter file (uploaded as separate supplementary file).

Per IJQHC guidelines: End-Matter must be uploaded as a separate file
with file type 'Online-only Supplementary Data', file name 'End Matter',
containing: Contributorship, Ethics and permissions, Funding,
Conflict of interests, Acknowledgments, Data Availability Statement.
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'documents', 'IJQHC')
os.makedirs(OUTPUT_DIR, exist_ok=True)

doc = Document()
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0
style.paragraph_format.space_after = Pt(0)


def add_para(text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p


def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13) if level == 1 else Pt(12)
    run.bold = True
    return p


# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("End Matter")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

add_para("Manuscript: Regional variation in anaesthesia practice in Japan: "
         "structural determinant or claims-audit artefact?", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)

# Contributorship
add_heading("Contributorship")
add_para(
    "[To be completed by the authors.] The lead author conceived the study, "
    "secured the data sources, designed the analytical framework, performed "
    "all statistical analyses and drafted the manuscript. Co-authors "
    "contributed to study conception and design, interpretation of results, "
    "and critical revision of the manuscript for important intellectual "
    "content. All authors approved the final version submitted. The lead "
    "author is the guarantor for the study.")
add_para(
    "Use of artificial intelligence tools: Devin (Cognition AI) was used as a "
    "coding assistant to support data processing, statistical modelling, "
    "geographic visualisation and manuscript preparation. All scientific "
    "interpretation, methodological decisions and the wording of the "
    "scientific claims were made by the human authors, who take full "
    "responsibility for the integrity and accuracy of the work.")
add_para(
    "No professional statistician outside the author list was consulted; "
    "the lead author took responsibility for the statistical analysis.")

# Ethics and other permissions
add_heading("Ethics and other permissions")
add_para(
    "This study used publicly available aggregate data only, with no access "
    "to individual-level data. Ethics committee approval was not required "
    "under the Ethical Guidelines for Medical and Biological Research "
    "Involving Human Subjects (Japan, 2021 revision). No additional "
    "permissions were required for the use of the data sources cited.")

# Funding
add_heading("Funding")
add_para("[To be completed by the authors.] This work received no specific "
         "grant from any funding agency in the public, commercial or "
         "not-for-profit sectors.")

# Conflict of interests
add_heading("Conflict of interests")
add_para("[To be completed by the authors.] No known conflict of interests.")

# Acknowledgements
add_heading("Acknowledgments")
add_para(
    "The authors thank the Cabinet Office Regional Variation Visualisation "
    "initiative, the Ministry of Health, Labour and Welfare, and the "
    "Ministry of Land, Infrastructure, Transport and Tourism for maintaining "
    "the public datasets used in this study.")

# Data Availability Statement
add_heading("Data Availability Statement")
add_para(
    "All data used in this study are publicly available. Standardised claim "
    "ratios: Cabinet Office Regional Variation Visualisation "
    "(https://www5.cao.go.jp/keizai-shimon/kaigi/special/reform/mieruka/). "
    "Physician statistics: Survey of Physicians, Dentists and Pharmacists, "
    "e-Stat (https://www.e-stat.go.jp). Geographic boundary data: National "
    "Land Numerical Information "
    "(https://nlftp.mlit.go.jp). Derived datasets and analysis code "
    "supporting the findings are available from the corresponding author on "
    "reasonable request.")

# Transparency declaration
add_heading("Transparency declaration")
add_para(
    "The lead author (manuscript guarantor) affirms that the manuscript is "
    "an honest, accurate and transparent account of the study being "
    "reported; that no important aspects of the study have been omitted; and "
    "that any discrepancies from the study as originally planned have been "
    "explained.")

out = os.path.join(OUTPUT_DIR, 'end_matter_IJQHC_EN.docx')
doc.save(out)
print(f"Saved: {out}")
