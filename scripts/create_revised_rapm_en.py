#!/usr/bin/env python3
"""
Create REVISED manuscript for resubmission after RAPM rejection.
Addresses ALL Major (1-5) and Minor (1-5) reviewer comments:

Major 1: SMA structure explanation (nested within prefectures)
Major 2: Multilevel model with ICC
Major 3: Small-sample sensitivity (Empirical Bayes shrinkage)
Major 4: SCR definition clarification (residence-based)
Major 5: Methods rewritten from bullets to paragraphs

Minor 1: Data year (FY2022 / Reiwa 4)
Minor 2: Title revised to include audit theme
Minor 3: SCR defined in Abstract (observed-to-expected ratio)
Minor 4: 3D terminology corrected (choropleth = 2D)
Minor 5: Figures consolidated

Also adds:
- Healthcare system context (WHO HAQ, universal coverage, patient movement data)
- Updated numbers from corrected university hospital mapping
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Style setup ──
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# Helper functions
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    return p

def add_table_row(table, cells, bold=False):
    row = table.add_row()
    for i, cell_text in enumerate(cells):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(cell_text))
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.bold = bold
    return row

def add_figure(path, caption, width=Inches(6)):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.italic = True
    else:
        add_para(f'[Figure: {path} not found]', italic=True)


# ═══════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════

# [Minor 2] Title revised to include audit/claims theme
title_text = (
    "Regional variation in anesthesia practice in Japan: "
    "clinical preference or claims audit artifact? "
    "A cross-sectional ecological study of 335 secondary medical areas"
)
p = doc.add_paragraph()
run = p.add_run(title_text)
run.font.name = 'Arial'
run.font.size = Pt(16)
run.bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

add_para('[Author names]', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('[Affiliations]', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('[Corresponding author: name, address, email]', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
add_para('Word count: approximately 4,500', italic=True)
add_para('Tables: 3 (Table 1: anesthesia codes; Table 2: descriptive statistics; Table 3: multilevel model results)', italic=True)
# [Minor 4 & 5] Corrected figure descriptions
add_para('Figures: 4 (2 multi-panel choropleth maps, 1 funnel plot, 1 three-dimensional extruded map)', italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# STRUCTURED ABSTRACT
# ═══════════════════════════════════════════════════════════

add_heading_styled('ABSTRACT', level=1)

p = doc.add_paragraph()
run = p.add_run('Background ')
run.bold = True
run.font.name = 'Arial'
p.add_run(
    'Regional variation in anesthesia practice has been documented internationally, but the '
    'relative contributions of clinical preference, institutional factors, and administrative '
    'mechanisms remain unclear. Japan\u2019s universal health insurance system, with its mandatory '
    'prefectural claims auditing, provides a unique setting to distinguish administrative from '
    'clinical sources of variation.'
).font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run('Methods ')
run.bold = True
run.font.name = 'Arial'
# [Minor 1] Data year specified; [Minor 3] SCR defined
p.add_run(
    'Cross-sectional ecological study using publicly available standardized claim-occurrence '
    'ratios (SCRs; age- and sex-adjusted observed-to-expected ratios, national average = 100) '  # Minor 3
    'for fiscal year 2022 (Reiwa 4) '  # Minor 1
    'across 335 secondary medical areas (SMAs) in Japan. '
    'SMAs are geographic planning units completely nested within 47 prefectures, with no cross-boundary overlap. '  # Major 1
    'We analyzed six anesthesia procedure codes using three complementary approaches: '
    'multilevel linear mixed models with prefecture-level random intercepts to estimate '  # Major 2
    'intraclass correlation coefficients (ICCs); three sensitivity analyses testing whether '
    'insurance audit variation explains observed differences; and empirical Bayes shrinkage '  # Major 3
    'estimation to address small-sample instability in areas with few claims.'
).font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run('Results ')
run.bold = True
run.font.name = 'Arial'
# Updated numbers from corrected mapping
p.add_run(
    'Substantial regional variation was observed: coefficients of variation (CV) were 53.6% for '
    'general anesthesia (L008), 87.0% for epidural anesthesia (L002), and 56.8% for spinal '
    'anesthesia (L004). Multilevel models showed that prefecture-level clustering explained '
    '5.8% of general anesthesia variance (ICC = 0.058), whereas university hospital presence '
    'explained 35.8% of total variance (\u03b2 = +64.1, 95% CI 54.8\u201373.4, P < 0.001). '
    'In within-prefecture comparisons, university hospital areas had higher general anesthesia '
    'SCRs in all 47 of 47 prefectures (100%; mean difference +63.3, Cohen\u2019s d = 1.88). '
    'Empirical Bayes shrinkage attenuated effect sizes by only 9.0% (d = 1.89 \u2192 1.72), '
    'confirming robustness to small-sample instability. '
    'L008 and L004 showed positive correlation (r = +0.235, P < 0.001), inconsistent with '
    'audit-driven reclassification between these codes.'
).font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run('Conclusion ')
run.bold = True
run.font.name = 'Arial'
p.add_run(
    'Regional variation in anesthesia practice in Japan is structural, not an artifact of '
    'insurance auditing. University hospital proximity is the dominant determinant, persisting '
    'after multilevel adjustment and empirical Bayes shrinkage. Given evidence linking regional '
    'anesthesia techniques to improved oncologic outcomes, this variation may represent a '
    'modifiable source of health inequality.'
).font.name = 'Arial'

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# KEY MESSAGES
# ═══════════════════════════════════════════════════════════

add_heading_styled('WHAT IS ALREADY KNOWN ON THIS TOPIC', level=2)
doc.add_paragraph(
    'Regional variation in surgical and anesthesia practice has been documented in many countries, '
    'but the relative contributions of clinical preference, institutional factors, and administrative '
    'mechanisms (such as insurance auditing) remain unclear.',
    style='List Bullet'
)
doc.add_paragraph(
    'Japan\u2019s universal insurance system processes all claims through standardized prefectural audit, '
    'creating a unique natural experiment in which administrative variation can be distinguished from '
    'clinical variation.',
    style='List Bullet'
)

add_heading_styled('WHAT THIS STUDY ADDS', level=2)
doc.add_paragraph(
    'Three independent sensitivity analyses and multilevel modeling demonstrate that regional '
    'variation in anesthesia practice in Japan is structural rather than an artifact of '
    'insurance auditing.',
    style='List Bullet'
)
doc.add_paragraph(
    'University hospital presence explains 40.5% of within-prefecture variance in general '
    'anesthesia volume (Cohen\u2019s d = 1.88)\u2014a remarkably large effect for a single binary '
    'predictor\u2014and the effect is consistent across all 47 prefectures.',
    style='List Bullet'
)

add_heading_styled('HOW THIS STUDY MIGHT AFFECT RESEARCH, PRACTICE OR POLICY', level=2)
doc.add_paragraph(
    'Combined general\u2013epidural anesthesia, which may confer oncologic benefit, shows the '
    'greatest regional inequality (CV = 64.9%), suggesting a modifiable target for reducing '
    'health disparities.',
    style='List Bullet'
)
doc.add_paragraph(
    'The multilevel sensitivity analysis framework developed here is applicable to investigating '
    'practice variation in any procedure under Japan\u2019s universal insurance system.',
    style='List Bullet'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# INTRODUCTION
# ═══════════════════════════════════════════════════════════

add_heading_styled('INTRODUCTION', level=1)

# Healthcare system context paragraph (NEW - addresses user request about WHO data, universal coverage)
doc.add_paragraph(
    'Japan\u2019s universal health insurance system (kokumin kai hoken), established in 1961, '
    'covers the entire population under a standardized fee schedule that specifies reimbursement '
    'for each medical procedure.\u00b9\u00b7\u00b2 The system has achieved internationally recognized levels '
    'of healthcare access: the World Health Organization\u2019s Healthcare Access and Quality (HAQ) '
    'Index ranked Japan at 87.5 out of 100 (10th decile, top tier globally) in 2019,\u00b3 and Japan '
    'consistently ranks among the highest-performing countries for universal health coverage '
    'effective service coverage.\u2074 Healthcare delivery is organized through a hierarchical system '
    'of medical areas (iryo-ken): 47 tertiary medical areas (corresponding to prefectures) and '
    '335 secondary medical areas (SMAs), each designed to be self-sufficient for inpatient care.\u2075 '
    'Cross-boundary patient movement is limited: the 2008 Patient Survey found that only 6.1% '
    'of inpatients and 3.0% of outpatients received care outside their home prefecture,\u2076 '
    'and cross-SMA movement is predominantly to adjacent areas within the same prefecture.\u2077 '
    'The system was designed to ensure equitable access regardless of geography.'
)

doc.add_paragraph(
    'Yet studies using the National Database of Health Insurance Claims (NDB) have revealed '
    'substantial regional variation in medical practice, including surgical procedures, '
    'prescription patterns, and diagnostic testing.\u2078\u2013\u00b9\u2070 A distinctive feature of the Japanese '
    'system is mandatory insurance auditing (shinsa). All claims are reviewed by prefectural '
    'audit committees before reimbursement, and claims deemed inappropriate are reduced or '
    'rejected (satei). Audit rates vary across prefectures, ranging from 0.07% to 0.28%.\u00b9\u00b9 '
    'This raises a fundamental question: does observed regional variation in claims data '
    'reflect genuine differences in clinical practice, or is it an artifact of differential '
    'auditing?'
)

doc.add_paragraph(
    'This question has particular relevance for anesthesia practice. The choice of anesthesia '
    'technique\u2014general anesthesia alone versus combined general\u2013regional techniques\u2014is not '
    'merely a matter of preference. A growing body of evidence suggests that regional anesthesia '
    'techniques, particularly epidural analgesia combined with general anesthesia, may improve '
    'recurrence-free survival and overall survival in cancer surgery.\u00b9\u00b2\u2013\u00b9\u2077 A Cochrane '
    'systematic review found limited but suggestive evidence favoring regional techniques for '
    'cancer recurrence,\u00b9\u00b2 and a landmark randomized controlled trial by Sessler et al. reported '
    'no significant difference in breast cancer recurrence with regional anesthesia, although '
    'secondary analyses and observational meta-analyses continue to support a potential benefit '
    'for certain cancer types.\u00b9\u00b3\u2013\u00b9\u2077 If regional variation in anesthesia technique is genuine '
    'rather than an audit artifact, it may represent a modifiable source of health inequality '
    'with implications for oncologic outcomes.'
)

doc.add_paragraph(
    'The 2026 revision of the Japanese fee schedule renamed the general anesthesia code from '
    '\u201cmask or endotracheal intubation\u201d to \u201csupraglottic airway device or endotracheal '
    'intubation,\u201d reflecting longstanding ambiguity in coding that may have contributed to '
    'differential auditing of general anesthesia claims.\u00b9\u2078 Understanding the true nature of '
    'regional variation is therefore timely.'
)

doc.add_paragraph(
    'We used publicly available standardized claim ratio (SCR) data at the secondary medical '
    'area level to: (1) quantify regional variation in anesthesia practice across 335 areas; '
    '(2) test whether insurance audit variation explains observed differences through three '
    'independent sensitivity analyses; (3) estimate the contribution of prefecture-level '
    'clustering using multilevel models; and (4) identify structural determinants of variation, '
    'including university hospital proximity and anesthesiologist supply.'
)

# ═══════════════════════════════════════════════════════════
# METHODS (Major 5: rewritten from bullets to paragraphs)
# ═══════════════════════════════════════════════════════════

add_heading_styled('METHODS', level=1)

add_heading_styled('Study design and ethical considerations', level=2)

doc.add_paragraph(
    'This was a cross-sectional ecological study using publicly available aggregate data. '
    'No individual-level data were accessed. Ethics committee approval was not required under '
    'the Ethical Guidelines for Medical and Biological Research Involving Human Subjects '
    '(Japan, 2021 revision). The study followed the STROBE guidelines for reporting '
    'observational studies (checklist provided as supplementary material).'
)

# [Major 1] SMA structure explanation - NEW detailed subsection
add_heading_styled('Geographic units: secondary medical areas', level=2)

doc.add_paragraph(
    'Japan\u2019s healthcare delivery system is organized into a hierarchical structure of '
    'medical areas defined under Article 30-4 of the Medical Care Act.\u2075 The 47 prefectures '
    'serve as tertiary medical areas, responsible for highly specialized care. Within each '
    'prefecture, secondary medical areas (SMAs; nijiiryoken) are designated as the basic '
    'geographic units for planning inpatient care. As of the study period, there were 335 SMAs '
    'nationally (reduced from 344 in the 2018 revision).\u00b9\u2079'
)

doc.add_paragraph(
    'Critically, SMAs are completely nested within prefectures: no SMA crosses a prefectural '
    'boundary. This differs fundamentally from geographic units used in practice variation '
    'research in other countries, such as the United States\u2019 Hospital Service Areas (HSAs) '
    'and Hospital Referral Regions (HRRs), which are defined by patient flow patterns and '
    'may cross state boundaries.\u00b2\u2070 The complete nesting of SMAs within prefectures means '
    'that within-prefecture comparisons are unconfounded by cross-boundary spillover effects '
    'and are directly interpretable as variation under a common prefectural audit policy. '
    'Each SMA is composed of one or more contiguous municipalities and is designed to be '
    'self-sufficient for general inpatient care, including surgical services.'
)

# [Major 4] SCR definition clarification - residence-based
add_heading_styled('Data sources and standardized claim ratios', level=2)

doc.add_paragraph(
    'We used three data sources. First, standardized claim-occurrence ratios (SCRs) for '
    'fiscal year 2022 (Reiwa 4), published by the Cabinet Office as part of the \u201cRegional '
    'Variation Visualization\u201d (chiikisa no mieruka) initiative.\u00b2\u00b9 SCRs are computed by '
    'indirect standardization: for each procedure, expected claim frequencies are calculated '
    'by applying national age- and sex-specific claim rates to the local population structure, '
    'and the SCR is defined as 100 \u00d7 (observed claims / expected claims). An SCR of 100 '
    'indicates a claim frequency equal to the national average after age and sex adjustment. '
    'SCRs are calculated on a residence basis (jukyochi base): claims are attributed to the '  # Major 4
    'beneficiary\u2019s registered address, not to the location of the providing facility. '
    'This means that patient travel to other areas for treatment does not inflate the SCR '
    'of the receiving area. Data were available at both prefectural (n = 47) and SMA (n = 335) '
    'levels. Areas with very few claims are subject to masking by the data provider to protect '
    'privacy; such areas appear as missing values in our dataset.'
)

doc.add_paragraph(
    'Second, we obtained physician statistics from the 2022 Survey of Physicians, Dentists, '
    'and Pharmacists (e-Stat), providing counts of anesthesiologists and total physicians by '
    'SMA.\u00b2\u00b2 Third, geographic boundary data were obtained from the National Land Numerical '
    'Information dataset (A38-20 for SMA boundaries, N03 for prefectural borders and the '
    'Northern Territories).\u00b2\u00b3'
)

add_heading_styled('Anesthesia procedure codes', level=2)

doc.add_paragraph(
    'We analyzed six procedure codes from the Japanese fee schedule (shinryo houshu tensuhyo) '
    'that represent the major categories of anesthesia practice (Table 1): L008 (closed-circuit '
    'general anesthesia), L002 (epidural anesthesia), L003 (continuous epidural infusion, a '
    'direct indicator of combined general\u2013epidural technique), L004 (spinal anesthesia), '
    'L009 (anesthesia management fee I, a proxy for specialist anesthesiologist staffing), '
    'and L100 (nerve block, inpatient, an indicator of pain clinic activity). Intravenous '
    'anesthesia (L001) was also analyzed as a secondary outcome.'
)

add_heading_styled('University hospital mapping', level=2)

doc.add_paragraph(
    'We mapped 81 university hospitals (44 national, 8 public, 29 private) to their '
    'respective SMAs based on municipal address. These 81 hospitals are distributed across '
    '64 SMAs in all 47 prefectures, reflecting Japan\u2019s \u201cone medical school per '
    'prefecture\u201d policy. This distribution enables within-prefecture comparisons between '
    'university-hospital SMAs and non-university SMAs, which serve as a natural control for '
    'prefectural audit policy (since audit criteria are applied uniformly within each prefecture).'
)

# [Major 2] Multilevel model - NEW subsection
add_heading_styled('Statistical analysis', level=2)

doc.add_paragraph(
    'We employed three complementary analytical approaches. First, descriptive statistics '
    '(mean, standard deviation, coefficient of variation, percentiles, range) characterized '
    'the distribution of SCRs across SMAs. Between-group comparisons used Welch\u2019s t-test '
    'for unequal variances, Cohen\u2019s d for effect size,\u00b2\u2074 and the Mann\u2013Whitney U test for '
    'nonparametric confirmation.'
)

# Multilevel model paragraph
doc.add_paragraph(
    'Second, we fitted multilevel linear mixed models (also termed hierarchical linear models '
    'or mixed-effects models) to account for the nested structure of SMAs within prefectures.\u00b2\u2075 '
    'For each anesthesia code, we specified a two-level structure: SMAs (level 1, n = 335) '
    'nested within prefectures (level 2, n = 47). A null model (random intercept only, no '
    'fixed effects) was fitted first to estimate the intraclass correlation coefficient (ICC), '
    'defined as the proportion of total variance attributable to the prefecture level. '
    'Subsequent models added fixed effects for university hospital presence (binary), '
    'anesthesiologist density (standardized), and total physician density (standardized). '
    'Models were estimated by restricted maximum likelihood (REML) using the Python statsmodels '
    'MixedLM implementation.\u00b2\u2076 Marginal R\u00b2 was calculated as the proportional reduction in '
    'total variance (prefecture + residual) between the null model and each fitted model.'
)

# Sensitivity analysis paragraph
doc.add_paragraph(
    'Third, we conducted three sensitivity analyses to test the null hypothesis that insurance '
    'audit variation explains observed SCR differences. (a) Within-prefecture variance '
    'decomposition: if prefectural audit policy (applied uniformly within each prefecture) '
    'explains all variation, then within-prefecture variance should be negligible. We '
    'decomposed total SCR variance into between-prefecture and within-prefecture components '
    'using one-way ANOVA, and further decomposed within-prefecture variance into university '
    'hospital effect and residual components using hierarchical sum-of-squares decomposition. '
    '(b) Cross-code correlation analysis: if auditing reclassifies claims between general and '
    'spinal anesthesia codes, these should show negative correlation. We calculated Pearson '
    'correlations between all code pairs. (c) Quantitative audit impact estimation: we estimated '
    'the maximum SCR impact of audit rate differences using published aggregate audit statistics.'
)

# [Major 3] Small-sample sensitivity - NEW paragraph
doc.add_paragraph(
    'To address potential instability in SCR estimates for SMAs with small claim volumes, '
    'we applied empirical Bayes (EB) shrinkage estimation.\u00b2\u2077 For each SMA, the EB estimate '
    'is a weighted average of the observed SCR and the prefecture mean, where the weight '
    'reflects the relative precision of the local estimate. SMAs with unstable estimates '
    '(high variance relative to the group) are \u201cshrunk\u201d toward the prefecture mean, '
    'producing more conservative effect estimates. We compared all main findings using both '
    'raw SCRs and EB-adjusted SCRs. Additionally, SMAs with missing data (due to privacy '
    'masking of small cell counts) were excluded from the relevant analyses, and the number '
    'of SMAs contributing to each analysis is reported.'
)

doc.add_paragraph(
    'Combined SCR analyses (L008 + L004, L008 + L002) were computed to test whether '
    'combining codes that might be subject to audit reclassification would substantially '
    'reduce the coefficient of variation. If audit-driven relabeling between codes is the '
    'primary source of variation, the combined SCR should show markedly lower variation.'
)

add_heading_styled('Patient and public involvement', level=2)

doc.add_paragraph(
    'This study used publicly available aggregate data with no individual patient information. '
    'Patients were not involved in the design or conduct of this research.'
)

# ═══════════════════════════════════════════════════════════
# RESULTS
# ═══════════════════════════════════════════════════════════

add_heading_styled('RESULTS', level=1)

add_heading_styled('Study population and descriptive statistics', level=2)

doc.add_paragraph(
    'The 335 SMAs were distributed across 47 prefectures (median 7 SMAs per prefecture, '
    'range 3\u201321). Of 335 SMAs, 64 (19.1%) contained at least one university hospital. '
    'General anesthesia (L008) SCR data were available for 334 of 335 SMAs (one masked); '
    'epidural anesthesia (L002) data for 307 SMAs (28 masked, reflecting lower volume in '
    'some rural areas). Table 2 summarizes the distribution of SCRs for each code.'
)

# Table 1: Anesthesia codes (brief)
add_para('Table 1. Anesthesia procedure codes analyzed.', bold=True)
table1 = doc.add_table(rows=1, cols=4)
table1.style = 'Table Grid'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr1 = table1.rows[0].cells
for i, text in enumerate(['Code', 'Procedure', 'Clinical significance', 'SMAs with data']):
    hdr1[i].paragraphs[0].add_run(text).bold = True
    hdr1[i].paragraphs[0].runs[0].font.size = Pt(9)

codes_data = [
    ['L008', 'Closed-circuit general anesthesia', 'Primary indicator of GA volume', '334'],
    ['L002', 'Epidural anesthesia', 'Regional technique (alone or combined)', '307'],
    ['L003', 'Continuous epidural infusion', 'Indicator of combined GA\u2013epidural technique', '331'],
    ['L004', 'Spinal anesthesia', 'Alternative to GA for suitable procedures', '334'],
    ['L009', 'Anesthesia management fee I', 'Proxy for specialist anesthesiologist staffing', '314'],
    ['L100', 'Nerve block (inpatient)', 'Indicator of pain clinic activity', '335'],
]
for row_data in codes_data:
    add_table_row(table1, row_data)
doc.add_paragraph()

# Table 2: Descriptive statistics (UPDATED numbers)
add_para('Table 2. Distribution of standardized claim ratios (SCRs) across secondary medical areas (national average = 100).', bold=True)
table2 = doc.add_table(rows=1, cols=8)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0].cells
for i, text in enumerate(['Code', 'n', 'Mean', 'SD', 'Median', 'IQR', 'Range', 'CV (%)']):
    hdr2[i].paragraphs[0].add_run(text).bold = True
    hdr2[i].paragraphs[0].runs[0].font.size = Pt(9)

# Updated from corrected analysis
table2_data = [
    ['L008', '334', '79.7', '42.7', '73.2', '47.9\u2013102.7', '2.3\u2013435.7', '53.6'],
    ['L002', '307', '95.6', '83.2', '73.9', '30.3\u2013133.3', '0.9\u2013519.1', '87.0'],
    ['L003', '331', '86.1', '55.9', '73.1', '44.6\u2013115.6', '1.5\u2013350.2', '64.9'],
    ['L004', '334', '93.8', '53.2', '84.3', '55.1\u2013121.7', '1.0\u2013346.5', '56.8'],
    ['L009', '314', '87.7', '49.4', '79.1', '51.8\u2013115.2', '1.2\u2013296.3', '56.4'],
    ['L100', '335', '99.7', '76.2', '72.2', '43.3\u2013133.7', '1.4\u2013482.5', '76.4'],
]
for row_data in table2_data:
    add_table_row(table2, row_data)
doc.add_paragraph()
add_para('CV = coefficient of variation; IQR = interquartile range; SD = standard deviation.', italic=True)

add_heading_styled('Regional variation in anesthesia practice', level=2)

doc.add_paragraph(
    'Substantial variation was observed across all anesthesia codes (Table 2). General '
    'anesthesia (L008) SCRs ranged from 2.3 to 435.7 across 334 SMAs (CV 53.6%), '
    'representing a 189-fold difference between the lowest- and highest-volume areas. '
    'Epidural anesthesia (L002) showed even greater relative variation (CV 87.0%), as did '
    'continuous epidural infusion (L003, CV 64.9%), the direct indicator of combined '
    'GA\u2013epidural technique. Spinal anesthesia (L004) had a CV of 56.8%. Figure 1 shows the '
    'geographic distribution of SCRs, revealing distinct regional clustering.'
)

# [Major 2] Multilevel model results - NEW
add_heading_styled('Multilevel model results', level=2)

# Table 3: Multilevel model results (NEW)
add_para('Table 3. Multilevel linear mixed model results: SCR as outcome, prefecture as random intercept.', bold=True)
table3 = doc.add_table(rows=1, cols=7)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr3 = table3.rows[0].cells
for i, text in enumerate(['Code', 'n', 'ICC (null)', '\u03b2 univ (95% CI)', 'P value', 'ICC (adj)', 'Marginal R\u00b2']):
    hdr3[i].paragraphs[0].add_run(text).bold = True
    hdr3[i].paragraphs[0].runs[0].font.size = Pt(9)

# Results from multilevel_analysis.py
table3_data = [
    ['L008', '334', '0.058', '+64.1 (54.8 to 73.4)', '<0.001', '0.048', '0.358'],
    ['L004', '334', '0.343', '+23.0 (11.0 to 35.1)', '<0.001', '0.360', '0.020'],
    ['L002', '307', '0.191', '+49.0 (27.2 to 70.8)', '<0.001', '0.057', '0.125'],
    ['L003', '331', '0.106', '+55.8 (43.0 to 68.7)', '<0.001', '0.142', '0.162'],
    ['L001', '327', '0.094', '+83.2 (66.3 to 100.0)', '<0.001', '0.133', '0.204'],
    ['L100 (inpt)', '335', '0.185', '+44.8 (25.9 to 63.6)', '<0.001', '0.218', '0.035'],
    ['L100 (outpt)', '335', '0.299', '+37.2 (22.4 to 52.0)', '<0.001', '0.322', '0.044'],
]
for row_data in table3_data:
    add_table_row(table3, row_data)
doc.add_paragraph()
add_para(
    'ICC = intraclass correlation coefficient (proportion of variance at prefecture level); '
    '\u03b2 univ = fixed effect coefficient for university hospital presence; '
    'Marginal R\u00b2 = proportional reduction in total variance from null model. '
    'All models include prefecture random intercept.',
    italic=True
)

doc.add_paragraph(
    'The null multilevel model showed that prefecture-level clustering accounted for a '
    'modest proportion of general anesthesia (L008) variance (ICC = 0.058), indicating that '
    '94.2% of variation occurs within prefectures, where audit policy is uniform. In contrast, '
    'spinal anesthesia (L004) showed much stronger prefecture-level clustering (ICC = 0.343), '
    'suggesting that prefectural factors (possibly including audit practice) play a larger role '
    'for this code. Nerve block SCR (outpatient) also showed substantial prefecture clustering '
    '(ICC = 0.299).'
)

doc.add_paragraph(
    'Adding university hospital presence as a fixed effect produced the largest improvement '
    'for general anesthesia (L008: marginal R\u00b2 = 0.358; \u03b2 = +64.1, 95% CI 54.8\u201373.4), '
    'meaning that university hospital presence alone explains 35.8% of total L008 variance. '
    'The effect was statistically significant for all codes (all P < 0.001) but varied in '
    'magnitude: largest for intravenous anesthesia (L001: \u03b2 = +83.2) and general anesthesia '
    '(L008: \u03b2 = +64.1), moderate for continuous epidural (L003: \u03b2 = +55.8) and epidural '
    '(L002: \u03b2 = +49.0), and smallest for spinal anesthesia (L004: \u03b2 = +23.0).'
)

add_heading_styled('Sensitivity analysis 1: within-prefecture variance decomposition', level=2)

doc.add_paragraph(
    'Hierarchical variance decomposition attributed 14.5% of L008 variance to '
    'between-prefecture differences, 40.5% to the university hospital effect (within '
    'prefecture), and 45.0% to residual within-group variation. The university hospital '
    'effect alone\u2014a single binary variable\u2014explained 47.4% of all within-prefecture '
    'variance in general anesthesia volume. This pattern was consistent across codes, '
    'with university hospital effect explaining 30.5% of within-prefecture L009 variance, '
    '21.5% of L002 variance, and 15.9% of L004 variance.'
)

doc.add_paragraph(
    'In within-prefecture comparisons (which control for audit policy), university hospital '
    'SMAs had higher L008 SCRs in all 47 of 47 prefectures (100%), with a mean within-'
    'prefecture difference of +63.3 points (SD 32.7, paired t = 13.28, P < 0.001). For '
    'anesthesia management fee I (L009, proxy for specialist staffing), the pattern held in '
    '41 of 47 prefectures (87%). For epidural anesthesia (L002), 40 of 47 prefectures (85%).'
)

add_heading_styled('Sensitivity analysis 2: cross-code correlation', level=2)

doc.add_paragraph(
    'General anesthesia (L008) and spinal anesthesia (L004) showed a positive correlation '
    '(r = +0.235, P < 0.001). This finding is inconsistent with the hypothesis that audit-driven '
    'reclassification shifts claims between these two codes, which would predict a negative '
    'correlation. The positive correlation instead suggests that areas with higher overall '
    'surgical/anesthesia capacity tend to have higher utilization of both techniques. '
    'General anesthesia (L008) and epidural anesthesia (L002) also showed a positive correlation '
    '(r = +0.319, P < 0.001), consistent with a common supply factor (anesthesia department '
    'capacity) rather than audit-driven substitution.'
)

add_heading_styled('Sensitivity analysis 3: quantitative audit impact', level=2)

doc.add_paragraph(
    'The maximum difference in prefectural audit rates was 0.21 percentage points '
    '(0.07% to 0.28%). Even if this entire difference were concentrated on a single '
    'anesthesia code, it would produce an SCR shift of approximately 0.2 points\u2014less '
    'than 0.3% of the observed interquartile range for L008 and less than 0.1% of the '
    'range for L002. Audit rate variation is quantitatively incapable of explaining the '
    'observed regional differences.'
)

add_heading_styled('Combined SCR analysis', level=2)

doc.add_paragraph(
    'Combining L008 and L004 SCRs reduced the CV from 53.6% to 43.4%\u2014a modest 19% '
    'reduction\u2014indicating that any potential spinal\u2013GA substitution (whether from audit '
    'reclassification, defensive undercoding, or clinical choice) accounts for a minority '
    'of GA variation. The university hospital effect remained strong on the combined measure: '
    'university areas had a mean combined SCR of 242.9 versus 157.6 for non-university areas '
    '(Cohen\u2019s d = 1.26, P < 0.001).'
)

# [Major 3] Empirical Bayes results
add_heading_styled('Small-sample sensitivity: empirical Bayes shrinkage', level=2)

doc.add_paragraph(
    'Empirical Bayes shrinkage toward prefecture means attenuated effect sizes modestly but '
    'did not materially alter any conclusions (Table 4). For general anesthesia (L008), the '
    'university hospital effect decreased from Cohen\u2019s d = 1.89 (raw) to d = 1.72 (EB), '
    'a 9.0% attenuation. For spinal anesthesia (L004), attenuation was 9.1% (d = 0.40 \u2192 0.36). '
    'For epidural anesthesia (L002), attenuation was 8.6% (d = 0.59 \u2192 0.54). In all cases, '
    'effect sizes remained large and statistically significant after shrinkage, confirming '
    'that the findings are not driven by unstable estimates in small-volume areas.'
)

doc.add_paragraph(
    'Two SMAs were identified as outliers (>3 SD from the mean L008 SCR): '
    'Awa in Chiba (SCR = 216.7) and Ku-Chuobu in Tokyo (SCR = 435.7, containing 5 university '
    'hospitals). For L002, five outlier SMAs were identified, including Kamikawachu in Hokkaido '
    '(SCR = 359.0, containing Asahikawa Medical University) and Iizuka in Fukuoka (SCR = 519.1). '
    'Excluding these outliers did not change the direction or statistical significance of any '
    'finding.'
)

add_heading_styled('University hospital effect', level=2)

doc.add_paragraph(
    'University hospital presence was the strongest single predictor of anesthesia practice '
    'patterns (Figure 2). For general anesthesia (L008), the 64 university hospital SMAs had '
    'a mean SCR of 132.1 (SD 47.5) compared with 67.3 (SD 30.5) for the 270 non-university '
    'SMAs (Cohen\u2019s d = 1.88, Welch\u2019s t = 10.40, P < 0.001; Mann\u2013Whitney P < 0.001). '
    'A dose\u2013response relationship was apparent in Tokyo, where the number of university '
    'hospitals per SMA ranged from 0 to 5: SCRs increased monotonically with university '
    'hospital count.'
)

add_heading_styled('Combined general\u2013epidural anesthesia', level=2)

doc.add_paragraph(
    'Continuous epidural infusion (L003), the direct indicator of combined GA\u2013epidural '
    'technique, showed a CV of 64.9%. L003 SCR was 1.73 times higher in university hospital '
    'SMAs than in non-university SMAs. Three audit-hypothesis tests applied to the epidural '
    'combination (L008 + L002) all rejected the audit explanation: positive L008\u2013L002 '
    'correlation (r = +0.319), and persistent university hospital effect (d = 1.26 on '
    'combined L008 + L004).'
)

# ═══════════════════════════════════════════════════════════
# DISCUSSION
# ═══════════════════════════════════════════════════════════

add_heading_styled('DISCUSSION', level=1)

add_heading_styled('Principal findings', level=2)

doc.add_paragraph(
    'This study demonstrates that regional variation in anesthesia practice across Japan\u2019s '
    '335 secondary medical areas is substantial and structural. Multilevel models showed that '
    'prefecture-level clustering accounts for only 5.8% of general anesthesia variance '
    '(ICC = 0.058), meaning that 94.2% of variation occurs within prefectures where audit '
    'policy is uniform. University hospital presence is the dominant determinant, explaining '
    '35.8% of total L008 variance in the multilevel model and 40.5% of within-prefecture '
    'variance in the hierarchical decomposition. The finding that this effect is present in '
    'all 47 prefectures\u2014despite differing audit policies\u2014provides compelling evidence for '
    'a supply-side, institutional mechanism. Empirical Bayes shrinkage attenuated the '
    'university hospital effect by only 9%, confirming that the results are not driven by '
    'unstable estimates in small-volume areas.'
)

add_heading_styled('The Japanese healthcare system context', level=2)

doc.add_paragraph(
    'Interpreting these findings requires understanding the Japanese healthcare system. '
    'Japan\u2019s universal health insurance covers the entire population under a single, '
    'nationally uniform fee schedule.\u00b9\u00b7\u00b2 The WHO HAQ Index ranked Japan at 87.5/100 '
    '(10th decile) in 2019,\u00b3 reflecting high-quality healthcare access. The 335 SMAs used '
    'as our unit of analysis are designed to be self-sufficient for inpatient care, and '
    'cross-boundary patient movement is limited: the 2008 Patient Survey found that only '
    '6.1% of inpatients received care outside their home prefecture.\u2076 Within prefectures, '
    'cross-SMA movement is predominantly to adjacent areas,\u2077 meaning that SCR variation '
    'between non-adjacent SMAs within the same prefecture predominantly reflects local practice '
    'patterns rather than patient sorting. The residence-based SCR calculation further mitigates '
    'this concern, as claims are attributed to the patient\u2019s home area regardless of where '
    'treatment was received.'
)

add_heading_styled('Comparison with existing literature', level=2)

doc.add_paragraph(
    'Our findings are consistent with the broader literature on medical practice variation. '
    'The Dartmouth Atlas project in the United States has documented extensive regional variation '
    'in surgical rates driven primarily by physician supply and practice style rather than '
    'patient need.\u00b2\u2078\u00b7\u00b2\u2079 Similar patterns have been reported in the United Kingdom,\u00b3\u2070 '
    'Germany,\u00b3\u00b9 and Australia.\u00b3\u00b2 Our study extends this literature by exploiting Japan\u2019s '
    'unique institutional structure to distinguish administrative from clinical sources of '
    'variation, and by providing the first multilevel analysis of anesthesia technique variation '
    'at a fine geographic scale (335 areas) in an East Asian universal insurance system.'
)

doc.add_paragraph(
    'The university hospital effect we document (Cohen\u2019s d = 1.88 for general anesthesia) is '
    'remarkably large compared with effect sizes typically reported in medical practice variation '
    'research. This likely reflects the Japanese medical education system, in which university '
    'departments (ikyoku) exert substantial influence over clinical practice in affiliated '
    'hospitals within their geographic sphere.\u00b3\u00b3\u00b7\u00b3\u2074 The ICC of 0.058 for L008 is notably '
    'lower than ICCs reported for surgical procedure rates in the United States (typically '
    '0.10\u20130.30),\u00b2\u2078 suggesting that prefectural-level factors (including audit policy) play '
    'a relatively minor role for general anesthesia compared with within-prefecture institutional '
    'factors.'
)

add_heading_styled('Implications for oncologic outcomes', level=2)

doc.add_paragraph(
    'The finding that combined GA\u2013epidural technique (L003) shows a CV of 64.9% is clinically '
    'important in light of evidence suggesting that regional anesthesia combined with general '
    'anesthesia may improve recurrence-free survival in cancer surgery.\u00b9\u00b2\u2013\u00b9\u2077 The 1.73-fold '
    'difference in L003 SCR between university and non-university hospital areas suggests that '
    'patients\u2019 access to this potentially beneficial technique depends substantially on where '
    'they live\u2014a finding that challenges the equity premise of Japan\u2019s universal insurance '
    'system.\u00b9\u00b7\u00b2'
)

add_heading_styled('Strengths and limitations', level=2)

doc.add_paragraph(
    'Strengths include the use of age- and sex-standardized ratios covering the entire national '
    'population; analysis at a fine geographic scale (335 areas); multilevel modeling that '
    'properly accounts for the nested data structure; empirical Bayes shrinkage to address '
    'small-sample instability; multiple independent sensitivity analyses; within-prefecture '
    'comparisons that control for audit policy; and residence-based SCR calculation that '
    'mitigates patient travel effects.'
)

doc.add_paragraph(
    'Limitations include the ecological fallacy inherent in area-level analysis;\u00b3\u2075 the '
    'inability to distinguish individual patient-level decisions from aggregate patterns; '
    'the post-audit nature of SCR data (which reflects reimbursed rather than intended '
    'practice, though our sensitivity analyses suggest this distinction is quantitatively '
    'minor); the cross-sectional design, which precludes causal inference; incomplete '
    'separation of the university hospital effect from urban concentration (university '
    'hospitals are typically in the largest city of each prefecture); the absence of '
    'code-specific audit rate data; and the inability to quantify \u201cdefensive undercoding\u201d '
    '\u2014where clinicians preemptively avoid claiming procedures they expect to be audited\u2014'
    'from claims data alone.\u00b3\u2076 Additionally, the multilevel models included only university '
    'hospital presence as a fixed effect; future studies with additional area-level covariates '
    '(population density, urbanization indices, bed density) may explain additional variance.'
)

add_heading_styled('Conclusions', level=2)

doc.add_paragraph(
    'Regional variation in anesthesia practice in Japan is genuine and large, driven primarily '
    'by institutional factors\u2014particularly university hospital proximity\u2014rather than by '
    'differential insurance auditing. Multilevel modeling confirms that prefecture-level '
    'clustering explains only 5.8% of general anesthesia variance, while university hospital '
    'presence explains 35.8%. Empirical Bayes shrinkage confirms that these findings are '
    'robust to small-sample instability. The variation in combined GA\u2013epidural anesthesia '
    'is of particular concern given emerging evidence of oncologic benefit. Policy responses '
    'could include targeted education programs, specialist outreach from university centers, '
    'and monitoring of anesthesia technique patterns as a quality indicator.\u00b3\u2077 The multilevel '
    'sensitivity analysis framework developed here is applicable to investigating practice '
    'variation in any procedure under Japan\u2019s universal insurance system.'
)

# ═══════════════════════════════════════════════════════════
# REFERENCES (expanded with new citations)
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_styled('REFERENCES', level=1)

refs = [
    # Healthcare system context
    '1. Ikegami N, Yoo B-K, Hashimoto H, et al. Japanese universal health coverage: evolution, '
    'achievements, and challenges. Lancet 2011;378:1106\u201315.',

    '2. Ikegami N. Universal Health Coverage for Inclusive and Sustainable Development: Lessons from '
    'Japan. Washington, DC: World Bank, 2014.',

    '3. GBD 2019 Healthcare Access and Quality Collaborators. Assessing performance of the Healthcare '
    'Access and Quality Index, overall and by select age groups, for 204 countries and territories, '
    '1990\u20132019: a systematic analysis from the Global Burden of Disease Study 2019. Lancet Glob '
    'Health 2022;10:e1715\u201343.',

    '4. World Health Organization. Tracking Universal Health Coverage: 2023 Global Monitoring Report. '
    'Geneva: WHO, 2023.',

    '5. Ministry of Health, Labour and Welfare. Medical Care Plan Guidelines (Iryo Keikaku Sakutei '
    'Shishin). Tokyo: MHLW, 2023.',

    '6. Ministry of Health, Labour and Welfare. Patient Survey 2008 (Kanja Chosa). '
    'https://www.mhlw.go.jp/toukei/saikin/hw/kanja/08/ (accessed 15 Mar 2026).',

    '7. Tokyo Metropolitan Government. Regional Medical Care Plan: Cross-Area Patient Movement '
    'Analysis. Tokyo: Bureau of Social Welfare and Public Health, 2025.',

    # NDB / claims data
    '8. Matsuda S, Fujimori K. The claim database in Japan. Asian Pac J Dis Manag 2012;6:55\u201359.',

    '9. Suto M, Iba A, Sugiyama T, et al. Literature review of studies using the National Database '
    'of Health Insurance Claims of Japan (NDB): limitations and strategies in using the NDB for '
    'research. JMA J 2023;7:10\u201320.',

    '10. Hamada H, Sekimoto M, Imanaka Y. Effects of the per diem prospective payment system with '
    'DRG-like grouping system (DPC/PDPS) on resource usage and healthcare quality in Japan. Health '
    'Policy 2012;107:194\u2013201.',

    # Audit data
    '11. Social Insurance Medical Fee Payment Fund (Shiharai Kikin). Review Statistics, Fiscal Year '
    '2022. https://www.ssk.or.jp/tokeijoho/shinsatokei/index.html (accessed 15 Mar 2026).',

    # Cancer outcomes
    '12. Pei L, Tan G, Wang L, et al. Anesthetic techniques for risk of malignant tumor recurrence. '
    'Cochrane Database Syst Rev 2019;(11):CD008877.',

    '13. Sessler DI, Pei L, Huang Y, et al. Recurrence of breast cancer after regional or general '
    'anesthesia: a randomized controlled trial. Lancet 2019;394:1807\u201315.',

    '14. Exadaktylos AK, Buggy DJ, Moriarty DC, et al. Can anesthetic technique for primary breast '
    'cancer surgery affect recurrence or metastasis? Anesthesiology 2006;105:660\u20134.',

    '15. Sun Y, Li T, Gan TJ. The effects of perioperative regional anesthesia and analgesia on cancer '
    'recurrence and survival after oncology surgery: a systematic review and meta-analysis. Reg '
    'Anesth Pain Med 2015;40:589\u201398.',

    '16. Xie S, Li L, Meng F, et al. Regional anesthesia might reduce recurrence and metastasis rates '
    'in adult patients with cancers after surgery: a meta-analysis. BMC Anesthesiol 2024;24:19.',

    '17. Zhang D, Jiang J, Liu J, et al. Effects of perioperative epidural analgesia on cancer '
    'recurrence and survival. Front Oncol 2021;11:798435.',

    # Fee schedule revision
    '18. Ministry of Health, Labour and Welfare. Revision of medical fee schedule, 2026 '
    '(Reiwa 8 nendo shinryo houshu kaitei). https://www.mhlw.go.jp (accessed 15 Mar 2026).',

    # SMA structure
    '19. Matsuda S. Secondary medical areas and regional healthcare planning in Japan (Nihon no '
    'ni-ji iryo-ken to chiiki iryo keikaku). J Health Care Soc 2018;28:137\u201348. [in Japanese]',

    # US comparison
    '20. Wennberg JE, Cooper M. The Dartmouth Atlas of Health Care 1999. Chicago: '
    'American Hospital Publishing, 1999.',

    # Data sources
    '21. Cabinet Office. Regional variation visualization (chiikisa no mieruka). '
    'https://www5.cao.go.jp/keizai-shimon/kaigi/special/reform/mieruka/ (accessed 10 Mar 2026).',

    '22. Ministry of Health, Labour and Welfare. Survey of Physicians, Dentists and Pharmacists, '
    '2022. e-Stat. https://www.e-stat.go.jp (accessed 12 Mar 2026).',

    '23. Ministry of Land, Infrastructure, Transport and Tourism. National Land Numerical Information '
    'download service. https://nlftp.mlit.go.jp (accessed 10 Mar 2026).',

    # Statistical methods
    '24. Cohen J. Statistical Power Analysis for the Behavioral Sciences. 2nd ed. '
    'Hillsdale, NJ: Lawrence Erlbaum, 1988.',

    '25. Rabe-Hesketh S, Skrondal A. Multilevel and Longitudinal Modeling Using Stata. 3rd ed. '
    'College Station, TX: Stata Press, 2012.',

    '26. Seabold S, Perktold J. Statsmodels: econometric and statistical modeling with Python. '
    'In: Proceedings of the 9th Python in Science Conference. 2010:92\u201396.',

    '27. Efron B, Morris C. Stein\u2019s estimation rule and its competitors\u2014an empirical Bayes '
    'approach. J Am Stat Assoc 1973;68:117\u201330.',

    # Practice variation literature
    '28. Wennberg JE. Tracking Medicine: A Researcher\u2019s Quest to Understand Health Care. '
    'New York: Oxford University Press, 2010.',

    '29. Wennberg JE, Fisher ES, Skinner JS. Geography and the debate over Medicare reform. '
    'Health Aff (Millwood) 2002;Suppl Web Exclusives:W96\u2013114.',

    '30. NHS RightCare. The NHS Atlas of Variation in Healthcare, 2015. London: Public Health England.',

    '31. Wengler A, Nimptsch U, Mansky T. Hip and knee replacement surgery (arthroplasty) and '
    'geographic variations in Germany. Dtsch Arztebl Int 2014;111:407\u201316.',

    '32. Australian Commission on Safety and Quality in Health Care. Australian Atlas of Healthcare '
    'Variation. Sydney: ACSQHC, 2015.',

    # Ikyoku system
    '33. Otsuka T. The ikyoku system of university orthopedic surgery departments: an in-hospital '
    'organizational system unique to Japan. J Orthop Sci 2012;17:513\u201314.',

    '34. Onishi H. History of Japanese medical education. Korean J Med Educ 2018;30:283\u201394.',

    # Ecological fallacy
    '35. Greenland S. Ecologic versus individual-level sources of bias in ecologic estimates of '
    'contextual health effects. Int J Epidemiol 2001;30:1343\u201350.',

    # Defensive undercoding
    '36. Sekimoto M, Imanaka Y, Kobayashi H, et al. Impact of hospital characteristics on the cost '
    'and practice patterns of acute myocardial infarction patients in Japan. Health Policy '
    '2006;78:34\u201345.',

    # Policy
    '37. Merry AF, Cooper JB, Soyannwo O, et al. International standards for a safe practice of '
    'anesthesia 2010. Can J Anesth 2010;57:1027\u201334.',
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.name = 'Arial'

# ═══════════════════════════════════════════════════════════
# FIGURES
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_styled('FIGURES', level=1)

# [Minor 4] Corrected terminology: choropleth = 2D, only extruded = 3D
# [Minor 5] Figures consolidated into multi-panel
figure_list = [
    ('/home/ubuntu/rapm_fig1_en.png',
     'Figure 1. Geographic distribution of anesthesia standardized claim ratios (SCRs) by secondary '
     'medical area, fiscal year 2022. (A) General anesthesia (L008) SCR. (B) Spinal anesthesia (L004) '
     'SCR. Choropleth maps with solid lines indicating prefectural boundaries and dashed lines '
     'indicating SMA boundaries. Northern Territories shown in white (no SMA designated). '
     'National average = 100 for each code.'),
    ('/home/ubuntu/rapm_fig2_en.png',
     'Figure 2. University hospital effect and combined SCR analysis. (A) University hospital presence '
     'by SMA: red circles indicate SMAs containing one or more university hospitals (64 of 335 SMAs). '
     '(B) Combined GA + spinal anesthesia (L008 + L004) SCR: combining these codes neutralizes potential '
     'audit-driven reclassification, yet substantial regional variation persists (CV = 43.4%). '
     'Choropleth maps.'),
    ('/home/ubuntu/rapm_fig3_en.png',
     'Figure 3. Combined GA\u2013epidural anesthesia variation. (A) Continuous epidural infusion (L003) '
     'SCR\u2014direct indicator of combined GA\u2013epidural technique (CV = 64.9%). (B) L003/L008 ratio '
     '(epidural combination rate per GA case). Higher values indicate greater use of combined '
     'GA\u2013epidural technique. Choropleth maps.'),
    ('/home/ubuntu/rapm_fig4_en.png',
     'Figure 4. Three-dimensional extruded map. Color represents L003/L008 ratio (combined epidural '
     'rate); height represents anesthesiologist count per SMA. Areas with both high epidural combination '
     'rate and high specialist count appear as tall, green-colored polygons, typically corresponding to '
     'university hospital areas in metropolitan regions. Northern Territories shown at base height.'),
]

for path, caption in figure_list:
    add_figure(path, caption, width=Inches(5.5))
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# STROBE CHECKLIST
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_styled('STROBE Checklist (Cross-sectional study)', level=1)

strobe_items = [
    ('1a', 'Title and abstract', 'Title page and structured abstract'),
    ('1b', 'Informative abstract', 'Structured abstract with Background/Methods/Results/Conclusion'),
    ('2', 'Background/rationale', 'Introduction paragraphs 1\u20134'),
    ('3', 'Objectives', 'Introduction paragraph 5'),
    ('4', 'Study design', 'Methods: Study design and ethical considerations'),
    ('5', 'Setting', 'Methods: Geographic units; Data sources (335 SMAs, FY2022)'),
    ('6a', 'Participants/eligibility', 'N/A (ecological study using aggregate data)'),
    ('7', 'Variables', 'Methods: Anesthesia procedure codes; Table 1'),
    ('8', 'Data sources', 'Methods: Data sources and standardized claim ratios'),
    ('9', 'Bias', 'Methods: Sensitivity analyses; Discussion: Limitations'),
    ('10', 'Study size', 'Methods: 335 SMAs, 47 prefectures'),
    ('11', 'Quantitative variables', 'Methods: Statistical analysis'),
    ('12', 'Statistical methods', 'Methods: Statistical analysis (multilevel models, ICC, EB shrinkage)'),
    ('13', 'Participants', 'Results: 334 SMAs with L008 data'),
    ('14', 'Descriptive data', 'Results: Table 2'),
    ('15', 'Outcome data', 'Results: Regional variation; Multilevel model results'),
    ('16', 'Main results', 'Results: Table 3; Sensitivity analyses 1\u20133; University hospital effect'),
    ('17', 'Other analyses', 'Results: Combined SCR; EB shrinkage; Combined GA\u2013epidural'),
    ('18', 'Key results', 'Discussion: Principal findings'),
    ('19', 'Limitations', 'Discussion: Strengths and limitations'),
    ('20', 'Interpretation', 'Discussion: Comparison with existing literature'),
    ('21', 'Generalizability', 'Discussion: Japanese healthcare system context; Conclusions'),
    ('22', 'Funding', '[To be completed]'),
]

strobe_table = doc.add_table(rows=1, cols=3)
strobe_table.style = 'Table Grid'
strobe_hdr = strobe_table.rows[0].cells
for i, text in enumerate(['Item', 'STROBE recommendation', 'Location in manuscript']):
    strobe_hdr[i].paragraphs[0].add_run(text).bold = True
    strobe_hdr[i].paragraphs[0].runs[0].font.size = Pt(9)

for item in strobe_items:
    add_table_row(strobe_table, item)

# ═══════════════════════════════════════════════════════════
# DECLARATIONS
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_styled('DECLARATIONS', level=1)

add_heading_styled('Ethics approval', level=2)
doc.add_paragraph(
    'This study used publicly available aggregate data. No individual-level data were accessed. '
    'Ethics committee approval was not required under the Ethical Guidelines for Medical and '
    'Biological Research Involving Human Subjects (Japan, 2021 revision).'
)

add_heading_styled('Data availability', level=2)
doc.add_paragraph(
    'All data used in this study are publicly available. SCR data: Cabinet Office regional variation '
    'visualization (https://www5.cao.go.jp/keizai-shimon/kaigi/special/reform/mieruka/). '
    'Physician statistics: e-Stat (https://www.e-stat.go.jp). GIS boundaries: National Land '
    'Numerical Information (https://nlftp.mlit.go.jp). Analysis code is available at '
    'https://github.com/bougtoir/wip.'
)

add_heading_styled('Competing interests', level=2)
doc.add_paragraph('[To be completed by authors.]')

add_heading_styled('Funding', level=2)
doc.add_paragraph('[To be completed by authors.]')

add_heading_styled('Contributors', level=2)
doc.add_paragraph('[To be completed by authors.]')

add_heading_styled('AI statement', level=2)
doc.add_paragraph(
    'Devin (Cognition AI) was used as a coding assistant for data analysis, statistical modeling, '
    'geographic information system visualization, and manuscript preparation. All scientific '
    'interpretation and clinical judgment were provided by the human authors. The AI-generated '
    'code and text were reviewed and validated by the authors.'
)

add_heading_styled('Transparency declaration', level=2)
doc.add_paragraph(
    'The lead author (the manuscript\u2019s guarantor) affirms that the manuscript is an honest, '
    'accurate, and transparent account of the study being reported; that no important aspects '
    'of the study have been omitted; and that any discrepancies from the study as planned have '
    'been explained.'
)

# ═══════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════

output_path = '/home/ubuntu/regional_anesthesia_REVISED_EN.docx'
doc.save(output_path)
print(f"Saved: {output_path}")
