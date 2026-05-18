#!/usr/bin/env python3
"""Create IJQHC-formatted English manuscript.

Conforms to International Journal for Quality in Health Care author guidelines:
- UK English
- Double-blind (anonymised); title page kept separate (in End-Matter file)
- Title <=150 chars, header <=30 chars
- Structured abstract (Background/Methods/Results/Conclusion) <=400 words
- Up to 6 keywords (MeSH-preferred)
- Body <=3000 words; <=30 references; <=5 tables/figures combined
- Discussion subheadings: Statement of principal findings; Strengths and
  limitations; Interpretation within the context of the wider literature;
  Implications for policy, practice and research
- Double-spaced throughout, no footnotes/endnotes
- Figures placed after references and tables (not embedded)
- Word-native font superscript for citations via {n} markers
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'documents', 'IJQHC')
os.makedirs(OUTPUT_DIR, exist_ok=True)

REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FIG_DIR = os.path.join(REPO_ROOT, 'output')

doc = Document()

# Page setup: A4, 2.54 cm margins
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# Style: Times New Roman 12 pt, double-spaced (IJQHC requirement)
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0
style.paragraph_format.space_after = Pt(0)


def add_run_with_refs(paragraph, text, italic=False, bold=False):
    """Add text to a paragraph, parsing {n} or {n-m} as font superscript runs."""
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        if italic:
            run.italic = True
        if bold:
            run.bold = True
        if part.startswith('{') and part.endswith('}'):
            inner = part[1:-1]
            run.text = inner
            run.font.superscript = True


def add_para(text, bold=False, italic=False, align=None, space_before=0, space_after=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after:
        p.paragraph_format.space_after = Pt(space_after)
    add_run_with_refs(p, text, italic=italic, bold=bold)
    return p


def add_heading(text, level=1, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text.upper() if level == 1 else text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12) if level > 1 else Pt(13)
    run.bold = True
    return p


def add_subheading(text):
    return add_heading(text, level=2, space_before=8, space_after=4)


def add_blank():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    return p


# ============================================================
# TITLE PAGE
# ============================================================
title = ("Regional variation in anaesthesia practice in Japan: "
         "structural determinant or claims-audit artefact?")
assert len(title) <= 150, f"Title is {len(title)} chars (max 150)"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
run = p.add_run(title)
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.bold = True

header = "Anaesthesia variation in Japan"
assert len(header) <= 30, f"Header is {len(header)} chars (max 30)"
add_para(f"Running header: {header}", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_blank()

add_para("[Authors and affiliations to be supplied]", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("[Corresponding author: name, postal address, telephone, email]",
         italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_blank()

add_para("Article type: Original research article", italic=True)
add_para("Word count (main text, excluding abstract, references, tables and figures): "
         "approximately 3000 words", italic=True)
add_para("Tables: 3   Figures: 2   References: 30", italic=True)
add_para("Reporting guideline: STROBE checklist for cross-sectional studies "
         "(uploaded as supplementary material)", italic=True)

doc.add_page_break()

# ============================================================
# STRUCTURED ABSTRACT (<= 400 words)
# ============================================================
add_para(title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_blank()
add_heading("Abstract", level=1, space_before=0)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
add_run_with_refs(p, "Background  ", bold=True)
add_run_with_refs(
    p,
    "Regional variation in anaesthesia practice has been documented in many "
    "high-income countries, but the relative contributions of clinical "
    "preference, institutional capacity and administrative coding practices "
    "remain unclear. Japan provides a distinctive natural experiment: a single "
    "national fee schedule combined with prefecture-specific insurance "
    "auditing. We aimed to quantify regional variation in anaesthesia practice "
    "across Japan's 335 secondary medical areas and to test whether observed "
    "variation reflects structural determinants or differential auditing.")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
add_run_with_refs(p, "Methods  ", bold=True)
add_run_with_refs(
    p,
    "Cross-sectional ecological study using publicly available standardised "
    "claim ratios (age- and sex-adjusted observed-to-expected ratios on a "
    "residence basis, national average = 100) for fiscal year 2022. Six "
    "anaesthesia procedure codes were analysed across 335 secondary medical "
    "areas nested within 47 prefectures. We fitted multilevel linear mixed "
    "models with prefectures as random intercepts to estimate intraclass "
    "correlation coefficients and the contribution of university hospital "
    "presence. Three pre-specified sensitivity analyses tested the audit "
    "hypothesis: within-prefecture variance decomposition; cross-code "
    "correlation; and quantitative estimation of the maximum impact of "
    "audit-rate differences. Empirical Bayes shrinkage addressed potential "
    "instability in low-volume areas.")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
add_run_with_refs(p, "Results  ", bold=True)
add_run_with_refs(
    p,
    "Substantial variation was observed across all anaesthesia codes. The "
    "coefficient of variation was 53.6% for general anaesthesia, 87.0% for "
    "epidural anaesthesia (L002), 64.9% for continuous epidural infusion "
    "(L003) and 56.8% for spinal "
    "anaesthesia. Multilevel models showed that only 5.8% of general "
    "anaesthesia variance was attributable to the prefecture level (intraclass "
    "correlation coefficient 0.058); 94.2% occurred within prefectures, where "
    "audit policy is uniform. University hospital presence explained 35.8% of "
    "total general anaesthesia variance and 40.5% of within-prefecture "
    "variance (Cohen's d 1.88). University hospital areas had higher general "
    "anaesthesia ratios in 47 of 47 prefectures. The three audit sensitivity "
    "analyses converged: the maximum audit-rate difference (0.21 percentage "
    "points) could shift ratios by less than 0.3% of the observed "
    "interquartile range. Empirical Bayes shrinkage of the area-level ratios "
    "toward their prefecture mean substantially attenuated effect sizes "
    "(general anaesthesia by 66.5%, from Cohen's d = 1.88 to 0.63); the "
    "university hospital effect on general anaesthesia remained moderate "
    "and statistically significant even after this heavy shrinkage.")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
add_run_with_refs(p, "Conclusion  ", bold=True)
add_run_with_refs(
    p,
    "Regional variation in anaesthesia practice across Japan is substantial "
    "and predominantly structural rather than an artefact of insurance "
    "auditing. University hospital proximity is the dominant determinant. The "
    "marked variation in epidural anaesthesia (L002) and continuous epidural "
    "infusion (L003), the publicly reported proxies for regional anaesthesia "
    "use, indicates a modifiable inequity in "
    "access to a potentially beneficial technique under a system designed for "
    "universal coverage.")

# Keywords
add_blank()
p = doc.add_paragraph()
add_run_with_refs(p, "Keywords: ", bold=True)
add_run_with_refs(
    p,
    "anaesthesia; small-area variation; multilevel analysis; "
    "health services research; quality of health care; Japan")

doc.add_page_break()

# ============================================================
# INTRODUCTION
# ============================================================
add_heading("Introduction", level=1, space_before=0)

add_para(
    "Japan's universal health insurance system, established in 1961, covers "
    "the entire population under a nationally uniform fee schedule that "
    "specifies reimbursement for every medical procedure.{1} The system "
    "consistently ranks among the highest performing globally for effective "
    "service coverage{2} and is delivered through a hierarchical structure of "
    "47 prefectural (tertiary) and 335 secondary medical areas, each designed "
    "to be self-sufficient for inpatient care.{3} Cross-boundary patient "
    "movement is limited: only 6.1% of inpatients receive care outside their "
    "home prefecture.{4} The system was intended to ensure equitable access "
    "regardless of geography.")

add_para(
    "Yet substantial regional variation in surgical procedures, prescribing "
    "and diagnostic testing has been documented through the National Database "
    "of Health Insurance Claims.{5} A distinctive feature of the Japanese "
    "system is mandatory prefectural insurance auditing (shinsa): all claims "
    "are reviewed before reimbursement and audit rates vary across "
    "prefectures from 0.07% to 0.28%.{6} This raises a fundamental question "
    "for the quality and safety community: does the variation observed in "
    "claims data reflect genuine differences in clinical practice, or is it "
    "an artefact of differential auditing? Distinguishing between these two "
    "explanations has different implications for quality improvement, payment "
    "policy and patient equity.")

add_para(
    "Anaesthesia practice is well suited to this question. The choice between "
    "general anaesthesia alone and techniques that add regional anaesthesia is "
    "clinically consequential: a body of evidence suggests that regional "
    "anaesthesia, particularly epidural analgesia combined with general "
    "anaesthesia, may improve recurrence-free and overall survival in some "
    "cancer surgeries,{7,8} although results from randomised trials are "
    "mixed.{9} The 2026 revision of the Japanese fee schedule renamed the "
    "general anaesthesia code to clarify longstanding ambiguity in airway "
    "definitions, an ambiguity that could plausibly have driven differential "
    "audit decisions.{10} Understanding whether observed regional variation "
    "is structural or administrative is therefore both timely and central to "
    "evaluating the equity of care under universal coverage.")

add_para(
    "We used publicly available standardised claim ratios at the secondary "
    "medical area level to address the following research questions: "
    "(i) How large is the regional variation in anaesthesia practice across "
    "335 areas? (ii) Is this variation explained by differential prefectural "
    "auditing? (iii) What proportion of variance lies between prefectures and "
    "between areas within prefectures? (iv) Which structural factors, "
    "including university hospital proximity, are associated with the "
    "observed variation?")

# ============================================================
# METHODS
# ============================================================
add_heading("Methods", level=1)

add_subheading("Study design and ethical considerations")
add_para(
    "This was a cross-sectional ecological study using publicly available "
    "aggregate data only. No individual-level data were accessed. Ethics "
    "committee approval was not required under the Ethical Guidelines for "
    "Medical and Biological Research Involving Human Subjects (Japan, 2021 "
    "revision). The study was reported in accordance with the Strengthening "
    "the Reporting of Observational Studies in Epidemiology (STROBE) "
    "checklist for cross-sectional studies,{30} which is provided as "
    "supplementary material.")

add_subheading("Geographic units")
add_para(
    "Japan's healthcare delivery system is organised into a hierarchical "
    "structure of medical areas defined under Article 30-4 of the Medical "
    "Care Act.{3} The 47 prefectures serve as tertiary medical areas; within "
    "each prefecture, secondary medical areas (n = 335 nationally) are "
    "designated as the basic units for inpatient care planning. Crucially, "
    "secondary medical areas are completely nested within prefectures and do "
    "not cross prefectural boundaries. This nesting differs from the Hospital "
    "Service Areas and Hospital Referral Regions used in the United States, "
    "which can cross state boundaries,{11} and makes within-prefecture "
    "comparisons unconfounded by spillover under a common audit policy.")

add_subheading("Data sources and standardised claim ratios")
add_para(
    "We used three data sources. First, standardised claim ratios for fiscal "
    "year 2022 published by the Cabinet Office under the \"Regional Variation "
    "Visualisation\" initiative.{12} Standardised claim ratios are computed "
    "by indirect standardisation: expected claim frequencies are calculated "
    "by applying national age- and sex-specific claim rates to the local "
    "population structure, and the ratio is defined as 100 \u00d7 (observed "
    "claims / expected claims). Ratios are calculated on a residence basis, "
    "so claims are attributed to the beneficiary's registered address rather "
    "than to the providing facility. Patient travel for treatment therefore "
    "does not inflate the ratio of the receiving area. Areas with very few "
    "claims are masked by the data provider to protect privacy and appear as "
    "missing values. Data were available at both prefectural and secondary "
    "medical area levels. Second, we obtained physician statistics from the "
    "2022 Survey of Physicians, Dentists and Pharmacists.{13} Third, "
    "geographic boundary data were obtained from the National Land Numerical "
    "Information dataset.{14}")

add_subheading("Anaesthesia procedure codes and university hospital mapping")
add_para(
    "Six procedure codes from the Japanese fee schedule were analysed "
    "(Table 1): L008 (closed-circuit general anaesthesia); L002 (epidural "
    "anaesthesia as main technique); L003 (continuous epidural infusion, "
    "largely billed as adjunct to general anaesthesia); L004 (spinal anaesthesia); "
    "L009 (anaesthesia management fee I, a proxy for specialist staffing); "
    "and L100 (inpatient nerve block). We mapped 81 university hospitals (44 "
    "national, 8 public and 29 private) to their respective secondary medical "
    "areas based on municipal address. The 81 university hospitals are "
    "distributed across 64 of the 335 secondary medical areas (19.1%) and "
    "across all 47 prefectures, which enables within-prefecture comparisons "
    "between university and non-university areas under a common audit policy.")

add_subheading("Statistical analysis")
add_para(
    "Three complementary approaches were used. First, we computed descriptive "
    "statistics including coefficients of variation. Between-group "
    "comparisons used Welch's t-test, Cohen's d for effect size{15} and the "
    "Mann-Whitney U test for non-parametric confirmation. Second, we fitted "
    "multilevel linear mixed models with secondary medical areas (level 1, "
    "n = 335) nested within prefectures (level 2, n = 47).{16} A null random-"
    "intercept model was fitted first to estimate the intraclass correlation "
    "coefficient. Subsequent models added fixed effects for university "
    "hospital presence, anaesthesiologist density (standardised) and total "
    "physician density (standardised). Models were estimated by restricted "
    "maximum likelihood using the Python statsmodels MixedLM "
    "implementation.{17} Marginal R\u00b2 was calculated as the proportional "
    "reduction in total variance from the null model. To address potential "
    "instability of ratios in low-volume areas, we applied empirical Bayes "
    "shrinkage estimation{18} and compared all main findings using both raw "
    "and shrunken ratios.")

add_para(
    "Third, three pre-specified sensitivity analyses tested the null "
    "hypothesis that insurance auditing explains observed variation. "
    "(a) Within-prefecture variance decomposition: under uniform prefectural "
    "auditing, within-prefecture variance should be small. We decomposed "
    "total variance using one-way analysis of variance and further partitioned "
    "the within-prefecture component into university hospital effect and "
    "residual using hierarchical sum-of-squares decomposition. (b) Cross-code "
    "correlation: audit-driven reclassification between general and spinal "
    "anaesthesia codes would produce negative correlations between them. "
    "(c) Quantitative audit-impact estimation: we calculated the maximum ratio "
    "shift attributable to the observed range of prefectural audit rates. "
    "Combined ratios (general plus spinal; general plus epidural) were also "
    "examined, since combining substitutable codes should largely absorb any "
    "audit-driven reclassification. Patients were not involved in the design, "
    "conduct or reporting of this study.")

# ============================================================
# RESULTS
# ============================================================
add_heading("Results", level=1)

add_subheading("Study population and variation in anaesthesia practice")
add_para(
    "The 335 secondary medical areas were distributed across 47 prefectures "
    "(median 7 areas per prefecture, range 3 to 21); 64 areas (19.1%) "
    "contained at least one university hospital. Standardised claim ratios "
    "for general anaesthesia were available for 334 of 335 areas; ratios for "
    "epidural anaesthesia were available for 307 areas, with 28 areas masked "
    "owing to low volume. Substantial variation was observed across all "
    "codes (Table 2). General anaesthesia ratios ranged from 2.3 to 435.7 "
    "(coefficient of variation 53.6%), a 189-fold difference. Epidural "
    "anaesthesia showed the greatest relative variation (coefficient of "
    "variation 87.0%); continuous epidural infusion 64.9% and spinal "
    "anaesthesia 56.8%. The geographic distribution of ratios is shown in "
    "Figure 1.")

add_subheading("Multilevel model and university hospital effect")
add_para(
    "The null multilevel model showed that only 5.8% of general anaesthesia "
    "variance was attributable to the prefecture level (intraclass correlation "
    "coefficient 0.058), indicating that 94.2% occurred within prefectures "
    "where audit policy is uniform. Spinal anaesthesia showed stronger "
    "prefecture-level clustering (intraclass correlation coefficient 0.343), "
    "consistent with a greater role of prefectural factors for that code. "
    "Adding university hospital presence as a fixed effect produced the "
    "largest improvement for general anaesthesia (marginal R\u00b2 0.358; "
    "\u03b2 = +64.1, 95% confidence interval 54.8 to 73.4) and was "
    "statistically significant for every code (all P < 0.001) (Table 3). "
    "University hospital areas had higher general anaesthesia ratios than "
    "non-university areas in 47 of 47 prefectures, with a mean within-"
    "prefecture difference of +63.3 points (paired t = 13.28, P < 0.001) "
    "(Figure 2A). Cohen's d for the university hospital effect on general "
    "anaesthesia was 1.88; on continuous epidural infusion 1.30; on epidural "
    "anaesthesia 0.59; and on spinal anaesthesia 0.40. A combined measure "
    "capturing the general-anaesthesia plus continuous-epidural workflow "
    "(mean of L008 and L003 SCR; 331 areas with data for both codes) "
    "showed a coefficient of variation of 55.5% and preserved a large "
    "university hospital effect (Cohen's d 1.46; university mean 128.3 "
    "versus non-university 70.7) (Figure 2B).")

add_subheading("Sensitivity analyses against the audit hypothesis")
add_para(
    "All three pre-specified sensitivity analyses converged in rejecting the "
    "audit hypothesis. First, within-prefecture variance decomposition "
    "attributed 14.5% of general anaesthesia variance to between-prefecture "
    "differences, 40.5% to the university hospital effect within prefecture, "
    "and 45.0% to residual within-group variation; the university hospital "
    "effect alone explained 47.4% of all within-prefecture variance. Second, "
    "general and spinal anaesthesia were positively, not negatively, "
    "correlated (r = +0.235, P < 0.001), as were general and epidural "
    "anaesthesia (r = +0.319, P < 0.001), inconsistent with audit-driven "
    "reclassification and consistent with a common supply factor. Third, the "
    "maximum prefectural audit-rate difference of 0.21 percentage points "
    "could produce a ratio shift of approximately 0.2 points, which is less "
    "than 0.3% of the observed interquartile range for general anaesthesia "
    "and less than 0.1% for epidural anaesthesia. Combining general and "
    "spinal ratios reduced the coefficient of variation only modestly from "
    "53.6% to 43.4% (a 19% reduction), and the university hospital effect "
    "remained large on the combined measure (Cohen's d 1.26).")

add_subheading("Robustness: empirical Bayes shrinkage and outliers")
add_para(
    "Empirical Bayes shrinkage toward prefecture means was implemented as "
    "a stress test of the area-level ratios under the explicit sampling "
    "model y_i ~ N(θ_i, σ²); θ_i ~ N(μ_pref, τ²_pref), with σ² taken as "
    "the residual within-prefecture variance from a null random-intercept "
    "mixed-effects model fitted separately for each code. This formulation "
    "is conservative because no information is borrowed from the observed "
    "prefecture-mean differences when shrinking individual areas. Under "
    "this model, shrinkage substantially compressed the area-level ratios "
    "toward their prefecture means. For general anaesthesia, Cohen's d for "
    "the university-hospital effect fell from 1.88 (raw) to 0.63 "
    "(shrunken), an attenuation of 66.5%; for epidural anaesthesia (L002) "
    "from 0.59 to 0.23 (60.6%); and for spinal anaesthesia (L004) from "
    "0.40 to 0.04 (89.9%). The general anaesthesia effect therefore "
    "remained of moderate magnitude despite heavy shrinkage, whereas the "
    "L002 and L004 effects were materially attenuated, indicating that "
    "those smaller raw effects are more vulnerable to sampling instability. "
    "These shrinkage estimates are complementary to, and consistent in "
    "direction with, the mixed-effects model coefficients reported in "
    "Table 3, which themselves incorporate prefecture-level shrinkage. Two "
    "areas "
    "were identified as outliers (more than three standard deviations from "
    "the mean general anaesthesia ratio): Awa in Chiba (ratio 216.7) and "
    "central Tokyo (ratio 435.7, containing five university hospitals). "
    "Excluding outliers did not alter the direction or statistical "
    "significance of any finding.")

# ============================================================
# DISCUSSION
# ============================================================
add_heading("Discussion", level=1)

add_subheading("Statement of principal findings")
add_para(
    "Regional variation in anaesthesia practice across Japan's 335 secondary "
    "medical areas is substantial, with coefficients of variation ranging from "
    "53% for general anaesthesia to 87% for epidural anaesthesia. Multilevel "
    "modelling showed that only 5.8% of general anaesthesia variance is "
    "attributable to prefectures, while 94.2% occurs within prefectures where "
    "audit policy is uniform. University hospital presence explained 35.8% of "
    "total variance and 40.5% of within-prefecture variance, and the effect "
    "was present in all 47 prefectures despite differing audit practices. "
    "Three independent sensitivity analyses converged in rejecting "
    "differential auditing as a plausible explanation. The observed variation "
    "is therefore predominantly structural, driven by institutional capacity "
    "and clinical-organisational factors rather than by administrative coding.")

add_subheading("Strengths and limitations")
add_para(
    "Strengths include the use of age- and sex-standardised ratios covering "
    "the entire national population; analysis at a fine geographic scale of "
    "335 areas; multilevel modelling that properly accounts for the nested "
    "data structure; empirical Bayes shrinkage to address low-volume "
    "instability; multiple pre-specified sensitivity analyses; within-"
    "prefecture comparisons that hold audit policy constant; and the "
    "residence-based ratio definition that mitigates patient travel effects. "
    "Limitations include the ecological fallacy inherent in area-level "
    "analysis;{19} the cross-sectional design, which precludes causal "
    "inference; the post-audit nature of the data, which reflects reimbursed "
    "rather than intended practice; the partial overlap between university "
    "hospital presence and urban concentration; the absence of code-specific "
    "audit-rate data; and the inability to quantify defensive undercoding "
    "from claims alone. A specific constraint of the publicly available "
    "dataset is that the combined anaesthesia add-on billed under L008 "
    "(the explicit code for general-plus-epidural technique) is aggregated "
    "with other anaesthesia add-ons in the regional variation data release "
    "and cannot be extracted separately at the secondary-medical-area "
    "level; patient-level National Database records would be required to "
    "measure the true rate of combined general-epidural anaesthesia. We "
    "therefore used L002 (epidural as main technique) and L003 (continuous "
    "epidural infusion) as the publicly reported proxies for regional "
    "technique use. Multilevel models included only university hospital "
    "presence as a structural fixed effect; future work should add "
    "additional area-level covariates such as urbanisation, bed density and "
    "case-mix indicators.")

add_subheading("Interpretation within the context of the wider literature")
add_para(
    "Our findings are consistent with the broader medical practice variation "
    "literature. The Dartmouth Atlas project documented extensive regional "
    "variation in surgical rates in the United States, driven primarily by "
    "physician supply and practice style.{21} Comparable patterns have been "
    "described for the National Health Service{22} and in Germany{23} and "
    "Australia.{24} Our study extends this evidence by exploiting Japan's "
    "uniquely uniform fee schedule combined with prefecture-specific auditing "
    "to disentangle administrative and clinical sources of variation, and by "
    "providing the first multilevel small-area analysis of anaesthesia "
    "technique under universal coverage in East Asia. The university hospital "
    "effect we report (Cohen's d = 1.88) is unusually large compared with "
    "effect sizes typically reported in this literature, plausibly reflecting "
    "the influence of the ikyoku (university medical office) system on "
    "clinical practice in affiliated hospitals.{25} The intraclass "
    "correlation coefficient of 0.058 for general anaesthesia is also "
    "noticeably lower than coefficients reported for surgical procedure rates "
    "in the United States,{21} indicating that within-prefecture institutional "
    "factors dominate over prefectural-level factors for this code.")

add_para(
    "Our findings and framework extend beyond Japan in four respects. First, "
    "the within-prefecture variance decomposition we describe is directly "
    "transportable to other universal-coverage systems that combine a "
    "centrally uniform fee schedule with regionally devolved claims auditing, "
    "including Taiwan's National Health Insurance,{20} South Korea's National "
    "Health Insurance Service, and the tariff-based systems of Germany, "
    "France and the English National Health Service.{22,23} Second, the "
    "dominance of tertiary-teaching-hospital concentration as a structural "
    "determinant of technique choice echoes analogous patterns for "
    "teaching-hospital anaesthesia mix reported in the US and United "
    "Kingdom,{21,22} and for other tertiary-intensive specialties across "
    "OECD countries;{26} the consistency of the university effect across "
    "47 of 47 prefectures in our data suggests that it is a robust feature "
    "of high-income health systems rather than a Japanese idiosyncrasy. "
    "Third, the variation we identify in epidural anaesthesia (L002, "
    "coefficient of variation 87.0%) and continuous epidural infusion "
    "(L003, 64.9%) speaks to an international equity agenda: where emerging "
    "evidence links regional anaesthesia to oncological and functional "
    "outcomes,{7-9} monitoring the distribution of anaesthesia techniques "
    "could serve as a system-level quality indicator across any universal "
    "coverage setting committed to equitable access to technology-rich "
    "care.{27,29} Fourth, extrapolation to low- and middle-income countries "
    "is more cautious: fee schedules and audit arrangements are typically "
    "heterogeneous and claims data often incomplete, which limits direct "
    "replication of the decomposition. Nevertheless, the general principle "
    "— that apparent small-area variation should be interrogated for "
    "administrative versus structural origins before policy action — is "
    "transferable, particularly as low- and middle-income countries "
    "progressively standardise their fee schedules under universal health "
    "coverage reforms.{2}")

add_subheading("Implications for policy, practice and research")
add_para(
    "The marked variation in continuous epidural infusion (L003; coefficient "
    "of variation 64.9%) and in epidural anaesthesia (L002; 87.0%) is "
    "clinically important because general anaesthesia supplemented with "
    "regional techniques has been linked, albeit inconsistently, with better "
    "oncological outcomes.{7-9} A 1.73-fold difference in continuous epidural "
    "infusion ratios between university and non-university areas implies that "
    "patients' access to a potentially beneficial technique depends "
    "substantially on where they live, which challenges the equity premise of "
    "Japan's universal coverage. Quality improvement responses could include "
    "targeted education, specialist outreach from university centres, and "
    "monitoring anaesthesia technique distribution as a system-level quality "
    "indicator,{27,29} aligned with international standards for safe practice "
    "of anaesthesia.{28} The within-prefecture variance decomposition and "
    "multilevel sensitivity framework developed here are generalisable to "
    "other procedures under universal insurance and provide a transferable "
    "tool for distinguishing administrative from structural sources of "
    "small-area variation.")

# ============================================================
# CONCLUSIONS
# ============================================================
add_heading("Conclusions", level=1)
add_para(
    "Regional variation in anaesthesia practice in Japan is large and "
    "predominantly structural rather than an artefact of insurance auditing. "
    "University hospital presence is the dominant determinant, explaining "
    "more than one third of total variance in general anaesthesia. The "
    "variation in epidural anaesthesia (L002) and continuous epidural "
    "infusion (L003) is of particular concern given emerging evidence of "
    "oncological benefit from regional techniques. Quality "
    "improvement and equity policy should address the supply and "
    "organisational determinants of anaesthesia practice rather than treat "
    "the variation as an unavoidable feature of claims processing.")

# ============================================================
# REFERENCES (Vancouver, numbered in order of appearance)
# ============================================================
add_heading("References", level=1)

REFERENCES = [
    "Ikegami N, Yoo BK, Hashimoto H, et al. Japanese universal health coverage: evolution, achievements, and challenges. Lancet 2011;378:1106-15.",
    "GBD 2019 Universal Health Coverage Collaborators. Measuring universal health coverage based on an index of effective coverage of health services in 204 countries and territories, 1990-2019. Lancet 2020;396:1250-84.",
    "Matsuda S. Health policy in Japan: current situation and future challenges. JMA J 2019;2:1-10.",
    "Ministry of Health, Labour and Welfare. Patient Survey 2008. Tokyo: MHLW, 2009.",
    "Hashimoto H, Ikegami N, Shibuya K, et al. Cost containment and quality of care in Japan: is there a trade-off? Lancet 2011;378:1174-82.",
    "Cabinet Office. Indicators of regional variation in healthcare and long-term care: insurance auditing rates by prefecture. Tokyo: Cabinet Office, 2023.",
    "Sessler DI, Pei L, Huang Y, et al. Recurrence of breast cancer after regional or general anaesthesia: a randomised controlled trial. Lancet 2019;394:1807-15.",
    "Chen WK, Miao CH. The effect of anesthetic technique on survival in human cancers: a meta-analysis of retrospective and prospective studies. PLoS One 2013;8:e56540.",
    "Weng M, Chen W, Hou W, et al. The effect of neuraxial anaesthesia on cancer recurrence and survival after cancer surgery: an updated meta-analysis. Oncotarget 2016;7:15262-73.",
    "Ministry of Health, Labour and Welfare. 2026 revision of the medical fee schedule: notice on anaesthesia code revisions. Tokyo: MHLW, 2025.",
    "Wennberg JE, Cooper MM, eds. The Dartmouth Atlas of Health Care. Chicago: American Hospital Publishing, 1999.",
    "Cabinet Office. Regional variation visualisation (chiikisa no mieruka). https://www5.cao.go.jp/keizai-shimon/kaigi/special/reform/mieruka/ (accessed 10 Mar 2026).",
    "Ministry of Health, Labour and Welfare. Survey of Physicians, Dentists and Pharmacists 2022 (e-Stat). https://www.e-stat.go.jp (accessed 12 Mar 2026).",
    "Ministry of Land, Infrastructure, Transport and Tourism. National Land Numerical Information download service. https://nlftp.mlit.go.jp (accessed 10 Mar 2026).",
    "Cohen J. Statistical Power Analysis for the Behavioral Sciences. 2nd ed. Hillsdale, NJ: Lawrence Erlbaum, 1988.",
    "Rabe-Hesketh S, Skrondal A. Multilevel and Longitudinal Modeling Using Stata. 3rd ed. College Station, TX: Stata Press, 2012.",
    "Seabold S, Perktold J. Statsmodels: econometric and statistical modeling with Python. In: Proceedings of the 9th Python in Science Conference. 2010:92-6.",
    "Efron B, Morris C. Stein's estimation rule and its competitors: an empirical Bayes approach. J Am Stat Assoc 1973;68:117-30.",
    "Greenland S. Ecologic versus individual-level sources of bias in ecologic estimates of contextual health effects. Int J Epidemiol 2001;30:1343-50.",
    "Cheng T-M. Taiwan's new national health insurance program: genesis and experience so far. Health Aff (Millwood) 2003;22:61-76.",
    "Wennberg JE, Fisher ES, Skinner JS. Geography and the debate over Medicare reform. Health Aff (Millwood) 2002;Suppl Web Exclusives:W96-114.",
    "NHS RightCare. The NHS Atlas of Variation in Healthcare. London: Public Health England, 2015.",
    "Wengler A, Nimptsch U, Mansky T. Hip and knee replacement surgery (arthroplasty) and geographic variations in Germany. Dtsch Arztebl Int 2014;111:407-16.",
    "Australian Commission on Safety and Quality in Health Care. Australian Atlas of Healthcare Variation. Sydney: ACSQHC, 2015.",
    "Otsuka T. The ikyoku system of university orthopaedic surgery departments: an in-hospital organisational system unique to Japan. J Orthop Sci 2012;17:513-14.",
    "OECD. Geographic Variations in Health Care: What Do We Know and What Can Be Done to Improve Health System Performance? Paris: OECD Publishing, 2014.",
    "Donabedian A. The quality of care: how can it be assessed? JAMA 1988;260:1743-8.",
    "Merry AF, Cooper JB, Soyannwo O, et al. International standards for a safe practice of anesthesia 2010. Can J Anesth 2010;57:1027-34.",
    "Mainz J. Defining and classifying clinical indicators for quality improvement. Int J Qual Health Care 2003;15:523-30.",
    "von Elm E, Altman DG, Egger M, et al. The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement: guidelines for reporting observational studies. Lancet 2007;370:1453-7.",
]
assert len(REFERENCES) == 30, f"Reference count is {len(REFERENCES)} (max 30)"

for i, ref in enumerate(REFERENCES, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(f"{i}. ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# ============================================================
# TABLES (after references, per IJQHC structure)
# ============================================================
doc.add_page_break()
add_heading("Tables", level=1, space_before=0)

# Table 1: Anaesthesia procedure codes
add_para("Table 1. Anaesthesia procedure codes analysed.", bold=True)
table1 = doc.add_table(rows=1, cols=3)
table1.style = 'Table Grid'
table1.alignment = 1
hdr = table1.rows[0].cells
hdr[0].text = "Code"
hdr[1].text = "Procedure"
hdr[2].text = "Notes"
for cell in hdr:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

table1_rows = [
    ("L008", "Closed-circuit general anaesthesia",
     "Primary indicator of overall general anaesthesia volume"),
    ("L002", "Epidural anaesthesia",
     "Regional technique claimed as main anaesthetic"),
    ("L003", "Continuous epidural infusion",
     "Regional technique largely billed as adjunct to general anaesthesia"),
    ("L004", "Spinal anaesthesia", "Alternative regional technique"),
    ("L009", "Anaesthesia management fee I",
     "Proxy for specialist anaesthesiologist staffing"),
    ("L100", "Nerve block, inpatient", "Indicator of pain-clinic activity"),
]
for code, proc, note in table1_rows:
    row = table1.add_row().cells
    row[0].text = code
    row[1].text = proc
    row[2].text = note
    for cell in row:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

add_blank()
add_para("Source: Japanese fee schedule (shinryo houshu tensuhyo), fiscal year "
         "2022. L codes correspond to anaesthesia-related procedures. "
         "Standardised claim ratios were computed for each code by indirect "
         "age- and sex-standardisation against the national average (= 100).",
         italic=True)

doc.add_page_break()

# Table 2: Distribution of SCRs
add_para("Table 2. Distribution of standardised claim ratios across 335 "
         "secondary medical areas (national average = 100), fiscal year 2022.",
         bold=True)
table2 = doc.add_table(rows=1, cols=8)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
for i, h in enumerate(["Code", "n", "Mean (SD)", "Median (IQR)", "Min", "Max",
                        "CV (%)", "Range ratio"]):
    hdr2[i].text = h
    for p in hdr2[i].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'

table2_rows = [
    ("L008", "334", "85.7 (45.9)", "76.5 (53.0-110.5)", "2.3", "435.7", "53.6", "189"),
    ("L002", "307", "61.4 (53.4)", "47.7 (24.5-83.4)", "0", "519.1", "87.0", "-"),
    ("L003", "307", "61.4 (39.8)", "53.5 (32.0-82.0)", "0", "224.7", "64.9", "-"),
    ("L004", "320", "82.5 (46.8)", "76.5 (49.5-108.5)", "0", "318.0", "56.8", "-"),
    ("L009", "330", "78.3 (43.0)", "71.0 (47.5-101.5)", "0", "230.0", "55.0", "-"),
    ("L100", "324", "60.5 (35.7)", "53.5 (35.5-78.0)", "0", "247.0", "59.0", "-"),
]
for r in table2_rows:
    row = table2.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val
        for p in row[i].paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

add_blank()
add_para("CV, coefficient of variation; IQR, interquartile range; SD, standard "
         "deviation. Range ratio (max / min) is reported only where the minimum "
         "is non-zero. Areas with low claim volume are masked by the data "
         "provider and appear as missing values; the n column reflects the "
         "number of secondary medical areas with non-missing values for each "
         "code.", italic=True)

doc.add_page_break()

# Table 3: Multilevel model results
add_para("Table 3. Multilevel linear mixed model results: standardised claim "
         "ratio as outcome, prefecture as random intercept.", bold=True)
table3 = doc.add_table(rows=1, cols=5)
table3.style = 'Table Grid'
hdr3 = table3.rows[0].cells
for i, h in enumerate(["Code", "Null model ICC",
                        "\u03b2 university (95% CI)", "P value",
                        "Marginal R\u00b2"]):
    hdr3[i].text = h
    for p in hdr3[i].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'

table3_rows = [
    ("L008", "0.058", "+64.1 (54.8 to 73.4)", "<0.001", "0.358"),
    ("L002", "0.082", "+49.0 (38.4 to 59.6)", "<0.001", "0.205"),
    ("L003", "0.071", "+55.8 (45.2 to 66.4)", "<0.001", "0.245"),
    ("L004", "0.343", "+23.0 (12.4 to 33.6)", "<0.001", "0.045"),
    ("L009", "0.205", "+47.6 (36.4 to 58.8)", "<0.001", "0.180"),
    ("L100", "0.299", "+18.4 (8.6 to 28.2)", "<0.001", "0.038"),
]
for r in table3_rows:
    row = table3.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val
        for p in row[i].paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)

add_blank()
add_para("\u03b2 university, fixed effect coefficient for university hospital "
         "presence (binary 0/1); CI, confidence interval; ICC, intraclass "
         "correlation coefficient (proportion of variance attributable to "
         "prefecture level in the null random-intercept model); Marginal "
         "R\u00b2, proportional reduction in total (prefecture + residual) "
         "variance from the null model after adding the university hospital "
         "fixed effect. Models were estimated by restricted maximum "
         "likelihood.", italic=True)

doc.add_page_break()

# ============================================================
# FIGURES (after tables, with legends; images included for review)
# ============================================================
add_heading("Figure legends", level=1, space_before=0)

add_para("Figure 1. Geographic distribution of anaesthesia standardised claim "
         "ratios across 335 secondary medical areas of Japan, fiscal year "
         "2022. (A) General anaesthesia (L008). (B) Spinal anaesthesia "
         "(L004). (C) Epidural anaesthesia as main anaesthetic (L002). "
         "(D) Continuous epidural infusion (L003). Choropleth maps shaded "
         "by quintile of the standardised claim ratio (national average = "
         "100). Red circles mark secondary medical areas containing at "
         "least one university hospital. Areas masked by the data provider "
         "owing to low volume are shown in grey.", bold=False)
add_blank()
add_para("Figure 2. University hospital presence and the combined "
         "general-anaesthesia plus continuous-epidural measure. "
         "(A) Distribution of secondary medical areas containing at least "
         "one university hospital (n = 64 of 335; red). (B) Choropleth map "
         "of the combined general-anaesthesia plus continuous-epidural "
         "standardised claim ratio (mean of L008 and L003 SCR; 331 areas "
         "with data for both codes), shaded by quintile. Red circles mark "
         "secondary medical areas containing at least one university "
         "hospital. Areas masked by the data provider for either code "
         "are shown in grey.", bold=False)

# Embed figure images at end (after legends) for reviewer convenience
for fig_path, caption in [
    (os.path.join(FIG_DIR, 'rapm_fig1_en.png'), 'Figure 1.'),
    (os.path.join(FIG_DIR, 'rapm_fig2_en.png'), 'Figure 2.'),
]:
    if os.path.exists(fig_path):
        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        doc.add_picture(fig_path, width=Inches(6.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================
# Save
# ============================================================
out = os.path.join(OUTPUT_DIR, 'regional_anaesthesia_IJQHC_EN.docx')
doc.save(out)
print(f"Saved: {out}")

# Word count check (excluding title page, abstract, references, tables, figures)
import collections
text_paras = []
in_body = False
in_refs = False
for p in doc.paragraphs:
    txt = p.text.strip()
    if not txt:
        continue
    if txt.upper() == 'INTRODUCTION':
        in_body = True
        continue
    if txt.upper() == 'REFERENCES':
        in_refs = True
        continue
    if txt.upper() == 'TABLES' or txt.upper().startswith('FIGURE LEGEND'):
        break
    if in_body and not in_refs:
        text_paras.append(txt)
body = '\n'.join(text_paras)
words = len(re.findall(r'\b\w+\b', body))
print(f"Body word count (intro through conclusion): {words}")
