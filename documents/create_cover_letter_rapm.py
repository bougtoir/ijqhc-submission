#!/usr/bin/env python3
"""
Create RAPM cover letter (English) for the anesthesia regional variation manuscript.
Emphasizes BMJ transfer recommendation, journal fit, and timeliness.
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Page setup
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

def add_para(text, bold=False, italic=False, align=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    return p

# Date
add_para('[Date]', align=WD_ALIGN_PARAGRAPH.LEFT)
doc.add_paragraph()

# Addressee
add_para('Editor-in-Chief')
add_para('Regional Anesthesia & Pain Medicine')
add_para('BMJ Publishing Group')
doc.add_paragraph()

# Subject line
add_para(
    'Re: Submission of manuscript transferred from The BMJ (BMJ-2026-090550)',
    bold=True
)
doc.add_paragraph()

# Salutation
add_para('Dear Editor,')
doc.add_paragraph()

# Body
add_para(
    'We are pleased to submit our manuscript entitled "Regional variation in anesthesia practice '
    'across 335 secondary medical areas in Japan: a cross-sectional analysis of national claims '
    'data" for consideration for publication in Regional Anesthesia & Pain Medicine. This '
    'manuscript was originally submitted to The BMJ (manuscript ID: BMJ-2026-090550) and has been '
    'transferred to your journal at the recommendation of the BMJ editorial team, who identified '
    'RAPM as the most appropriate venue for this work.'
)

add_para(
    'Using publicly available standardized claim ratios from Japan\u2019s National Database of Health '
    'Insurance Claims, we quantify regional variation in anesthesia practice across all 335 '
    'secondary medical areas\u2014the finest geographic resolution available for the entire Japanese '
    'population. Our study addresses a question of direct relevance to the RAPM readership: does '
    'geographic variation in anesthesia technique selection reflect genuine differences in clinical '
    'practice, or is it an artifact of administrative mechanisms such as insurance auditing?'
)

add_para(
    'Three independent sensitivity analyses converge on the conclusion that regional variation is '
    'structural rather than an audit artifact. The dominant determinant is university hospital '
    'proximity, which alone explains 38.5% of variance in general anesthesia volume (Cohen\u2019s '
    'd\u2009=\u20091.78)\u2014an effect present in all 47 prefectures. Of particular relevance to your '
    'readership, combined general\u2013epidural anesthesia technique shows the greatest regional '
    'inequality (CV\u2009=\u200964.9%), with a 1.73-fold difference between university and non-university '
    'hospital areas. Given emerging evidence linking regional anesthesia techniques to improved '
    'oncologic outcomes, this variation may represent a modifiable source of health inequality.'
)

add_para(
    'We believe this work is well suited to RAPM for several reasons:'
)

doc.add_paragraph(
    'It provides the first systematic, population-level analysis of anesthesia technique variation '
    'at a fine geographic scale in an East Asian universal insurance system, extending the practice '
    'variation literature beyond the Dartmouth Atlas tradition.',
    style='List Bullet'
)

doc.add_paragraph(
    'The sensitivity analysis framework we develop\u2014within-prefecture variance decomposition, '
    'cross-code correlation, and quantitative audit impact estimation\u2014is novel and applicable to '
    'investigating practice variation in any procedure under universal insurance systems.',
    style='List Bullet'
)

doc.add_paragraph(
    'The finding that combined GA\u2013epidural technique shows the greatest geographic inequality '
    'is directly relevant to the ongoing debate in RAPM about regional anesthesia and cancer '
    'outcomes.',
    style='List Bullet'
)

doc.add_paragraph(
    'The study is timely: the 2026 Japanese fee schedule revision renamed the general anesthesia '
    'code to reflect supraglottic airway device use, and our findings provide baseline data for '
    'evaluating the impact of this policy change.',
    style='List Bullet'
)

doc.add_paragraph()

add_para(
    'The manuscript reports an ecological cross-sectional study using exclusively publicly available '
    'aggregate data. No individual patient data were accessed, and ethics committee approval was not '
    'required under Japanese ethical guidelines. The study adheres to STROBE reporting guidelines, '
    'and a completed checklist is included.'
)

add_para(
    'This manuscript has not been published previously and is not under consideration elsewhere. '
    'All authors have approved the submitted version and agree to be accountable for all aspects '
    'of the work.'
)

add_para(
    'We appreciate the opportunity to have this work considered by Regional Anesthesia & Pain '
    'Medicine and look forward to your editorial decision.'
)

doc.add_paragraph()
add_para('Sincerely,')
doc.add_paragraph()
add_para('[Corresponding author name]')
add_para('[Affiliation]')
add_para('[Email address]')
add_para('[ORCID]')

output_path = '/home/ubuntu/cover_letter_RAPM_EN.docx'
doc.save(output_path)
print(f"Saved: {output_path}")
