#!/usr/bin/env python3
"""
Create RAPM-format docx for regional variation in anesthesia practice in Japan.
- American English (RAPM is ASRA publication)
- STROBE compliant
- Structured abstract: Background/Methods/Results/Conclusion
- 3 key message boxes: What is already known / What this study adds / How this study might affect
- Tables + Figures combined <= 6 (2 tables + 4 multi-panel figures)
- Vancouver references
- Figures as multi-panel composites to meet RAPM limit
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Style setup ──
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# Helper functions
def add_heading_rapm(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    return p

def add_table_row(table, cells, bold=False):
    row = table.add_row()
    for i, cell_text in enumerate(cells):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(cell_text))
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.bold = bold
    return row

def add_figure(path, caption, width=Inches(6)):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.italic = True
    else:
        add_para(f'[Figure: {path} not found]', italic=True)


# ═══════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════

title_text = (
    "Regional variation in anesthesia practice across 335 secondary medical areas in Japan: "
    "a cross-sectional analysis of national claims data"
)
p = doc.add_paragraph()
run = p.add_run(title_text)
run.font.name = 'Arial'
run.font.size = Pt(16)
run.bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

add_para('[Author names]', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('[Affiliations]', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('[Corresponding author: name, address, email]', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
add_para('Word count: approximately 3,800', italic=True)
add_para('Tables: 2', italic=True)
add_para('Figures: 4 (3 multi-panel choropleth maps + 1 three-dimensional extruded map)', italic=True)
add_para('Supplementary figures: 5 (2 additional 3D maps + 3 individual 2D maps)', italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# STRUCTURED ABSTRACT (RAPM format: Background/Methods/Results/Conclusion)
# ═══════════════════════════════════════════════════════════

add_heading_rapm('ABSTRACT', level=1)

p = doc.add_paragraph()
run = p.add_run('Background ')
run.bold = True
run.font.name = 'Arial'
p.add_run(
    'Regional variation in anesthesia practice has been documented internationally, but the '
    'relative contributions of clinical preference, institutional factors, and administrative '
    'mechanisms remain unclear. Japan\u2019s universal insurance system, with its mandatory '
    'prefectural claims auditing, provides a unique setting to distinguish administrative from '
    'clinical sources of variation.'
).font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run('Methods ')
run.bold = True
run.font.name = 'Arial'
p.add_run(
    'Cross-sectional ecological study using publicly available standardized claim ratios (SCRs) '
    'for fiscal year 2022 across 335 secondary medical areas (SMAs). We analyzed six anesthesia '
    'procedure codes: general anesthesia (L008), epidural anesthesia (L002), continuous epidural '
    'infusion (L003), spinal anesthesia (L004), anesthesia management fee (L009), and nerve block '
    '(L100). Three sensitivity analyses tested whether insurance audit variation explained observed '
    'differences: within-prefecture variance decomposition, cross-code correlation analysis, and '
    'quantitative audit impact estimation.'
).font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run('Results ')
run.bold = True
run.font.name = 'Arial'
p.add_run(
    'Substantial regional variation was observed: coefficients of variation (CV) were 54.6% for '
    'general anesthesia (L008), 83.2% for epidural anesthesia (L002), and 64.9% for continuous '
    'epidural infusion (L003). Three sensitivity analyses indicated that audit variation is '
    'insufficient to explain observed differences: 85.8% of L008 variance occurred within '
    'prefectures (where audit policy is uniform); general and spinal anesthesia showed inverse '
    'correlation (r=\u22120.506); and audit rate differences could explain less than 1% of observed '
    'SCR variation. University hospital presence was the strongest single predictor, explaining '
    '38.5% of total L008 variance (Cohen\u2019s d=1.78). Combined GA\u2013epidural technique (L003) was '
    '1.73 times higher in university hospital areas.'
).font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run('Conclusion ')
run.bold = True
run.font.name = 'Arial'
p.add_run(
    'Regional variation in anesthesia practice in Japan is structural, not an artifact of insurance '
    'auditing. University hospital proximity is the dominant determinant. Given evidence linking '
    'regional anesthesia techniques to improved oncologic outcomes, this variation may represent '
    'a modifiable source of health inequality.'
).font.name = 'Arial'

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# KEY MESSAGES (RAPM 3-box format)
# ═══════════════════════════════════════════════════════════

add_heading_rapm('WHAT IS ALREADY KNOWN ON THIS TOPIC', level=2)
doc.add_paragraph(
    'Regional variation in surgical and anesthesia practice has been documented in many countries, '
    'but the relative contributions of clinical preference, institutional factors, and administrative '
    'mechanisms (such as insurance auditing) remain unclear.',
    style='List Bullet'
)
doc.add_paragraph(
    'Japan\u2019s universal insurance system processes all claims through standardized prefectural audit, '
    'creating a unique natural experiment in which administrative variation can be distinguished from '
    'clinical variation.',
    style='List Bullet'
)

add_heading_rapm('WHAT THIS STUDY ADDS', level=2)
doc.add_paragraph(
    'Three independent sensitivity analyses demonstrate that regional variation in anesthesia practice '
    'in Japan is structural rather than an artifact of insurance auditing.',
    style='List Bullet'
)
doc.add_paragraph(
    'University hospital presence explains 38.5% of variance in general anesthesia volume\u2014a '
    'remarkably large effect for a single binary predictor\u2014and the effect is consistent across '
    'all 47 prefectures.',
    style='List Bullet'
)

add_heading_rapm('HOW THIS STUDY MIGHT AFFECT RESEARCH, PRACTICE OR POLICY', level=2)
doc.add_paragraph(
    'Combined general\u2013epidural anesthesia, which may confer oncologic benefit, shows the greatest '
    'regional inequality (CV=64.9%), suggesting a modifiable target for reducing health disparities.',
    style='List Bullet'
)
doc.add_paragraph(
    'The sensitivity analysis framework developed here\u2014within-prefecture variance decomposition, '
    'cross-code correlation, and quantitative audit impact estimation\u2014is applicable to investigating '
    'practice variation in any procedure under Japan\u2019s universal insurance system.',
    style='List Bullet'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# INTRODUCTION
# ═══════════════════════════════════════════════════════════

add_heading_rapm('INTRODUCTION', level=1)

doc.add_paragraph(
    'Japan\u2019s universal health insurance system, established in 1961, covers the entire population '
    'under a standardized fee schedule (shinryo houshu tensuhyo) that specifies reimbursement for '
    'each medical procedure.\u00b9\u00b7\u00b2 This system was designed to ensure equitable access to healthcare '
    'regardless of geography. Yet studies using the National Database of Health Insurance Claims '
    '(NDB) have revealed substantial regional variation in medical practice, including surgical '
    'procedures, prescription patterns, and diagnostic testing.\u00b3\u2013\u2075'
)

doc.add_paragraph(
    'A distinctive feature of the Japanese system is mandatory insurance auditing (shinsa). All '
    'claims are reviewed by prefectural audit committees before reimbursement, and claims deemed '
    'inappropriate are reduced or rejected (satei). Audit rates vary across prefectures, ranging '
    'from 0.07% to 0.28%.\u2076 This raises a fundamental question: does observed regional variation '
    'in claims data reflect genuine differences in clinical practice, or is it an artifact of '
    'differential auditing?'
)

doc.add_paragraph(
    'This question has particular relevance for anesthesia practice. The choice of anesthesia '
    'technique\u2014general anesthesia alone versus combined general\u2013regional techniques\u2014is '
    'not merely a matter of preference. A growing body of evidence suggests that regional anesthesia '
    'techniques, particularly epidural analgesia combined with general anesthesia, may improve '
    'recurrence-free survival and overall survival in cancer surgery.\u2077\u2013\u00b9\u00b2 '
    'A Cochrane systematic review found limited but suggestive evidence favoring regional techniques '
    'for cancer recurrence,\u2077 and the landmark randomized controlled trial by Sessler et al. '
    'reported no significant difference in breast cancer recurrence with regional anesthesia, '
    'although secondary analyses and observational meta-analyses continue to support a potential '
    'benefit for certain cancer types.\u2078\u2013\u00b9\u00b2 '
    'If regional variation in anesthesia technique is genuine rather than an audit artifact, it may '
    'represent a modifiable source of health inequality with implications for oncologic outcomes.'
)

doc.add_paragraph(
    'The 2026 revision of the Japanese fee schedule renamed the general anesthesia code from '
    '"mask or endotracheal intubation" to "supraglottic airway device or endotracheal intubation," '
    'reflecting longstanding ambiguity in coding that may have contributed to differential auditing '
    'of general anesthesia claims.\u00b9\u00b3 Understanding the true nature of regional variation is therefore '
    'timely.'
)

doc.add_paragraph(
    'We used publicly available standardized claim ratio (SCR) data at the secondary medical area '
    'level to: (1) quantify regional variation in anesthesia practice across 335 areas; '
    '(2) test whether insurance audit variation explains observed differences through three '
    'independent sensitivity analyses; and (3) identify structural determinants of variation, '
    'including university hospital proximity and pain clinic activity.'
)

# ═══════════════════════════════════════════════════════════
# METHODS
# ═══════════════════════════════════════════════════════════

add_heading_rapm('METHODS', level=1)

add_heading_rapm('Study design and data sources', level=2)

doc.add_paragraph(
    'This was a cross-sectional ecological study using publicly available data. We used three data sources:'
)

doc.add_paragraph(
    'Standardized claim ratios (SCRs) for fiscal year 2022 (Reiwa 4), published by the Cabinet Office '
    'as part of the "Regional Variation Visualization" initiative.\u00b9\u2074 SCRs are age- and sex-adjusted '
    'ratios of observed to expected claim frequencies, where the national average equals 100. '
    'Data were available at both the prefectural level (47 prefectures) and the secondary medical area '
    '(SMA) level (335 areas).',
    style='List Bullet'
)

doc.add_paragraph(
    'Physician statistics from the 2022 Survey of Physicians, Dentists, and Pharmacists (e-Stat), '
    'providing the number of anesthesiologists by secondary medical area.\u00b9\u2075',
    style='List Bullet'
)

doc.add_paragraph(
    'Geographic data from the National Land Numerical Information dataset (A38-20) for SMA boundaries, '
    'and administrative boundary data (N03) for prefectural borders and the Northern Territories.\u00b9\u2076',
    style='List Bullet'
)

add_heading_rapm('Anesthesia codes analyzed', level=2)

doc.add_paragraph(
    'We analyzed the following procedure codes from the Japanese fee schedule (Table 1):'
)

# Table 1: Codes
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
for i, text in enumerate(['Code', 'Procedure', 'Clinical significance', 'SMAs with data']):
    hdr[i].paragraphs[0].add_run(text).bold = True
    hdr[i].paragraphs[0].runs[0].font.size = Pt(9)

codes_data = [
    ['L008', 'Closed-circuit general anesthesia', 'Primary indicator of GA volume', '334'],
    ['L002', 'Epidural anesthesia', 'Regional technique (alone or combined)', '307'],
    ['L003', 'Continuous epidural infusion', 'Direct indicator of GA+epidural combination', '331'],
    ['L004', 'Spinal anesthesia', 'Alternative to GA for suitable procedures', '334'],
    ['L009', 'Anesthesia management fee I', 'Proxy for specialist anesthesiologist staffing', '314'],
    ['L100', 'Nerve block (inpatient)', 'Indicator of pain clinic activity', '335'],
]
for row_data in codes_data:
    add_table_row(table, row_data)

doc.add_paragraph()
add_para('Table 1. Anesthesia procedure codes analyzed.', italic=True)

add_heading_rapm('University hospital mapping', level=2)

doc.add_paragraph(
    'We mapped 81 university hospitals (44 national, 8 public, 29 private) to 64 SMAs based on '
    'municipal address. Under Japan\u2019s "one medical school per prefecture" policy, all 47 prefectures '
    'have at least one university hospital, enabling within-prefecture comparisons that control for '
    'prefectural audit policy.'
)

add_heading_rapm('Sensitivity analyses for the audit hypothesis', level=2)

doc.add_paragraph(
    'We tested the null hypothesis that "all regional variation in anesthesia SCRs is explained by '
    'differential insurance auditing" using three independent approaches:'
)

doc.add_paragraph(
    'Test 1 (within-prefecture variance): If audit policy (applied uniformly within each prefecture) '
    'explains all variation, then within-prefecture variance should be negligible. We decomposed total '
    'SCR variance into between-prefecture and within-prefecture components using one-way ANOVA.',
    style='List Bullet'
)

doc.add_paragraph(
    'Test 2 (cross-code correlation): If auditing reclassifies general anesthesia claims as spinal '
    'anesthesia (or vice versa), these codes may show negative correlation at the ecological level '
    '(though clinical substitution would also produce this pattern). We calculated Pearson correlations '
    'between anesthesia code pairs and interpreted them in conjunction with the quantitative audit '
    'impact estimate (Test 3).',
    style='List Bullet'
)

doc.add_paragraph(
    'Test 3 (quantitative impact): We estimated the maximum SCR impact of audit rate differences '
    'using published aggregate audit statistics (range 0.07\u20130.28% across prefectures).\u2076',
    style='List Bullet'
)

add_heading_rapm('Combined SCR analysis', level=2)

doc.add_paragraph(
    'To further test the audit reclassification hypothesis, we computed combined SCRs (L008+L004 for '
    'GA\u2013spinal substitution; L008+L002 for GA\u2013epidural substitution). If audit-driven '
    'reclassification is the primary source of variation, combining the reclassified codes should '
    'substantially reduce the coefficient of variation.'
)

add_heading_rapm('Statistical analysis', level=2)

doc.add_paragraph(
    'We report descriptive statistics (mean, standard deviation, coefficient of variation, percentiles) '
    'for SCRs. Between-group comparisons used Welch\u2019s t-test and Cohen\u2019s d for effect size.\u00b9\u2077 Variance '
    'decomposition used one-way ANOVA (between-prefecture) and hierarchical decomposition '
    '(between-prefecture, university hospital effect, residual). Correlations are Pearson r. '
    'Standardized claim ratios were computed using indirect standardization, a standard method in '
    'small-area variation research.\u00b9\u2078 '
    'All analyses were conducted in Python 3.12 using pandas, scipy, and geopandas. Three-dimensional '
    'extruded choropleth maps were created using Plotly Mesh3d with Delaunay triangulation of polygon '
    'boundaries.\u00b9\u2079 Two-dimensional maps were created using matplotlib and geopandas.'
)

add_heading_rapm('Patient and public involvement', level=2)

doc.add_paragraph(
    'This study used publicly available aggregate data with no individual patient information. '
    'Patients were not involved in the design or conduct of this research.'
)

# ═══════════════════════════════════════════════════════════
# RESULTS
# ═══════════════════════════════════════════════════════════

add_heading_rapm('RESULTS', level=1)

add_heading_rapm('Regional variation in anesthesia practice', level=2)

doc.add_paragraph(
    'Substantial variation was observed across all anesthesia codes (Table 2). General anesthesia '
    '(L008) SCRs ranged from 2.3 to 458.9 across 334 SMAs (coefficient of variation [CV] 54.6%). '
    'Epidural anesthesia (L002) showed even greater variation (CV 83.2%), as did continuous epidural '
    'infusion (L003, CV 64.9%), the direct indicator of combined GA\u2013epidural technique. '
    'Spinal anesthesia (L004) had a CV of 56.3%. Figure 1 shows the geographic distribution of '
    'general anesthesia and spinal anesthesia SCRs, revealing an inverse spatial pattern.'
)

# Table 2
table2 = doc.add_table(rows=1, cols=7)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0].cells
for i, text in enumerate(['Code', 'Procedure', 'P10', 'P50', 'P90', 'CV (%)', 'n']):
    hdr2[i].paragraphs[0].add_run(text).bold = True
    hdr2[i].paragraphs[0].runs[0].font.size = Pt(9)

table2_data = [
    ['L008', 'General anesthesia', '32.3', '73.2', '131.5', '54.6', '334'],
    ['L002', 'Epidural anesthesia', '7.6', '73.9', '184.0', '83.2', '307'],
    ['L003', 'Continuous epidural infusion', '19.0', '73.1', '150.0', '64.9', '331'],
    ['L004', 'Spinal anesthesia', '31.1', '84.3', '169.0', '56.3', '334'],
    ['L009', 'Anesthesia management fee I', '32.5', '79.1', '149.1', '57.2', '314'],
    ['L100', 'Nerve block (inpatient)', '23.4', '72.2', '174.7', '72.1', '335'],
]
for row_data in table2_data:
    add_table_row(table2, row_data)

doc.add_paragraph()
add_para(
    'Table 2. Distribution of standardized claim ratios for anesthesia procedures across '
    'secondary medical areas (national average = 100).',
    italic=True
)

add_heading_rapm('Sensitivity analysis 1: within-prefecture variance', level=2)

doc.add_paragraph(
    'For general anesthesia (L008), 85.8% of total variance occurred within prefectures, where '
    'audit policy is uniform. This finding was consistent across codes: within-prefecture '
    'variance accounted for 72.7% (epidural), 58.0% (spinal), and 65.1% (nerve block) of total '
    'variation. Within individual prefectures, L008 SCRs varied enormously: from 61.4 to 458.9 '
    'within Tokyo (12 SMAs), and from 2.3 to 170.9 within Kumamoto (10 SMAs). Audit policy, '
    'which is applied uniformly within each prefecture, can explain at most 14.2% of L008 variation.'
)

add_heading_rapm('Sensitivity analysis 2: cross-code correlation', level=2)

doc.add_paragraph(
    'General anesthesia (L008) and spinal anesthesia (L004) showed a strong inverse correlation '
    '(r=\u22120.506, P<0.001), indicating a substitution pattern: areas with high GA volumes tend to '
    'have low spinal volumes and vice versa. This inverse correlation is consistent with genuine '
    'differences in clinical technique selection. It is also compatible in principle with audit-driven '
    'reclassification (where cases are shifted between the two codes) or defensive undercoding '
    '(where clinicians anticipate audit outcomes). However, given that the maximum audit rate '
    'difference across prefectures is only 0.21 percentage points (Test 3), reclassification of this '
    'magnitude is insufficient to produce the observed correlation. The most parsimonious '
    'interpretation is that the substitution pattern reflects clinical preference and regional '
    'practice style.'
)

doc.add_paragraph(
    'General anesthesia (L008) and epidural anesthesia (L002) showed a positive correlation '
    '(r=+0.319, P<0.001), indicating that areas with more GA also perform more epidural anesthesia. '
    'This is consistent with a common supply factor (anesthesia department capacity) rather than '
    'audit-driven substitution.'
)

add_heading_rapm('Sensitivity analysis 3: quantitative audit impact', level=2)

doc.add_paragraph(
    'The maximum difference in prefectural audit rates was 0.21 percentage points (0.07% to 0.28%). '
    'Even if this entire difference were concentrated on a single anesthesia code, it would produce '
    'an SCR shift of approximately 0.2 points\u2014less than 0.3% of the observed 62-point interquartile '
    'range for L008 and less than 0.1% of the 198-point range for L002. Audit rate variation is '
    'quantitatively incapable of explaining the observed regional differences.'
)

add_heading_rapm('Combined SCR analysis', level=2)

doc.add_paragraph(
    'Combining L008 and L004 SCRs reduced the CV from 54.6% to 43.4%\u2014a modest 19% reduction, '
    'indicating that spinal\u2013GA substitution accounts for a minority of GA variation (Figure 2B). '
    'Combining L008 and L002 actually increased the CV from 50.2% to 58.3%, as these codes are '
    'positively correlated. In both cases, the university hospital effect remained strong: combined '
    'L008+L004 showed d=1.16 (P<0.001) and combined L008+L002 showed d=1.00 (P<0.001) for the '
    'university hospital comparison.'
)

add_heading_rapm('University hospital effect', level=2)

doc.add_paragraph(
    'University hospital presence was the strongest single predictor of anesthesia practice patterns '
    '(Figure 2A). For general anesthesia (L008), the 64 university hospital SMAs had a mean SCR of '
    '130.2 compared with 67.7 for the 270 non-university SMAs (Cohen\u2019s d=1.78, t=12.49). '
    'Remarkably, in within-prefecture comparisons (which control for audit policy), university '
    'hospital SMAs had higher L008 SCRs in all 47 of 47 prefectures (100%), with a mean within-'
    'prefecture difference of +61.4 points (SD 33.7, t=12.49).'
)

doc.add_paragraph(
    'Hierarchical variance decomposition attributed 14.5% of L008 variance to between-prefecture '
    'differences, 38.5% to the university hospital effect (within prefecture), and 47.0% to residual '
    'within-prefecture variation. The university hospital effect alone\u2014a single binary variable\u2014'
    'explained nearly 40% of all variation in general anesthesia volume across Japan.'
)

doc.add_paragraph(
    'A dose\u2013response relationship was apparent in Tokyo, where the number of university hospitals '
    'per SMA ranged from 0 to 5: SCRs increased from 62\u2013110 (0 hospitals) to 83\u2013119 (2 hospitals) '
    'to 168.5 (3 hospitals) to 435.7 (5 hospitals).'
)

add_heading_rapm('Combined general\u2013epidural anesthesia', level=2)

doc.add_paragraph(
    'Continuous epidural infusion (L003), the direct indicator of combined GA\u2013epidural technique, '
    'showed a CV of 64.9%\u2014the largest variation among major anesthesia codes (Figure 3). L003 SCR '
    'was 1.73 times higher in university hospital SMAs (126.4 vs 73.2, d=0.96, P<0.001) and '
    'correlated strongly with L008 (r=0.753, P<0.001). Three audit-hypothesis tests applied to the '
    'epidural combination (L008+L002) all rejected the audit explanation: positive L008\u2013L002 '
    'correlation (r=+0.319), increased CV on combination (50.2%\u219258.3%), and persistent university '
    'hospital effect (d=1.00).'
)

add_heading_rapm('Pain clinic spillover', level=2)

doc.add_paragraph(
    'The hypothesis that pain clinic activity (proxied by nerve block SCR, L100) would predict '
    'regional anesthesia use in surgery was only weakly supported. The correlation between nerve '
    'block SCR and a composite regional anesthesia index was r=0.153 (P<0.01). However, nerve block '
    'activity correlated more strongly with GA volume (r=0.307, P<0.001), suggesting that pain clinic '
    'activity reflects overall anesthesia department capacity rather than a specific spillover of '
    'regional anesthesia skills.'
)

# ═══════════════════════════════════════════════════════════
# DISCUSSION
# ═══════════════════════════════════════════════════════════

add_heading_rapm('DISCUSSION', level=1)

add_heading_rapm('Principal findings', level=2)

doc.add_paragraph(
    'This study demonstrates that regional variation in anesthesia practice across Japan\u2019s 335 '
    'secondary medical areas is substantial and structural. Three independent sensitivity analyses '
    'converge on the conclusion that insurance audit variation explains, at most, a trivial fraction '
    'of observed differences. The dominant determinant is university hospital proximity, which alone '
    'explains nearly 40% of all variation in general anesthesia volume. The finding that this effect '
    'is present in all 47 prefectures\u2014despite differing audit policies\u2014provides compelling '
    'evidence for a supply-side, institutional mechanism.'
)

add_heading_rapm('Comparison with existing literature', level=2)

doc.add_paragraph(
    'Our findings are consistent with the broader literature on medical practice variation. '
    'The Dartmouth Atlas project in the United States has documented extensive regional variation in '
    'surgical rates that is driven primarily by physician supply and practice style rather than '
    'patient need.\u00b2\u2070\u00b7\u00b2\u00b9 Similar patterns have been reported in the United Kingdom,\u00b2\u00b2 '
    'Germany,\u00b2\u00b3 and Australia.\u00b2\u2074 Our study extends this literature in two ways. '
    'First, we exploit Japan\u2019s unique institutional structure\u2014uniform fee schedule, prefectural '
    'audit committees\u2014to distinguish administrative from clinical sources of variation. Second, '
    'we provide the first systematic analysis of anesthesia technique variation at a fine geographic '
    'scale (335 areas) in an East Asian universal insurance system.'
)

doc.add_paragraph(
    'The university hospital effect we document (d=1.78 for general anesthesia) is remarkably large '
    'compared with effect sizes typically reported in medical practice variation research. This likely '
    'reflects the Japanese medical education system, in which university departments (ikyoku) exert '
    'substantial influence over clinical practice in affiliated hospitals within their geographic '
    'sphere.\u00b2\u2075\u00b7\u00b2\u2076'
)

doc.add_paragraph(
    'Our findings are directly relevant to the ongoing debate about anesthesia technique and cancer '
    'outcomes. Settled and unsettled questions about local anesthetic dosing variation in regional '
    'anesthesia have been highlighted in this journal,\u00b2\u2077 and our study extends the variation framework '
    'to the macro level of technique selection across an entire national health system.'
)

add_heading_rapm('Implications for oncologic outcomes', level=2)

doc.add_paragraph(
    'The finding that combined GA\u2013epidural technique (L003) shows a CV of 64.9% is clinically '
    'important in light of evidence suggesting that regional anesthesia combined with general '
    'anesthesia may improve recurrence-free survival in cancer surgery.\u2077\u2013\u00b9\u00b2 The 1.73-fold '
    'difference in L003 SCR between university and non-university hospital areas suggests that '
    'patients\u2019 access to this potentially beneficial technique depends substantially on where they '
    'live\u2014a finding that challenges the equity premise of Japan\u2019s universal insurance system.\u00b9\u00b7\u00b2'
)

add_heading_rapm('The 2026 fee schedule revision', level=2)

doc.add_paragraph(
    'The renaming of the general anesthesia code in 2026 (from "mask or endotracheal intubation" '
    'to "supraglottic airway device or endotracheal intubation") reflects acknowledged ambiguity '
    'in the previous coding, which may have led to differential audit decisions regarding supraglottic '
    'airway use. Our finding that audit-driven variation is quantitatively small suggests that this '
    'revision, while appropriate for coding clarity, is unlikely to substantially alter the geographic '
    'patterns we have documented.'
)

add_heading_rapm('Strengths and limitations', level=2)

doc.add_paragraph(
    'Strengths include the use of age- and sex-standardized ratios that cover the entire national '
    'population, analysis at a fine geographic scale (335 areas), multiple independent sensitivity '
    'analyses, and within-prefecture comparisons that control for audit policy. The ecological design '
    'avoids individual consent requirements and allows comprehensive geographic coverage.'
)

doc.add_paragraph(
    'Limitations include the ecological fallacy inherent in area-level analysis;\u00b2\u2078 the inability to '
    'distinguish individual patient-level decisions from aggregate patterns; the post-audit nature '
    'of SCR data (which reflects reimbursed rather than intended practice, though our sensitivity '
    'analyses suggest this distinction is quantitatively minor); the cross-sectional design, which '
    'precludes causal inference; incomplete separation of the university hospital effect from urban '
    'concentration (university hospitals are typically in the largest city of each prefecture); and '
    'the absence of code-specific audit rate data. Additionally, "defensive undercoding"\u2014where '
    'clinicians preemptively avoid claiming procedures they expect to be audited\u2014cannot be '
    'quantified from claims data.\u00b2\u2079'
)

add_heading_rapm('Conclusions and policy implications', level=2)

doc.add_paragraph(
    'Regional variation in anesthesia practice in Japan is genuine and large, driven primarily by '
    'institutional factors\u2014particularly university hospital proximity\u2014rather than by '
    'differential insurance auditing. The 1.73-fold variation in combined GA\u2013epidural anesthesia '
    'between university and non-university hospital areas is of particular concern given emerging '
    'evidence of oncologic benefit.\u2077\u2013\u00b9\u00b2 Policy responses could include targeted education programs, '
    'specialist outreach from university centers, and monitoring of anesthesia technique patterns '
    'as a quality indicator.\u00b3\u2070 The sensitivity analysis framework developed here\u2014using within-'
    'prefecture variance decomposition, cross-code correlation, and quantitative audit impact '
    'estimation\u2014is applicable to investigating practice variation in any procedure under Japan\u2019s '
    'universal insurance system.'
)

# ═══════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_rapm('REFERENCES', level=1)

refs = [
    '1. Ikegami N, Yoo B-K, Hashimoto H, et al. Japanese universal health coverage: evolution, '
    'achievements, and challenges. Lancet 2011;378:1106\u201315.',
    '2. Ikegami N. Universal Health Coverage for Inclusive and Sustainable Development: Lessons from '
    'Japan. Washington, DC: World Bank, 2014.',
    '3. Matsuda S, Fujimori K. The claim database in Japan. Asian Pac J Dis Manag 2012;6:55\u201359.',
    '4. Suto M, Iba A, Sugiyama T, et al. Literature review of studies using the National Database '
    'of Health Insurance Claims of Japan (NDB): limitations and strategies in using the NDB for '
    'research. JMA J 2023;7:10\u201320.',
    '5. Hamada H, Sekimoto M, Imanaka Y. Effects of the per diem prospective payment system with '
    'DRG-like grouping system (DPC/PDPS) on resource usage and healthcare quality in Japan. Health '
    'Policy 2012;107:194\u2013201.',
    '6. Social Insurance Medical Fee Payment Fund (Shiharai Kikin). Review Statistics, Fiscal Year '
    '2022. https://www.ssk.or.jp/tokeijoho/shinsatokei/index.html (accessed 15 Mar 2026).',
    '7. Pei L, Tan G, Wang L, et al. Anesthetic techniques for risk of malignant tumor recurrence. '
    'Cochrane Database Syst Rev 2019;(11):CD008877.',
    '8. Sessler DI, Pei L, Huang Y, et al. Recurrence of breast cancer after regional or general '
    'anesthesia: a randomized controlled trial. Lancet 2019;394:1807\u201315.',
    '9. Exadaktylos AK, Buggy DJ, Moriarty DC, et al. Can anesthetic technique for primary breast '
    'cancer surgery affect recurrence or metastasis? Anesthesiology 2006;105:660\u20134.',
    '10. Sun Y, Li T, Gan TJ. The effects of perioperative regional anesthesia and analgesia on cancer '
    'recurrence and survival after oncology surgery: a systematic review and meta-analysis. Reg '
    'Anesth Pain Med 2015;40:589\u201398.',
    '11. Xie S, Li L, Meng F, et al. Regional anesthesia might reduce recurrence and metastasis rates '
    'in adult patients with cancers after surgery: a meta-analysis. BMC Anesthesiol 2024;24:19.',
    '12. Zhang D, Jiang J, Liu J, et al. Effects of perioperative epidural analgesia on cancer '
    'recurrence and survival. Front Oncol 2021;11:798435.',
    '13. Ministry of Health, Labour and Welfare. Revision of medical fee schedule, 2026 '
    '(Reiwa 8 nendo shinryo houshu kaitei). https://www.mhlw.go.jp (accessed 15 Mar 2026).',
    '14. Cabinet Office. Regional variation visualization (chiikisa no mieruka). '
    'https://www5.cao.go.jp/keizai-shimon/kaigi/special/reform/mieruka/ (accessed 10 Mar 2026).',
    '15. Ministry of Health, Labour and Welfare. Survey of Physicians, Dentists and Pharmacists, '
    '2022. e-Stat. https://www.e-stat.go.jp (accessed 12 Mar 2026).',
    '16. Ministry of Land, Infrastructure, Transport and Tourism. National Land Numerical Information '
    'download service. https://nlftp.mlit.go.jp (accessed 10 Mar 2026).',
    '17. Cohen J. Statistical Power Analysis for the Behavioral Sciences. 2nd ed. '
    'Hillsdale, NJ: Lawrence Erlbaum, 1988.',
    '18. Julious SA, Nicholl J, George S. Why do we continue to use standardized mortality ratios for '
    'small area comparisons? J Public Health Med 2001;23:40\u20136.',
    '19. Plotly Technologies Inc. Plotly: collaborative data science. Montreal, QC, 2015. '
    'https://plotly.com (accessed 15 Mar 2026).',
    '20. Wennberg JE. Tracking Medicine: A Researcher\u2019s Quest to Understand Health Care. '
    'New York: Oxford University Press, 2010.',
    '21. Wennberg JE, Fisher ES, Skinner JS. Geography and the debate over Medicare reform. '
    'Health Aff (Millwood) 2002;Suppl Web Exclusives:W96\u2013114.',
    '22. NHS RightCare. The NHS Atlas of Variation in Healthcare, 2015. London: Public Health England.',
    '23. Wengler A, Nimptsch U, Mansky T. Hip and knee replacement surgery (arthroplasty) and '
    'geographic variations in Germany. Dtsch Arztebl Int 2014;111:407\u201316.',
    '24. Australian Commission on Safety and Quality in Health Care. Australian Atlas of Healthcare '
    'Variation. Sydney: ACSQHC, 2015.',
    '25. Otsuka T. The ikyoku system of university orthopedic surgery departments: an in-hospital '
    'organizational system unique to Japan. J Orthop Sci 2012;17:513\u201314.',
    '26. Onishi H. History of Japanese medical education. Korean J Med Educ 2018;30:283\u201394.',
    '27. Koscielniak-Nielsen ZJ, Helbo-Hansen HS. Settled science or unwarranted variation in local '
    'anesthetic dosing? Reg Anesth Pain Med 2019;44:998\u20131000.',
    '28. Greenland S. Ecologic versus individual-level sources of bias in ecologic estimates of '
    'contextual health effects. Int J Epidemiol 2001;30:1343\u201350.',
    '29. Sekimoto M, Imanaka Y, Kobayashi H, et al. Impact of hospital characteristics on the cost '
    'and practice patterns of acute myocardial infarction patients in Japan. Health Policy '
    '2006;78:34\u201345.',
    '30. Merry AF, Cooper JB, Soyannwo O, et al. International standards for a safe practice of '
    'anesthesia 2010. Can J Anesth 2010;57:1027\u201334.',
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.name = 'Arial'

# ═══════════════════════════════════════════════════════════
# FIGURES
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_rapm('FIGURES', level=1)

# Multi-panel figures
figure_list = [
    ('/home/ubuntu/rapm_fig1_en.png',
     'Figure 1. Geographic distribution of anesthesia standardized claim ratios by secondary medical area. '
     '(A) General anesthesia (L008) SCR. (B) Spinal anesthesia (L004) SCR. Note the inverse spatial pattern: '
     'areas with high GA tend to have low spinal anesthesia volumes and vice versa (r=\u22120.506). '
     'Solid lines indicate prefectural boundaries; dashed lines indicate SMA boundaries. '
     'Northern Territories shown in white (no SMA designated).'),
    ('/home/ubuntu/rapm_fig2_en.png',
     'Figure 2. University hospital effect and combined SCR analysis. '
     '(A) University hospital presence by SMA: red circles indicate SMAs containing one or more '
     'university hospitals (64 of 335 SMAs). (B) Combined GA + spinal anesthesia (L008+L004) SCR: '
     'combining these codes neutralizes potential audit-driven reclassification, yet substantial '
     'regional variation persists (CV=43.4%).'),
    ('/home/ubuntu/rapm_fig3_en.png',
     'Figure 3. Combined GA\u2013epidural anesthesia variation. '
     '(A) Continuous epidural infusion (L003) SCR\u2014direct indicator of combined GA\u2013epidural technique '
     '(CV=64.9%). (B) L003/L008 ratio (epidural combination rate per GA case), adjusted for surgery '
     'volume. Higher values indicate greater use of combined GA\u2013epidural technique.'),
    ('/home/ubuntu/rapm_fig4_en.png',
     'Figure 4. Three-dimensional extruded choropleth map. Color represents L003/L008 ratio '
     '(combined epidural rate); height represents anesthesiologist count per SMA. Areas with both '
     'high epidural combination rate and high specialist count appear as tall, green-colored polygons, '
     'typically corresponding to university hospital areas in metropolitan regions.'),
]

for path, caption in figure_list:
    add_figure(path, caption, width=Inches(5.5))
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# STROBE CHECKLIST
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_rapm('STROBE Checklist (Cross-sectional study)', level=1)

strobe_items = [
    ('1a', 'Title and abstract', 'Title page and structured abstract'),
    ('1b', 'Informative abstract', 'Structured abstract with all required elements'),
    ('2', 'Background/rationale', 'Introduction paragraphs 1\u20133'),
    ('3', 'Objectives', 'Introduction paragraph 5'),
    ('4', 'Study design', 'Methods: Study design and data sources'),
    ('5', 'Setting', 'Methods: 335 SMAs, FY2022 data'),
    ('6a', 'Participants/eligibility', 'N/A (ecological study using aggregate data)'),
    ('7', 'Variables', 'Methods: Anesthesia codes analyzed (Table 1)'),
    ('8', 'Data sources', 'Methods: Study design and data sources'),
    ('9', 'Bias', 'Methods: Sensitivity analyses; Discussion: Limitations'),
    ('10', 'Study size', 'Methods: 335 SMAs, 47 prefectures'),
    ('11', 'Quantitative variables', 'Methods: Statistical analysis'),
    ('12', 'Statistical methods', 'Methods: Statistical analysis'),
    ('13', 'Participants', 'Results: 334 SMAs with L008 data'),
    ('14', 'Descriptive data', 'Results: Table 2'),
    ('15', 'Outcome data', 'Results: Regional variation section'),
    ('16', 'Main results', 'Results: Sensitivity analyses 1\u20133, University hospital effect'),
    ('17', 'Other analyses', 'Results: Combined SCR, Pain clinic spillover'),
    ('18', 'Key results', 'Discussion: Principal findings'),
    ('19', 'Limitations', 'Discussion: Strengths and limitations'),
    ('20', 'Interpretation', 'Discussion: Comparison with existing literature'),
    ('21', 'Generalizability', 'Discussion: Conclusions and policy implications'),
    ('22', 'Funding', '[To be completed]'),
]

strobe_table = doc.add_table(rows=1, cols=3)
strobe_table.style = 'Table Grid'
strobe_hdr = strobe_table.rows[0].cells
for i, text in enumerate(['Item', 'STROBE recommendation', 'Location in manuscript']):
    strobe_hdr[i].paragraphs[0].add_run(text).bold = True
    strobe_hdr[i].paragraphs[0].runs[0].font.size = Pt(9)

for item in strobe_items:
    add_table_row(strobe_table, item)

# ═══════════════════════════════════════════════════════════
# DECLARATIONS
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_rapm('DECLARATIONS', level=1)

add_heading_rapm('Ethics approval', level=2)
doc.add_paragraph(
    'This study used publicly available aggregate data. No individual-level data were accessed. '
    'Ethics committee approval was not required under the Ethical Guidelines for Medical and '
    'Biological Research Involving Human Subjects (Japan, 2021 revision).'
)

add_heading_rapm('Data availability', level=2)
doc.add_paragraph(
    'All data used in this study are publicly available. SCR data: Cabinet Office regional variation '
    'visualization (https://www5.cao.go.jp/keizai-shimon/kaigi/special/reform/mieruka/). '
    'Physician statistics: e-Stat (https://www.e-stat.go.jp). GIS boundaries: National Land '
    'Numerical Information (https://nlftp.mlit.go.jp). Analysis code is available from the '
    'corresponding author on request.'
)

add_heading_rapm('Competing interests', level=2)
doc.add_paragraph('[To be completed by authors.]')

add_heading_rapm('Funding', level=2)
doc.add_paragraph('[To be completed by authors.]')

add_heading_rapm('Contributors', level=2)
doc.add_paragraph('[To be completed by authors.]')

add_heading_rapm('AI statement', level=2)
doc.add_paragraph(
    '[To be completed. Per RAPM/BMJ policy, authors must provide details of any AI technology used '
    'in the preparation of this manuscript, including what AI technology was used, the reason for '
    'its use, and what task it performed.]'
)

add_heading_rapm('Transparency declaration', level=2)
doc.add_paragraph(
    'The lead author (the manuscript\u2019s guarantor) affirms that the manuscript is an honest, '
    'accurate, and transparent account of the study being reported; that no important aspects '
    'of the study have been omitted; and that any discrepancies from the study as planned have '
    'been explained.'
)

# ═══════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════

output_path = '/home/ubuntu/regional_anesthesia_RAPM_EN.docx'
doc.save(output_path)
print(f"Saved: {output_path}")
