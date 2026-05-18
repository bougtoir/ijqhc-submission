#!/usr/bin/env python3
"""
Create RAPM-format docx (Japanese version) for regional variation in anesthesia practice in Japan.
- Japanese translation of the RAPM English manuscript
- Same structure: STROBE, structured abstract, 3 key message boxes
- Tables + Figures combined <= 6 (2 tables + 4 multi-panel figures)
- Japanese panel figures used
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Style setup ──
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# Helper functions
def add_heading_rapm(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    return p

def add_table_row(table, cells, bold=False):
    row = table.add_row()
    for i, cell_text in enumerate(cells):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(cell_text))
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.bold = bold
    return row

def add_figure(path, caption, width=Inches(6)):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.italic = True
    else:
        add_para(f'[Figure: {path} not found]', italic=True)


# ═══════════════════════════════════════════════════════════
# 表題ページ
# ═══════════════════════════════════════════════════════════

title_text = (
    "日本の335二次医療圏における麻酔実施の地域差："
    "全国保険請求データの横断研究"
)
p = doc.add_paragraph()
run = p.add_run(title_text)
run.font.name = 'Arial'
run.font.size = Pt(16)
run.bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

add_para('[著者名]', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('[所属機関]', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('[責任著者: 氏名、住所、メールアドレス]', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
add_para('語数: 約3,800語', italic=True)
add_para('表: 2', italic=True)
add_para('図: 4（マルチパネルコロプレスマップ3枚 + 三次元押出マップ1枚）', italic=True)
add_para('補足図: 5（追加3Dマップ2枚 + 追加2Dマップ3枚）', italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 構造化抄録（RAPM形式: Background/Methods/Results/Conclusion）
# ═══════════════════════════════════════════════════════════

add_heading_rapm('抄録', level=1)

p = doc.add_paragraph()
run = p.add_run('背景 ')
run.bold = True
run.font.name = 'Arial'
p.add_run(
    '麻酔実施の地域差は国際的に報告されているが、臨床的選好、施設要因、行政的メカニズムの'
    '相対的寄与は不明である。日本の国民皆保険制度は、都道府県別の保険審査（査定）を義務づけて'
    'おり、行政的変動と臨床的変動を区別するユニークな環境を提供する。'
).font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run('方法 ')
run.bold = True
run.font.name = 'Arial'
p.add_run(
    '2022年度（令和4年度）の公開標準化レセプト出現比（SCR）を用いた横断的生態学的研究。'
    '335二次医療圏における6種の麻酔関連手技コードを分析した：全身麻酔（L008）、硬膜外麻酔'
    '（L002）、持続硬膜外注入（L003）、脊椎麻酔（L004）、麻酔管理料I（L009）、神経ブロック'
    '（L100）。保険審査の地域差が観測された差異を説明するか否かを、3つの感度分析で検証した：'
    '県内分散分解、コード間相関分析、定量的査定影響推定。'
).font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run('結果 ')
run.bold = True
run.font.name = 'Arial'
p.add_run(
    '顕著な地域差が観測された：変動係数（CV）は全身麻酔（L008）54.6%、硬膜外麻酔（L002）'
    '83.2%、持続硬膜外注入（L003）64.9%であった。3つの感度分析は、査定の変動では観測された'
    '差異を説明できないことを示した：L008分散の85.8%は県内（査定方針が統一されている）で発生；'
    '全身麻酔と脊椎麻酔は逆相関（r=\u22120.506）；査定率の差は観測されたSCR変動の1%未満しか'
    '説明できなかった。大学病院の存在が最も強い単一予測因子であり、L008総分散の38.5%を説明した'
    '（Cohen\u2019s d=1.78）。全身麻酔\u2013硬膜外併用手技（L003）は大学病院圏で1.73倍高かった。'
).font.name = 'Arial'

p = doc.add_paragraph()
run = p.add_run('結論 ')
run.bold = True
run.font.name = 'Arial'
p.add_run(
    '日本における麻酔実施の地域差は構造的なものであり、保険査定のアーティファクトではない。'
    '大学病院の近接性が主要な決定因子である。区域麻酔手技が腫瘍学的転帰の改善と関連する'
    'エビデンスを考慮すると、この地域差は修正可能な健康格差の源泉である可能性がある。'
).font.name = 'Arial'

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# キーメッセージ（RAPM 3ボックス形式）
# ═══════════════════════════════════════════════════════════

add_heading_rapm('既に知られていること', level=2)
doc.add_paragraph(
    '手術・麻酔実施の地域差は多くの国で報告されているが、臨床的選好、施設要因、行政的メカニズム'
    '（保険審査など）の相対的寄与は不明のままである。',
    style='List Bullet'
)
doc.add_paragraph(
    '日本の国民皆保険制度は全ての請求を都道府県別の標準化された審査に付しており、行政的変動と'
    '臨床的変動を区別できるユニークな自然実験の場を提供する。',
    style='List Bullet'
)

add_heading_rapm('本研究が加えるもの', level=2)
doc.add_paragraph(
    '3つの独立した感度分析により、日本における麻酔実施の地域差は保険査定のアーティファクトではなく'
    '構造的なものであることが示された。',
    style='List Bullet'
)
doc.add_paragraph(
    '大学病院の存在は全身麻酔量の分散の38.5%を説明する\u2014単一の二値予測因子としては顕著に'
    '大きな効果\u2014であり、この効果は47都道府県すべてで一貫している。',
    style='List Bullet'
)

add_heading_rapm('本研究が研究・臨床・政策に与えうる影響', level=2)
doc.add_paragraph(
    '腫瘍学的利益をもたらしうる全身麻酔\u2013硬膜外併用は、最大の地域格差（CV=64.9%）を示しており、'
    '健康格差縮小の修正可能なターゲットである。',
    style='List Bullet'
)
doc.add_paragraph(
    '本研究で開発した感度分析フレームワーク\u2014県内分散分解、コード間相関、定量的査定影響推定\u2014は、'
    '日本の国民皆保険制度下の任意の手技における実施変動の調査に適用可能である。',
    style='List Bullet'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 緒言
# ═══════════════════════════════════════════════════════════

add_heading_rapm('緒言', level=1)

doc.add_paragraph(
    '1961年に確立された日本の国民皆保険制度は、全国民を標準化された診療報酬点数表（診療報酬'
    '点数表）の下でカバーし、各医療行為に対する報酬を規定している。\u00b9\u00b7\u00b2 この制度は地理的条件に'
    'かかわらず医療への公平なアクセスを保障するために設計された。しかし、レセプト情報・特定健診'
    '等情報データベース（NDB）を用いた研究は、手術手技、処方パターン、検査を含む医療実施の'
    '顕著な地域差を明らかにしている。\u00b3\u2013\u2075'
)

doc.add_paragraph(
    '日本の制度の特徴的な点は、保険審査（審査）の義務化である。全ての請求は支払前に都道府県の'
    '審査委員会により審査され、不適切と判断された請求は減点または査定される（査定）。査定率は'
    '都道府県間で0.07%から0.28%まで変動する。\u2076 これは根本的な問いを提起する：請求データに'
    'おける観測された地域差は臨床実施の真の差異を反映しているのか、それとも差異的な査定の'
    'アーティファクトなのか？'
)

doc.add_paragraph(
    'この問いは麻酔実施にとって特に関連がある。麻酔手技の選択\u2014全身麻酔単独か全身麻酔\u2013'
    '区域麻酔併用か\u2014は単なる選好の問題ではない。区域麻酔手技、特に全身麻酔と硬膜外鎮痛の'
    '併用が、がん手術における無再発生存期間と全生存期間を改善しうるというエビデンスが蓄積して'
    'いる。\u2077\u2013\u00b9\u00b2 Cochrane系統的レビューは区域麻酔手技ががん再発に対して限定的だが示唆的な'
    'エビデンスを見出し、\u2077 Sesslerらのランドマーク無作為化比較試験は乳がん再発における区域麻酔'
    'の有意差を報告しなかったが、二次解析と観察研究のメタアナリシスは特定のがん種における潜在的'
    '利益を支持し続けている。\u2078\u2013\u00b9\u00b2 麻酔手技の地域差が査定アーティファクトではなく真のもので'
    'あるならば、それは腫瘍学的転帰に対する含意を持つ修正可能な健康格差の源泉である可能性がある。'
)

doc.add_paragraph(
    '2026年の診療報酬改定は全身麻酔コードの名称を「マスク又は気管内挿管」から「声門上器具使用'
    '又は気管内挿管」に変更し、声門上器具の使用に関する差異的な査定に寄与した可能性のある'
    'コーディングの曖昧さを反映している。\u00b9\u00b3 地域差の真の性質を理解することは時宜にかなっている。'
)

doc.add_paragraph(
    '我々は、二次医療圏レベルの公開SCRデータを用いて：（1）335地域における麻酔実施の地域差を'
    '定量化し；（2）3つの独立した感度分析により保険査定の変動が観測された差異を説明するかを'
    '検証し；（3）大学病院の近接性やペインクリニック活動を含む変動の構造的決定因子を同定した。'
)

# ═══════════════════════════════════════════════════════════
# 方法
# ═══════════════════════════════════════════════════════════

add_heading_rapm('方法', level=1)

add_heading_rapm('研究デザインとデータソース', level=2)

doc.add_paragraph(
    '本研究は公開データを用いた横断的生態学的研究である。3つのデータソースを使用した：'
)

doc.add_paragraph(
    '2022年度（令和4年度）の標準化レセプト出現比（SCR）。内閣府の「地域差の見える化」'
    'イニシアティブの一環として公表。\u00b9\u2074 SCRは、全国平均を100とする、年齢・性別調整済みの'
    '観測値と期待値の比率である。データは都道府県レベル（47都道府県）と二次医療圏レベル'
    '（335地域）の両方で利用可能であった。',
    style='List Bullet'
)

doc.add_paragraph(
    '2022年医師・歯科医師・薬剤師統計（e-Stat）からの医師統計。二次医療圏別の麻酔科医数を'
    '提供。\u00b9\u2075',
    style='List Bullet'
)

doc.add_paragraph(
    '国土数値情報データセット（A38-20）からの二次医療圏境界の地理データ、および都道府県境と'
    '北方領土の行政区域データ（N03）。\u00b9\u2076',
    style='List Bullet'
)

add_heading_rapm('分析した麻酔コード', level=2)

doc.add_paragraph(
    '日本の診療報酬点数表から以下の手技コードを分析した（表1）：'
)

# 表1
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
for i, text in enumerate(['コード', '手技', '臨床的意義', 'データのあるSMA数']):
    hdr[i].paragraphs[0].add_run(text).bold = True
    hdr[i].paragraphs[0].runs[0].font.size = Pt(9)

codes_data = [
    ['L008', '閉鎖循環式全身麻酔', '全身麻酔量の一次指標', '334'],
    ['L002', '硬膜外麻酔', '区域麻酔手技（単独または併用）', '307'],
    ['L003', '持続硬膜外注入', '全身麻酔+硬膜外併用の直接指標', '331'],
    ['L004', '脊椎麻酔', '適切な手術に対する全身麻酔の代替', '334'],
    ['L009', '麻酔管理料I', '専門麻酔科医配置の代理指標', '314'],
    ['L100', '神経ブロック（入院）', 'ペインクリニック活動の指標', '335'],
]
for row_data in codes_data:
    add_table_row(table, row_data)

doc.add_paragraph()
add_para('表1. 分析した麻酔手技コード。', italic=True)

add_heading_rapm('大学病院マッピング', level=2)

doc.add_paragraph(
    '81大学病院（国立44、公立8、私立29）を市区町村住所に基づき64の二次医療圏にマッピングした。'
    '日本の「一県一医大」政策のもと、全47都道府県に少なくとも1つの大学病院があり、都道府県の'
    '査定方針を制御した県内比較が可能となる。'
)

add_heading_rapm('査定仮説の感度分析', level=2)

doc.add_paragraph(
    '「麻酔SCRの全ての地域差は差異的な保険査定により説明される」という帰無仮説を3つの独立した'
    'アプローチで検証した：'
)

doc.add_paragraph(
    '検定1（県内分散）：査定方針（各県内で統一的に適用）が全ての変動を説明するなら、県内分散は'
    '無視できるはずである。一元配置分散分析により総SCR分散を県間成分と県内成分に分解した。',
    style='List Bullet'
)

doc.add_paragraph(
    '検定2（コード間相関）：査定が全身麻酔請求を脊椎麻酔に再分類する（またはその逆）なら、'
    'これらのコードは生態学的レベルで負の相関を示しうる（臨床的代替も同様のパターンを生じうるが）。'
    '麻酔コードペア間のPearson相関を計算し、定量的査定影響推定（検定3）と併せて解釈した。',
    style='List Bullet'
)

doc.add_paragraph(
    '検定3（定量的影響）：公表された査定統計の集計値（都道府県間の範囲0.07\u20130.28%）を用いて、'
    '査定率差のSCRへの最大影響を推定した。\u2076',
    style='List Bullet'
)

add_heading_rapm('合計SCR分析', level=2)

doc.add_paragraph(
    '査定による再分類仮説をさらに検証するため、合計SCRを算出した（全身麻酔\u2013脊椎代替のための'
    'L008+L004；全身麻酔\u2013硬膜外代替のためのL008+L002）。査定による再分類が変動の主要な'
    '源泉であれば、再分類されたコードを合計することで変動係数が大幅に減少するはずである。'
)

add_heading_rapm('統計解析', level=2)

doc.add_paragraph(
    'SCRの記述統計量（平均、標準偏差、変動係数、パーセンタイル値）を報告する。群間比較には'
    'Welchのt検定とCohen\u2019s dを効果量として使用した。\u00b9\u2077 分散分解は一元配置分散分析（県間）と'
    '階層的分解（県間、大学病院効果、残差）を使用した。相関はPearsonのrである。標準化レセプト'
    '出現比は間接標準化を用いて算出され、小地域変動研究の標準的手法である。\u00b9\u2078 '
    '全ての解析はPython 3.12（pandas、scipy、geopandas）で実施した。三次元押出コロプレスマップは'
    'Plotly Mesh3dとポリゴン境界のDelaunay三角分割を用いて作成した。\u00b9\u2079 二次元マップは'
    'matplotlibとgeopandasを用いて作成した。'
)

add_heading_rapm('患者・市民の参画', level=2)

doc.add_paragraph(
    '本研究は個人の患者情報を含まない公開集計データを使用した。患者は本研究のデザインまたは'
    '実施に関与していない。'
)

# ═══════════════════════════════════════════════════════════
# 結果
# ═══════════════════════════════════════════════════════════

add_heading_rapm('結果', level=1)

add_heading_rapm('麻酔実施の地域差', level=2)

doc.add_paragraph(
    '全ての麻酔コードで顕著な変動が観測された（表2）。全身麻酔（L008）SCRは334の二次医療圏で'
    '2.3から458.9の範囲であった（変動係数[CV] 54.6%）。硬膜外麻酔（L002）はさらに大きな変動を'
    '示し（CV 83.2%）、全身麻酔\u2013硬膜外併用の直接指標である持続硬膜外注入（L003）も同様で'
    'あった（CV 64.9%）。脊椎麻酔（L004）のCVは56.3%であった。図1は全身麻酔と脊椎麻酔SCRの'
    '地理的分布を示し、逆の空間パターンを明らかにしている。'
)

# 表2
table2 = doc.add_table(rows=1, cols=7)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0].cells
for i, text in enumerate(['コード', '手技', 'P10', 'P50', 'P90', 'CV (%)', 'n']):
    hdr2[i].paragraphs[0].add_run(text).bold = True
    hdr2[i].paragraphs[0].runs[0].font.size = Pt(9)

table2_data = [
    ['L008', '全身麻酔', '32.3', '73.2', '131.5', '54.6', '334'],
    ['L002', '硬膜外麻酔', '7.6', '73.9', '184.0', '83.2', '307'],
    ['L003', '持続硬膜外注入', '19.0', '73.1', '150.0', '64.9', '331'],
    ['L004', '脊椎麻酔', '31.1', '84.3', '169.0', '56.3', '334'],
    ['L009', '麻酔管理料I', '32.5', '79.1', '149.1', '57.2', '314'],
    ['L100', '神経ブロック（入院）', '23.4', '72.2', '174.7', '72.1', '335'],
]
for row_data in table2_data:
    add_table_row(table2, row_data)

doc.add_paragraph()
add_para(
    '表2. 二次医療圏における麻酔手技の標準化レセプト出現比の分布（全国平均=100）。',
    italic=True
)

add_heading_rapm('感度分析1：県内分散', level=2)

doc.add_paragraph(
    '全身麻酔（L008）では、総分散の85.8%が査定方針が統一されている県内で発生した。この所見は'
    'コード間で一貫しており、県内分散は硬膜外72.7%、脊椎58.0%、神経ブロック65.1%を占めた。'
    '個々の都道府県内でL008 SCRは大きく変動した：東京都内（12二次医療圏）で61.4から458.9、'
    '熊本県内（10二次医療圏）で2.3から170.9。各県内で統一的に適用される査定方針は、L008変動の'
    '最大14.2%しか説明できない。'
)

add_heading_rapm('感度分析2：コード間相関', level=2)

doc.add_paragraph(
    '全身麻酔（L008）と脊椎麻酔（L004）は強い逆相関を示した（r=\u22120.506, P<0.001）。これは'
    '代替パターンを示している：全身麻酔量の多い地域は脊椎麻酔量が少なく、その逆も然りである。'
    'この逆相関は臨床的手技選択の真の差異と整合的である。原理的には査定による再分類（症例が'
    '2つのコード間で移動する）や防衛的過少請求（臨床家が査定結果を予期する）とも両立しうる。'
    'しかし、都道府県間の最大査定率差がわずか0.21ポイント（検定3）であることを考えると、'
    'この規模の再分類では観測された相関を生じさせることは不可能である。最も節約的な解釈は、'
    '代替パターンが臨床的選好と地域の診療スタイルを反映しているというものである。'
)

doc.add_paragraph(
    '全身麻酔（L008）と硬膜外麻酔（L002）は正の相関を示した（r=+0.319, P<0.001）。全身麻酔量'
    'の多い地域は硬膜外麻酔もより多く実施している。これは査定による代替ではなく、共通の供給因子'
    '（麻酔科の体制）と整合的である。'
)

add_heading_rapm('感度分析3：定量的査定影響', level=2)

doc.add_paragraph(
    '都道府県間の査定率の最大差は0.21ポイント（0.07%から0.28%）であった。この差が単一の麻酔'
    'コードに集中したとしても、約0.2ポイントのSCR変動しか生じない\u2014L008の観測された四分位範囲'
    '62ポイントの0.3%未満、L002の198ポイントの範囲の0.1%未満。査定率の変動は観測された地域差を'
    '定量的に説明することが不可能である。'
)

add_heading_rapm('合計SCR分析', level=2)

doc.add_paragraph(
    'L008とL004のSCRを合計するとCVは54.6%から43.4%に減少した\u2014控えめな19%の減少であり、'
    '脊椎\u2013全身麻酔の代替がGA変動の少数派を占めることを示す（図2B）。L008とL002を合計すると'
    'CVは50.2%から58.3%に実際に増加した（これらのコードは正の相関があるため）。いずれの場合も'
    '大学病院効果は強く残存した：L008+L004合計でd=1.16（P<0.001）、L008+L002合計でd=1.00'
    '（P<0.001）。'
)

add_heading_rapm('大学病院効果', level=2)

doc.add_paragraph(
    '大学病院の存在は麻酔実施パターンの最も強い単一予測因子であった（図2A）。全身麻酔（L008）'
    'では、64の大学病院所在二次医療圏の平均SCRは130.2であり、270の非大学病院圏の67.7と比較して'
    '大きかった（Cohen\u2019s d=1.78, t=12.49）。注目すべきことに、県内比較（査定方針を制御）において'
    '大学病院所在圏は47都道府県中47（100%）でL008 SCRが高く、平均県内差は+61.4ポイント'
    '（SD 33.7, t=12.49）であった。'
)

doc.add_paragraph(
    '階層的分散分解では、L008分散の14.5%を県間差、38.5%を大学病院効果（県内）、47.0%を残差に'
    '帰属させた。大学病院効果のみ\u2014単一の二値変数\u2014が日本全国の全身麻酔量の全変動の約40%を'
    '説明した。'
)

doc.add_paragraph(
    '東京都では用量反応関係が明確であった。二次医療圏あたりの大学病院数は0から5の範囲で、'
    'SCRは62\u2013110（0病院）から83\u2013119（2病院）、168.5（3病院）、435.7（5病院）へと増加した。'
)

add_heading_rapm('全身麻酔\u2013硬膜外併用', level=2)

doc.add_paragraph(
    '全身麻酔\u2013硬膜外併用の直接指標である持続硬膜外注入（L003）はCV 64.9%を示した\u2014'
    '主要麻酔コード中最大の変動である（図3）。L003 SCRは大学病院圏で1.73倍高く（126.4 vs 73.2, '
    'd=0.96, P<0.001）、L008と強く相関した（r=0.753, P<0.001）。硬膜外併用（L008+L002）に'
    '適用した3つの査定仮説検定は全て査定説明を棄却した：L008\u2013L002の正の相関（r=+0.319）、'
    '合計によるCVの増加（50.2%\u219258.3%）、持続する大学病院効果（d=1.00）。'
)

add_heading_rapm('ペインクリニックの波及効果', level=2)

doc.add_paragraph(
    'ペインクリニック活動（神経ブロックSCR, L100で代理）が手術における区域麻酔使用を予測する'
    'という仮説は弱い支持のみであった。神経ブロックSCRと複合区域麻酔指数の相関はr=0.153'
    '（P<0.01）であった。しかし、神経ブロック活動はGA量とより強く相関しており（r=0.307, P<0.001）、'
    'ペインクリニック活動が区域麻酔技術の特定の波及ではなく、麻酔科全体の体制を反映していることを'
    '示唆する。'
)

# ═══════════════════════════════════════════════════════════
# 考察
# ═══════════════════════════════════════════════════════════

add_heading_rapm('考察', level=1)

add_heading_rapm('主要な所見', level=2)

doc.add_paragraph(
    '本研究は、日本の335二次医療圏における麻酔実施の地域差が顕著かつ構造的であることを実証した。'
    '3つの独立した感度分析は、保険査定の変動が観測された差異のごく一部しか説明できないという'
    '結論に収束する。主要な決定因子は大学病院の近接性であり、これだけで全身麻酔量の全変動の'
    '約40%を説明する。この効果が査定方針の異なる47都道府県全てに存在するという所見は、供給側の'
    '制度的メカニズムの説得力のあるエビデンスを提供する。'
)

add_heading_rapm('既存文献との比較', level=2)

doc.add_paragraph(
    '我々の所見は医療実施変動に関する広範な文献と整合的である。米国のDartmouth Atlas projectは'
    '手術率の広範な地域差を文書化しており、それは患者のニーズよりも医師の供給と診療スタイルに'
    '主に起因する。\u00b2\u2070\u00b7\u00b2\u00b9 同様のパターンが英国、\u00b2\u00b2 ドイツ、\u00b2\u00b3 オーストラリアでも報告'
    'されている。\u00b2\u2074 本研究はこの文献を2つの点で拡張する。第一に、日本のユニークな制度構造'
    '\u2014統一的な診療報酬点数表、都道府県の審査委員会\u2014を活用して行政的変動と臨床的変動を'
    '区別する。第二に、東アジアの国民皆保険制度における細かい地理的スケール（335地域）での'
    '麻酔手技変動の初の体系的分析を提供する。'
)

doc.add_paragraph(
    '我々が文書化した大学病院効果（全身麻酔でd=1.78）は、医療実施変動研究で通常報告される'
    '効果量と比較して顕著に大きい。これはおそらく日本の医学教育制度を反映しており、大学の'
    '医局（ikyoku）が地理的影響圏内の関連病院における臨床実施に実質的な影響力を及ぼしている。'
    '\u00b2\u2075\u00b7\u00b2\u2076'
)

doc.add_paragraph(
    '我々の所見は麻酔手技とがん転帰に関する進行中の議論に直接関連する。区域麻酔における'
    '局所麻酔薬投与量の変動に関する確定した問題と未確定の問題がRAPM誌で強調されており、\u00b2\u2077 '
    '本研究は変動のフレームワークを国全体の健康システムにおける手技選択のマクロレベルに拡張する。'
)

add_heading_rapm('腫瘍学的転帰への含意', level=2)

doc.add_paragraph(
    '全身麻酔\u2013硬膜外併用手技（L003）がCV 64.9%を示すという所見は、区域麻酔と全身麻酔の'
    '併用ががん手術における無再発生存期間を改善しうるというエビデンスに照らして臨床的に重要'
    'である。\u2077\u2013\u00b9\u00b2 大学病院圏と非大学病院圏の間のL003 SCRの1.73倍の差は、この潜在的に'
    '有益な手技へのアクセスが居住地に実質的に依存することを示唆しており、日本の国民皆保険制度の'
    '公平性の前提に疑問を呈する。\u00b9\u00b7\u00b2'
)

add_heading_rapm('2026年診療報酬改定', level=2)

doc.add_paragraph(
    '2026年の全身麻酔コードの名称変更（「マスク又は気管内挿管」から「声門上器具使用又は気管内'
    '挿管」）は、声門上器具の使用に関する差異的な査定判断につながった可能性のある以前のコーディング'
    'の曖昧さを反映している。査定による変動が定量的に小さいという我々の所見は、この改定が'
    'コーディングの明確化としては適切であるものの、我々が文書化した地理的パターンを大幅に変更'
    'する可能性は低いことを示唆する。'
)

add_heading_rapm('強みと限界', level=2)

doc.add_paragraph(
    '強みには、全国民をカバーする年齢・性別標準化比率の使用、細かい地理的スケール（335地域）での'
    '分析、複数の独立した感度分析、査定方針を制御する県内比較が含まれる。生態学的デザインは'
    '個人の同意要件を回避し、包括的な地理的カバレッジを可能にする。'
)

doc.add_paragraph(
    '限界には、地域レベル分析に固有の生態学的誤謬；\u00b2\u2078 個々の患者レベルの意思決定と集計'
    'パターンを区別できないこと；SCRデータの査定後の性質（意図された実施ではなく保険償還された'
    '実施を反映するが、我々の感度分析はこの区別が定量的に軽微であることを示唆する）；因果推論を'
    '排除する横断的デザイン；大学病院効果と都市集中の不完全な分離（大学病院は通常各県の最大'
    '都市に所在する）；コード別査定率データの欠如が含まれる。さらに、「防衛的過少請求」\u2014'
    '臨床家が査定されると予期する手技の請求を予防的に回避する\u2014は請求データから定量化'
    'できない。\u00b2\u2079'
)

add_heading_rapm('結論と政策的含意', level=2)

doc.add_paragraph(
    '日本における麻酔実施の地域差は真実かつ大きく、差異的な保険査定ではなく、制度的要因\u2014'
    '特に大学病院の近接性\u2014に主に起因する。大学病院圏と非大学病院圏の間の全身麻酔\u2013硬膜外'
    '併用の1.73倍の差は、腫瘍学的利益のエビデンスが蓄積する中で特に懸念される。\u2077\u2013\u00b9\u00b2 '
    '政策的対応として、対象を絞った教育プログラム、大学センターからの専門家アウトリーチ、品質'
    '指標としての麻酔手技パターンのモニタリングが考えられる。\u00b3\u2070 ここで開発した感度分析'
    'フレームワーク\u2014県内分散分解、コード間相関、定量的査定影響推定\u2014は、日本の国民皆保険制度'
    '下の任意の手技における実施変動の調査に適用可能である。'
)

# ═══════════════════════════════════════════════════════════
# 参考文献
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_rapm('参考文献', level=1)

refs = [
    '1. Ikegami N, Yoo B-K, Hashimoto H, et al. Japanese universal health coverage: evolution, '
    'achievements, and challenges. Lancet 2011;378:1106\u201315.',
    '2. Ikegami N. Universal Health Coverage for Inclusive and Sustainable Development: Lessons from '
    'Japan. Washington, DC: World Bank, 2014.',
    '3. Matsuda S, Fujimori K. The claim database in Japan. Asian Pac J Dis Manag 2012;6:55\u201359.',
    '4. Suto M, Iba A, Sugiyama T, et al. Literature review of studies using the National Database '
    'of Health Insurance Claims of Japan (NDB): limitations and strategies in using the NDB for '
    'research. JMA J 2023;7:10\u201320.',
    '5. Hamada H, Sekimoto M, Imanaka Y. Effects of the per diem prospective payment system with '
    'DRG-like grouping system (DPC/PDPS) on resource usage and healthcare quality in Japan. Health '
    'Policy 2012;107:194\u2013201.',
    '6. Social Insurance Medical Fee Payment Fund (Shiharai Kikin). Review Statistics, Fiscal Year '
    '2022. https://www.ssk.or.jp/tokeijoho/shinsatokei/index.html (accessed 15 Mar 2026).',
    '7. Pei L, Tan G, Wang L, et al. Anesthetic techniques for risk of malignant tumor recurrence. '
    'Cochrane Database Syst Rev 2019;(11):CD008877.',
    '8. Sessler DI, Pei L, Huang Y, et al. Recurrence of breast cancer after regional or general '
    'anesthesia: a randomized controlled trial. Lancet 2019;394:1807\u201315.',
    '9. Exadaktylos AK, Buggy DJ, Moriarty DC, et al. Can anesthetic technique for primary breast '
    'cancer surgery affect recurrence or metastasis? Anesthesiology 2006;105:660\u20134.',
    '10. Sun Y, Li T, Gan TJ. The effects of perioperative regional anesthesia and analgesia on cancer '
    'recurrence and survival after oncology surgery: a systematic review and meta-analysis. Reg '
    'Anesth Pain Med 2015;40:589\u201398.',
    '11. Xie S, Li L, Meng F, et al. Regional anesthesia might reduce recurrence and metastasis rates '
    'in adult patients with cancers after surgery: a meta-analysis. BMC Anesthesiol 2024;24:19.',
    '12. Zhang D, Jiang J, Liu J, et al. Effects of perioperative epidural analgesia on cancer '
    'recurrence and survival. Front Oncol 2021;11:798435.',
    '13. Ministry of Health, Labour and Welfare. Revision of medical fee schedule, 2026 '
    '(Reiwa 8 nendo shinryo houshu kaitei). https://www.mhlw.go.jp (accessed 15 Mar 2026).',
    '14. Cabinet Office. Regional variation visualization (chiikisa no mieruka). '
    'https://www5.cao.go.jp/keizai-shimon/kaigi/special/reform/mieruka/ (accessed 10 Mar 2026).',
    '15. Ministry of Health, Labour and Welfare. Survey of Physicians, Dentists and Pharmacists, '
    '2022. e-Stat. https://www.e-stat.go.jp (accessed 12 Mar 2026).',
    '16. Ministry of Land, Infrastructure, Transport and Tourism. National Land Numerical Information '
    'download service. https://nlftp.mlit.go.jp (accessed 10 Mar 2026).',
    '17. Cohen J. Statistical Power Analysis for the Behavioral Sciences. 2nd ed. '
    'Hillsdale, NJ: Lawrence Erlbaum, 1988.',
    '18. Julious SA, Nicholl J, George S. Why do we continue to use standardized mortality ratios for '
    'small area comparisons? J Public Health Med 2001;23:40\u20136.',
    '19. Plotly Technologies Inc. Plotly: collaborative data science. Montreal, QC, 2015. '
    'https://plotly.com (accessed 15 Mar 2026).',
    '20. Wennberg JE. Tracking Medicine: A Researcher\u2019s Quest to Understand Health Care. '
    'New York: Oxford University Press, 2010.',
    '21. Wennberg JE, Fisher ES, Skinner JS. Geography and the debate over Medicare reform. '
    'Health Aff (Millwood) 2002;Suppl Web Exclusives:W96\u2013114.',
    '22. NHS RightCare. The NHS Atlas of Variation in Healthcare, 2015. London: Public Health England.',
    '23. Wengler A, Nimptsch U, Mansky T. Hip and knee replacement surgery (arthroplasty) and '
    'geographic variations in Germany. Dtsch Arztebl Int 2014;111:407\u201316.',
    '24. Australian Commission on Safety and Quality in Health Care. Australian Atlas of Healthcare '
    'Variation. Sydney: ACSQHC, 2015.',
    '25. Otsuka T. The ikyoku system of university orthopedic surgery departments: an in-hospital '
    'organizational system unique to Japan. J Orthop Sci 2012;17:513\u201314.',
    '26. Onishi H. History of Japanese medical education. Korean J Med Educ 2018;30:283\u201394.',
    '27. Koscielniak-Nielsen ZJ, Helbo-Hansen HS. Settled science or unwarranted variation in local '
    'anesthetic dosing? Reg Anesth Pain Med 2019;44:998\u20131000.',
    '28. Greenland S. Ecologic versus individual-level sources of bias in ecologic estimates of '
    'contextual health effects. Int J Epidemiol 2001;30:1343\u201350.',
    '29. Sekimoto M, Imanaka Y, Kobayashi H, et al. Impact of hospital characteristics on the cost '
    'and practice patterns of acute myocardial infarction patients in Japan. Health Policy '
    '2006;78:34\u201345.',
    '30. Merry AF, Cooper JB, Soyannwo O, et al. International standards for a safe practice of '
    'anesthesia 2010. Can J Anesth 2010;57:1027\u201334.',
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.name = 'Arial'

# ═══════════════════════════════════════════════════════════
# 図版
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_rapm('図版', level=1)

figure_list = [
    ('/home/ubuntu/rapm_fig1_jp.png',
     '図1. 二次医療圏別麻酔標準化レセプト出現比の地理的分布。'
     '（A）全身麻酔（L008）SCR。（B）脊椎麻酔（L004）SCR。逆の空間パターンに注意：'
     '全身麻酔量の多い地域は脊椎麻酔量が少なく、その逆も然り（r=\u22120.506）。'
     '実線は都道府県境、点線は二次医療圏境。北方領土は白色（二次医療圏未設定）。'),
    ('/home/ubuntu/rapm_fig2_jp.png',
     '図2. 大学病院効果と合計SCR分析。'
     '（A）二次医療圏別大学病院所在：赤丸は1つ以上の大学病院が所在する二次医療圏（335中64）。'
     '（B）全身+脊椎麻酔合計（L008+L004）SCR：これらのコードを合計することで査定による'
     '再分類を中和するが、顕著な地域差は残存する（CV=43.4%）。'),
    ('/home/ubuntu/rapm_fig3_jp.png',
     '図3. 全身麻酔\u2013硬膜外併用の変動。'
     '（A）持続硬膜外注入（L003）SCR\u2014全身麻酔\u2013硬膜外併用の直接指標（CV=64.9%）。'
     '（B）L003/L008比（全身麻酔あたりの硬膜外併用率）、手術量で調整。値が高いほど'
     '全身麻酔\u2013硬膜外併用の使用が多いことを示す。'),
    ('/home/ubuntu/rapm_fig4_jp.png',
     '図4. 三次元押出コロプレスマップ。色はL003/L008比（硬膜外併用率）、高さは二次医療圏'
     'あたりの麻酔科医数を表す。硬膜外併用率と専門医数の両方が高い地域は、高く緑色の'
     'ポリゴンとして表示され、通常は都市部の大学病院圏に対応する。'),
]

for path, caption in figure_list:
    add_figure(path, caption, width=Inches(5.5))
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════
# STROBEチェックリスト
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_rapm('STROBEチェックリスト（横断研究）', level=1)

strobe_items = [
    ('1a', 'タイトルと抄録', '表題ページと構造化抄録'),
    ('1b', '情報量のある抄録', '必要な全要素を含む構造化抄録'),
    ('2', '背景/根拠', '緒言第1\u20133段落'),
    ('3', '目的', '緒言第5段落'),
    ('4', '研究デザイン', '方法：研究デザインとデータソース'),
    ('5', 'セッティング', '方法：335二次医療圏、2022年度データ'),
    ('6a', '参加者/適格性', '該当なし（集計データを用いた生態学的研究）'),
    ('7', '変数', '方法：分析した麻酔コード（表1）'),
    ('8', 'データソース', '方法：研究デザインとデータソース'),
    ('9', 'バイアス', '方法：感度分析；考察：限界'),
    ('10', '研究サイズ', '方法：335二次医療圏、47都道府県'),
    ('11', '定量的変数', '方法：統計解析'),
    ('12', '統計手法', '方法：統計解析'),
    ('13', '参加者', '結果：L008データのある334二次医療圏'),
    ('14', '記述的データ', '結果：表2'),
    ('15', 'アウトカムデータ', '結果：地域差セクション'),
    ('16', '主要結果', '結果：感度分析1\u20133、大学病院効果'),
    ('17', 'その他の分析', '結果：合計SCR、ペインクリニック波及効果'),
    ('18', '主要な所見', '考察：主要な所見'),
    ('19', '限界', '考察：強みと限界'),
    ('20', '解釈', '考察：既存文献との比較'),
    ('21', '一般化可能性', '考察：結論と政策的含意'),
    ('22', '資金', '［記入予定］'),
]

strobe_table = doc.add_table(rows=1, cols=3)
strobe_table.style = 'Table Grid'
strobe_hdr = strobe_table.rows[0].cells
for i, text in enumerate(['項目', 'STROBE推奨', '原稿中の位置']):
    strobe_hdr[i].paragraphs[0].add_run(text).bold = True
    strobe_hdr[i].paragraphs[0].runs[0].font.size = Pt(9)

for item in strobe_items:
    add_table_row(strobe_table, item)

# ═══════════════════════════════════════════════════════════
# 宣言事項
# ═══════════════════════════════════════════════════════════

doc.add_page_break()
add_heading_rapm('宣言事項', level=1)

add_heading_rapm('倫理承認', level=2)
doc.add_paragraph(
    '本研究は公開集計データを使用した。個人レベルのデータにはアクセスしていない。「人を対象とする'
    '生命科学・医学系研究に関する倫理指針」（日本、2021年改訂）に基づき、倫理委員会の承認は'
    '不要であった。'
)

add_heading_rapm('データの利用可能性', level=2)
doc.add_paragraph(
    '本研究で使用した全てのデータは公開されている。SCRデータ：内閣府地域差の見える化'
    '（https://www5.cao.go.jp/keizai-shimon/kaigi/special/reform/mieruka/）。'
    '医師統計：e-Stat（https://www.e-stat.go.jp）。GIS境界：国土数値情報'
    '（https://nlftp.mlit.go.jp）。分析コードは責任著者への要求により入手可能。'
)

add_heading_rapm('利益相反', level=2)
doc.add_paragraph('［著者が記入予定。］')

add_heading_rapm('資金', level=2)
doc.add_paragraph('［著者が記入予定。］')

add_heading_rapm('著者の貢献', level=2)
doc.add_paragraph('［著者が記入予定。］')

add_heading_rapm('AI使用に関する声明', level=2)
doc.add_paragraph(
    '［記入予定。RAPM/BMJの方針に従い、著者は本原稿の作成に使用したAI技術の詳細（使用した'
    'AI技術、使用理由、実行したタスク）を提供する必要がある。］'
)

add_heading_rapm('透明性に関する宣言', level=2)
doc.add_paragraph(
    '筆頭著者（本原稿の保証人）は、本原稿が報告されている研究の正直、正確、かつ透明な説明で'
    'あること、研究の重要な側面が省略されていないこと、および計画された研究からの逸脱が説明'
    'されていることを確認する。'
)

# ═══════════════════════════════════════════════════════════
# 保存
# ═══════════════════════════════════════════════════════════

output_path = '/home/ubuntu/regional_anesthesia_RAPM_JP.docx'
doc.save(output_path)
print(f"Saved: {output_path}")
